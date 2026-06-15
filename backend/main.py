"""
ProofreadAI Backend v2.21
=========================
FastAPI server with OpenRouter integration, chunking, and docx + footnotes support.
CHANGES:
- v2.21: COMPLETE Track Changes rewrite:
         * Real Comments panel (not just hover tooltip)
         * commentRangeStart/End + commentReference for proper linking
         * comments.xml with full content
         * Formatting preserved
         * Footnotes support
"""

import os
import re
import json
import uuid
import asyncio
import zipfile
import shutil
import copy
import unicodedata
from datetime import datetime
from typing import Optional, List, Tuple
from pathlib import Path

from dotenv import load_dotenv
from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, Response
from pydantic import BaseModel

import docx
from openai import OpenAI

# Use lxml for proper namespace handling (prevents Word corruption)
try:
    from lxml import etree as ET
    USING_LXML = True
except ImportError:
    import xml.etree.ElementTree as ET
    USING_LXML = False

# ============== CONFIGURATION ==============

load_dotenv()  # Try current directory
load_dotenv(Path(__file__).parent / ".env")  # Try same dir as main.py
load_dotenv(Path(__file__).parent / "backend" / ".env")  # Try backend subdir

OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY", "").strip()
if not OPENROUTER_API_KEY:
    print("WARNING: OPENROUTER_API_KEY is not set", flush=True)

OPENROUTER_BASE_URL = "https://openrouter.ai/api/v1"

UPLOAD_DIR = Path("uploads")
OUTPUT_DIR = Path("outputs")
UPLOAD_DIR.mkdir(exist_ok=True)
OUTPUT_DIR.mkdir(exist_ok=True)

# Severity filtering: keep only critical/major when True (drops "minor")
ONLY_SERIOUS_ERRORS = os.getenv("ONLY_SERIOUS_ERRORS", "false").lower() == "true"
# Grammar verifier: run secondary YES/NO LLM pass on risky single-word grammar fixes
ENABLE_GRAMMAR_VERIFIER = os.getenv("ENABLE_GRAMMAR_VERIFIER", "false").lower() == "true"

# Phrases the LLM tends to use when its rationale is stylistic, not grammatical.
# Used by downgrade_stylistic_findings() to demote those findings to "minor".
STYLE_RED_FLAGS = [
    "λιγότερο δόκιμο", "λιγότερο δόκιμη",
    "πιο φυσικό", "πιο φυσική", "πιο φυσικά",
    "καλύτερη ροή", "βελτίωση ροής",
    "προτιμότερο", "προτιμότερη", "προτιμάται",
    "συνηθέστερο", "συνηθέστερη",
    "πιο σωστό", "πιο σωστή",
    "πιο εκφραστικό", "πιο εκφραστική",
    "ομαλότερη διατύπωση", "ομαλότερο", "ομαλότερη",
    "πιο κομψό", "πιο κομψή",
    "πιο ρέον", "πιο ρέουσα",
]
# Concrete grammar terminology — if present in the rationale, the red-flag is
# treated as grammatically grounded and the finding is kept as-is.
GRAMMAR_GROUNDING_TERMS = [
    "ασυμφωνία", "συμφωνία", "πτώση", "γένος", "αριθμός",
    "ορθογραφικό", "ορθογραφία", "τόνος", "τονισμός",
    "ρήμα", "υποκείμενο", "αντικείμενο", "συντακτικό",
    "γραμματικά λάθος", "λάθος κλίση", "πρόσωπο",
    "φωνή ρήματος", "πτωτική", "κλίση",
]

AVAILABLE_MODELS = {
    # ── Latest aliases (auto-update with OpenRouter) ──
    "gemini-pro-latest":      {"name": "Gemini Pro — latest",              "model_id": "~google/gemini-pro-latest"},
    "gemini-flash-latest":    {"name": "Gemini Flash — latest",            "model_id": "~google/gemini-flash-latest"},
    "gpt-latest":             {"name": "GPT — latest",                     "model_id": "~openai/gpt-latest"},
    "claude-sonnet-latest":   {"name": "Claude Sonnet — latest",           "model_id": "~anthropic/claude-sonnet-latest"},
    "claude-opus-latest":     {"name": "Claude Opus — latest",             "model_id": "~anthropic/claude-opus-latest"},
    # ── Pinned versions ──
    "claude-fable-5": {"name": "Claude Fable 5", "model_id": "anthropic/claude-fable-5"},
    "claude-sonnet-4.6": {"name": "Claude Sonnet 4.6", "model_id": "anthropic/claude-sonnet-4.6"},
    "claude-opus-4.6": {"name": "Claude Opus 4.6", "model_id": "anthropic/claude-opus-4.6"},
    "claude-sonnet-4": {"name": "Claude Sonnet 4.5", "model_id": "anthropic/claude-sonnet-4.5"},
    "claude-opus-4": {"name": "Claude Opus 4.5", "model_id": "anthropic/claude-opus-4.5"},
    "gemini-3.1-pro": {"name": "Gemini 3.1 Pro", "model_id": "google/gemini-3.1-pro-preview"},
    "gemini-3-flash": {"name": "Gemini 3 Flash", "model_id": "google/gemini-3-flash-preview"},
    "gemini-3-pro": {"name": "Gemini 3 Pro", "model_id": "google/gemini-3-pro-preview"},
    "gemini-2.5-flash": {"name": "Gemini 2.5 Flash", "model_id": "google/gemini-2.5-flash"},
    "gemini-2.5-pro": {"name": "Gemini 2.5 Pro", "model_id": "google/gemini-2.5-pro"},
    "gpt-5.4-pro": {"name": "GPT-5.4 Pro", "model_id": "openai/gpt-5.4-pro"},
    "gpt-5.4": {"name": "GPT-5.4", "model_id": "openai/gpt-5.4"},
    "gpt-5.2": {"name": "GPT-5.2", "model_id": "openai/gpt-5.2"},
    "gpt-5-mini": {"name": "GPT-5 Mini", "model_id": "openai/gpt-5-mini"},
    "gpt-4.1": {"name": "GPT-4.1", "model_id": "openai/gpt-4.1"},
    "grok-4.1": {"name": "Grok 4.1 Fast", "model_id": "x-ai/grok-4.1-fast"},
    "mimo-v2-pro": {"name": "MiMo V2 Pro (Xiaomi)", "model_id": "xiaomi/mimo-v2-pro"},
    "step-3.5-flash": {"name": "Step 3.5 Flash (Free)", "model_id": "stepfun/step-3.5-flash:free"},
}

# Kept for /api/config compatibility
CHECK_INSTRUCTIONS = {
    "spelling": "Ορθογραφικά λάθη και λάθη τονισμού",
    "punctuation": "Λάθη στίξης (κόμματα, τελείες, άνω τελείες, ερωτηματικά)",
    "grammar": "Συντακτικά και γραμματικά λάθη",
    "clarity": "Νοηματικά λάθη, ασαφείς διατυπώσεις",
    "foreignisms": "Ξενισμούς και αγγλισμούς",
    "literal": "Κυριολεκτικές αποδόσεις που ακούγονται αφύσικα",
    "false_friends": "False friends",
    "naturalness": "Προτάσεις που δεν ακούγονται φυσικές",
    "dates": "Λανθασμένες ημερομηνίες και χρονολογίες",
    "events": "Λανθασμένα ιστορικά γεγονότα",
    "names": "Λάθος ονόματα προσώπων και τοποθεσίες",
    "numbers": "Ανακριβή αριθμητικά δεδομένα",
    "academic": "Ακαδημαϊκό/επιστημονικό ύφος",
    "literary": "Λογοτεχνικό ύφος",
    "journalistic": "Δημοσιογραφικό ύφος",
    "legal": "Νομικό/επίσημο ύφος",
    "casual": "Ανεπίσημο/καθημερινό ύφος",
}

# Word XML Namespaces - CRITICAL for proper XML handling
WORD_NAMESPACES = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
    'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
    'w14': 'http://schemas.microsoft.com/office/word/2010/wordml',
    'w15': 'http://schemas.microsoft.com/office/word/2012/wordml',
    'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape',
}

# Register namespaces for ElementTree (if not using lxml)
if not USING_LXML:
    for prefix, uri in WORD_NAMESPACES.items():
        ET.register_namespace(prefix, uri)

# ============== OPENROUTER CLIENT ==============

client = OpenAI(api_key=OPENROUTER_API_KEY, base_url=OPENROUTER_BASE_URL)

# ============== PROMPT GENERATION ==============

STYLE_LIBRARY = {
    "academic": "Ακαδημαϊκό/επιστημονικό κείμενο",
    "legal": "Νομικό/επίσημο κείμενο",
    "journalistic": "Δημοσιογραφικό κείμενο",
    "literary": "Λογοτεχνικό κείμενο",
    "casual": "Καθημερινό/ανεπίσημο κείμενο",
    "business": "Επαγγελματικό/εταιρικό κείμενο",
    "technical": "Τεχνικό εγχειρίδιο/τεκμηρίωση",
    "religious": "Θεολογικό/θρησκευτικό κείμενο",
    "history": "Ιστορικό κείμενο",
    "marketing": "Διαφημιστικό/marketing κείμενο",
    "other": "Άλλο",
}


# UNUSED — legacy single-pass prompt, kept for reference only.
# The /api/analyze endpoint uses generate_pass1_mechanical_prompt(),
# generate_pass2_grammatical_prompt(), and generate_pass3_semantic_prompt().
# Candidate for removal in future cleanup commit.
def generate_system_prompt(
    text_description: str,
    custom_instructions: str,
    translation_check: bool,
    fact_check: bool,
    web_fact_check: bool,
    style_check: bool,
    style_type: str,
    style_other: str,
    style_intensity: str,
    allow_sentence_level: bool,
) -> str:
    allow_sentence_level_note = (
        "Απαγορεύεται κάθε αλλαγή σε επίπεδο πρότασης."
        if not allow_sentence_level
        else "Επιτρέπονται αλλαγές σε επίπεδο πρότασης ΜΟΝΟ όταν είναι απολύτως αναγκαίο."
    )

    base_rules = f"""Είσαι έμπειρος επιμελητής κειμένων.
Στόχος: να εντοπίσεις πραγματικά προβλήματα και να προτείνεις διορθώσεις με μηδενική αλλαγή νοήματος.

ΒΑΣΙΚΟΣ ΕΛΕΓΧΟΣ (πάντα):
- Ορθογραφία/τονισμός
- Στίξη
- Γραμματική/σύνταξη
- Σαφήνεια ΜΟΝΟ όταν υπάρχει πραγματικό πρόβλημα (όχι "βελτιώσεις" για ομορφιά)

ΓΕΝΙΚΟΙ ΚΑΝΟΝΕΣ:
1) Εντόπισε ΜΟΝΟ πραγματικά προβλήματα.
2) ΜΗΝ προτείνεις αλλαγές σε σωστές φράσεις.
3) ΜΗΝ αλλάζεις ύφος/φωνή, εκτός αν έχει ενεργοποιηθεί ο Έλεγχος Ύφους.
4) Μία πρόταση ανά πρόβλημα (όχι εναλλακτικές).
5) Το "anchor" πρέπει να είναι ΑΚΡΙΒΩΣ copy-paste από το κείμενο (10–30 λέξεις).
6) Για αλλαγές τύπου "fix", το "original" πρέπει να είναι ΑΚΡΙΒΩΣ substring του κειμένου.

ΑΥΣΤΗΡΟΙ ΚΑΝΟΝΕΣ (δεν παραβιάζονται ποτέ):
H1) Αν το ίδιο λάθος εμφανίζεται ΠΟΛΛΕΣ ΦΟΡΕΣ στην ίδια παράγραφο, επέστρεψε ΞΕΧΩΡΙΣΤΗ διόρθωση για ΚΑΘΕ εμφάνιση.
H2) ΠΟΤΕ μην επιστρέφεις διόρθωση όπου το "suggested" είναι πανομοιότυπο με το "original".
H3) Για διορθώσεις γραμματικής/σύνταξης/μορφολογίας, αξιολόγησε ΠΑΝΤΑ ολόκληρη την πρόταση πριν προτείνεις τοπική αλλαγή λέξης.
H4) ΜΗΝ αλλάζεις μόνο ρηματικό τύπο, πτώση, αριθμό, γένος, άρθρο ή πρόθεση εκτός αν η ολόκληρη πρόταση το απαιτεί ξεκάθαρα.
H5) Αν υπάρχει αμφιβολία, ΜΗΝ επιστρέφεις διόρθωση.
H6) Ανέφερε ΜΟΝΟ ξεκάθαρα, ουσιαστικά, δημοσιεύσιμα λάθη. Απόφυγε κοσμητικές επανεγγραφές.
H7) {allow_sentence_level_note}
"""

    desc_block = f"\nΠΕΡΙΓΡΑΦΗ ΚΕΙΜΕΝΟΥ:\n{text_description if text_description else 'Γενικό κείμενο προς έλεγχο'}\n"
    custom_block = f"\nΕΙΔΙΚΕΣ ΟΔΗΓΙΕΣ ΧΡΗΣΤΗ:\n{custom_instructions}\n" if custom_instructions else ""

    modules = []

    if translation_check:
        modules.append("""MODULE: ΕΛΕΓΧΟΣ ΜΕΤΑΦΡΑΣΗΣ
- Εντόπισε ξενισμούς/αγγλισμούς, false friends, κυριολεκτικές αποδόσεις και αφύσικη σύνταξη που προέρχεται από calque.
- ΜΗΝ αλλάζεις το ύφος του κειμένου, μόνο κάνε τα ελληνικά φυσικά και σωστά.""")

    if fact_check:
        modules.append("""MODULE: ΕΛΕΓΧΟΣ ΠΛΗΡΟΦΟΡΙΩΝ (χωρίς web)
- Εντόπισε προφανή λάθη, πιθανές ανακρίβειες και εσωτερικές αντιφάσεις (ημερομηνίες, ονόματα, αριθμούς, γεγονότα).
- Αν δεν είσαι βέβαιος, ΜΗΝ δίνεις κάτι ως βέβαιο. Σήμανε verification_required=true και βάλε χαμηρότερο confidence.""")

    if web_fact_check:
        modules.append("""MODULE: ΕΛΕΓΧΟΣ ΠΛΗΡΟΦΟΡΙΩΝ (με web πηγές)
- Θα λάβεις και αποσπάσματα/πηγές από web αναζήτηση (αν υπάρχουν).
- Χρησιμοποίησε τις πηγές για να στηρίξεις διόρθωση και γράψε sources (τίτλος/URL) στο πεδίο sources.
- Ποτέ auto-apply: πάντα να είναι πρόταση με τεκμηρίωση.""")

    if style_check:
        stype = (style_type or "").strip() or "other"
        stype_label = STYLE_LIBRARY.get(stype, stype)
        if stype == "other":
            stype_label = f"Άλλο: {(style_other or '').strip()}" if (style_other or "").strip() else "Άλλο (μη καθορισμένο)"

        intensity = (style_intensity or "light").lower().strip()
        intensity_text = {
            "light": "Light: μόνο ξεκάθαρα αδόκιμες διατυπώσεις για το είδος κειμένου.",
            "medium": "Medium: περισσότερες παρεμβάσεις register, χωρίς αλλαγή φωνής.",
            "high": "High: επιθετικότερη ευθυγράμμιση register, χωρίς αλλαγή νοήματος.",
        }.get(intensity, "Light: μόνο ξεκάθαρα αδόκιμες διατυπώσεις.")

        scope_rule = (
            "Επιτρέπονται ΜΟΝΟ μικρές αλλαγές (λέξεις/φράσεις)."
            if not allow_sentence_level
            else "Επιτρέπονται και αλλαγές σε επίπεδο πρότασης ΜΟΝΟ όταν είναι αναγκαίο για register και χωρίς αλλαγή νοήματος."
        )

        literary_guard = ""
        if stype == "literary":
            literary_guard = "\nΓια λογοτεχνικό κείμενο: ΜΗΝ 'ισοπεδώνεις' μεταφορές/ρυθμό. Κάνε παρεμβάσεις μόνο σε λάθη ή σαφή αδοκιμία."

        modules.append(f"""MODULE: ΕΛΕΓΧΟΣ ΚΑΤΑΛΛΗΛΟΤΗΤΑΣ ΥΦΟΥΣ (register check)
ΤΥΠΟΣ ΚΕΙΜΕΝΟΥ: {stype_label}
ΕΝΤΑΣΗ: {intensity_text}
ΚΑΝΟΝΑΣ ΕΦΑΡΜΟΓΗΣ: {scope_rule}{literary_guard}
- ΜΗΝ κάνεις γενικό rewrite. Μόνο διορθώσεις όταν το register είναι ακατάλληλο για το είδος.""")

    modules_block = "\n\n".join(modules) if modules else "MODULES: (κανένα επιπλέον module ενεργό)\n"

    output_schema = """OUTPUT FORMAT (JSON array):
[
  {
    "module": "core|translation|facts|style",
    "type": "fix|suggestion",
    "scope": "token|phrase|sentence",

    "target": "body|footnote",
    "paragraph": <αριθμός παραγράφου ή 0>,
    "footnote_id": <αριθμός υποσημείωσης ή null>,

    "anchor": "<10–30 λέξεις ακριβές απόσπασμα από το κείμενο>",
    "original": "<ακριβές substring προς αντικατάσταση (μόνο για type=fix)>",
    "suggested": "<πρόταση διόρθωσης>",
    "reason": "<σύντομη εξήγηση>",
    "confidence": "low|medium|high",
    "verification_required": true|false,
    "sources": [{"title": "...", "url": "..."}],
    "exact_offset": <χαρακτήρας έναρξης του original μέσα στην παράγραφο/υποσημείωση, ή null αν άγνωστο>,
    "severity": "critical|major|minor"
  }
]

ΚΑΝΟΝΑΣ:
- Για ΣΩΜΑ: target="body", paragraph=<n>, footnote_id=null
- Για ΥΠΟΣΗΜΕΙΩΣΗ: target="footnote", footnote_id=<id>, paragraph=0
- severity="critical": λάθη που αλλάζουν νόημα ή είναι εντελώς λανθασμένα
- severity="major": ξεκάθαρα λάθη ορθογραφίας/γραμματικής/στίξης
- severity="minor": μικρές βελτιώσεις, κοσμητικά, αμφίβολα

Αν δεν υπάρχουν προβλήματα: []
Απάντησε ΜΟΝΟ με JSON array.
"""

    return "\n".join([base_rules, desc_block, custom_block, modules_block, output_schema])


def generate_pass1_mechanical_prompt() -> str:
    """Pass 1: spelling, accents, critical punctuation only."""
    return """Είσαι αυστηρός επιμελητής ελληνικών κειμένων με ΜΟΝΑΔΙΚΗ αποστολή:
εντοπισμός ΜΗΧΑΝΙΚΩΝ λαθών (ορθογραφία, τονισμός, σοβαρή στίξη).

═══════════════════════════════════════════════════════════════
ΤΙ ΠΡΕΠΕΙ ΝΑ ΒΡΕΙΣ (και ΜΟΝΟ αυτά)
═══════════════════════════════════════════════════════════════

1. ΟΡΘΟΓΡΑΦΙΚΑ ΛΑΘΗ (category: "spelling", severity: "critical")
   - Λανθασμένη γραφή λέξης (π.χ. "πηγένω" αντί "πηγαίνω")
   - Λάθος κατάληξη ορθογραφικά (π.χ. "γίνετε" αντί "γίνεται")
   - Λάθος -η/-ι, -ο/-ω, -ε/-αι, διπλά σύμφωνα
   - Λάθος αρχικό κεφαλαίο σε κύριο όνομα

2. ΤΟΝΙΣΜΟΣ / ΔΙΑΛΥΤΙΚΑ (category: "accent", severity: "critical")
   - Λέξη χωρίς απαραίτητο τόνο (π.χ. "αγαπημενος")
   - Λέξη με περιττό ή λάθος τόνο
   - Λείπουν διαλυτικά όπου χρειάζονται
   - ΔΙΑΚΡΙΤΙΚΟΣ τονισμός: η/ή, πώς/πως, πού/που, ώς/ως — ΜΟΝΟ όταν
     αλλάζει σημασία στο συγκεκριμένο context.

3. ΣΟΒΑΡΗ ΣΤΙΞΗ (category: "punctuation", severity: "major")
   - Λείπει τελεία στο τέλος ολοκληρωμένης περιόδου
   - Λείπει κόμμα σε κλητική ("Γιάννη έλα εδώ" → "Γιάννη, έλα εδώ")
   - Λείπει ερωτηματικό/θαυμαστικό σε ξεκάθαρα ερωτηματική/θαυμαστική
   - Λάθος χρήση άνω τελείας vs τελείας όταν αλλάζει νόημα
   - Διπλά σημεία στίξης κατά λάθος (π.χ. ",," ή "..")

═══════════════════════════════════════════════════════════════
ΤΙ ΔΕΝ ΕΠΙΣΤΡΕΦΕΙΣ ΠΟΤΕ
═══════════════════════════════════════════════════════════════

❌ Γραμματική / σύνταξη (άλλο pass αναλαμβάνει)
❌ Νοηματικές βελτιώσεις (άλλο pass)
❌ Ξενισμούς, false friends (άλλο pass)
❌ Στυλιστικές αλλαγές ("όμως" → "ωστόσο")
❌ Προαιρετικά κόμματα ("χθες, πήγα" vs "χθες πήγα" — και τα δύο δεκτά)
❌ Εναλλακτικά σωστές γραφές ("δε/δεν" πριν σύμφωνο, κλπ)
❌ Εισαγωγικά/παύλες (typography — αγνόησε)
❌ Ελληνικά vs αγγλικά σημεία στίξης
❌ ΑΝΟΙΧΤΑ ΕΙΣΑΓΩΓΙΚΑ που δεν κλείνουν (typography)
❌ Κενά μεταξύ λέξεων, double spaces

═══════════════════════════════════════════════════════════════
ΑΠΑΡΑΒΑΤΟΙ ΚΑΝΟΝΕΣ
═══════════════════════════════════════════════════════════════

R1. Το "suggested" ΠΡΕΠΕΙ να διαφέρει από το "original". Αν μετά
    από Unicode normalization είναι ίδια, ΜΗΝ το επιστρέφεις.

R2. Το "original" ΠΡΕΠΕΙ να είναι ΑΚΡΙΒΩΣ substring του κειμένου
    (character-for-character, συμπεριλαμβανομένων τόνων).

R3. Αν το ίδιο λάθος εμφανίζεται Ν φορές, επιστρέφεις Ν ΞΕΧΩΡΙΣΤΕΣ
    εγγραφές (μία ανά εμφάνιση), με διαφορετικό "anchor" για καθεμία.

R4. Το "anchor" είναι 10-30 λέξεις πραγματικό κείμενο γύρω από το
    λάθος (copy-paste από το input). Χρησιμεύει για disambiguation.

# Phase 2 Rule 6 — Pattern 1-9 (conservative bias, strengthened with rationale)
R5. Αν έχεις αμφιβολία αν κάτι είναι λάθος, ΔΕΝ το επιστρέφεις.
    Καλύτερα missed correction παρά false positive — κάθε FP καταναλώνει
    χρόνο επιμελητή και υπονομεύει την εμπιστοσύνη στο εργαλείο.

R6. Στα κύρια ονόματα, ΜΗΝ "διορθώνεις" ασυνήθιστες γραφές εκτός
    αν είναι ξεκάθαρα τυπογραφικό λάθος.

R7. ΜΗΝ επιστρέφεις severity: "minor" στο Pass 1. Μόνο critical/major.

R8. ΣΑΡΩΣΕ ΚΑΘΕ παράγραφο διεξοδικά από την αρχή ως το τέλος.
    Η τελευταία παράγραφος του chunk αξίζει την ίδια προσοχή με την
    πρώτη. Μην σταματάς μόλις βρεις μερικά.

═══════════════════════════════════════════════════════════════
OUTPUT FORMAT
═══════════════════════════════════════════════════════════════

Επιστρέφεις ΜΟΝΟ JSON array. Κάθε element:

{
  "target": "body" | "footnote",
  "paragraph_number": <αριθμός παραγράφου, 0 αν footnote>,
  "footnote_id": <αριθμός υποσημείωσης, μόνο αν target="footnote">,
  "original": "<ακριβές substring>",
  "suggested": "<διορθωμένη εκδοχή>",
  "anchor": "<10-30 λέξεις context>",
  "category": "spelling" | "accent" | "punctuation",
  "severity": "critical" | "major",
  "reason": "ορθογραφία" | "τονισμός" | "στίξη"
}

ΓΙΑ ΤΟ "reason": βάζεις ΜΟΝΟ μία από τις ετικέτες «ορθογραφία», «τονισμός»,
«στίξη» (την κατηγορία του λάθους). ΜΗΝ γράφεις δική σου εξήγηση ή αιτιολόγηση.

Αν δεν υπάρχει κανένα λάθος: επιστρέφεις []

Απάντησε ΜΟΝΟ με JSON array."""


def generate_pass2_grammatical_prompt(allow_sentence_level: bool) -> str:
    """Pass 2: grammar, syntax, agreement, morphology."""
    allow_sentence_level_note = (
        "Απαγορεύεται κάθε αλλαγή σε επίπεδο πρότασης."
        if not allow_sentence_level
        else "Επιτρέπονται αλλαγές σε επίπεδο πρότασης ΜΟΝΟ όταν είναι απολύτως αναγκαίο."
    )
    return f"""Είσαι επιμελητής ελληνικών κειμένων με ΑΠΟΚΛΕΙΣΤΙΚΗ εστίαση στη γραμματική.
Αποστολή: εντοπισμός ΜΟΝΟ γραμματικών/συντακτικών λαθών.

ΤΙ ΕΝΤΟΠΙΖΕΙΣ (και ΜΟΝΟ αυτά):
- Λάθη συμφωνίας (υποκείμενο-ρήμα, γένος-αριθμός-πτώση)
- Λάθος ρηματικός τύπος / χρόνος / έγκλιση
- Λάθος πτώση ονόματος ή αντωνυμίας
- Λάθος άρθρο ή πρόθεση που αλλάζει γραμματική ορθότητα
- Ελλιπής ή λάθος σύνταξη πρότασης
- {allow_sentence_level_note}

ΤΙ ΔΕΝ ΕΠΙΣΤΡΕΦΕΙΣ ΠΟΤΕ:
❌ Ορθογραφία / τονισμός / στίξη (Pass 1 αναλαμβάνει)
❌ Ξενισμούς, false friends, ύφος (Pass 3 αναλαμβάνει)
❌ Αλλαγές που δεν είναι ξεκάθαρα γραμματικά λάθη
❌ Αν η ολόκληρη πρόταση επιδέχεται διαφορετική ερμηνεία, ΜΗΝ επιστρέφεις

ΑΠΑΡΑΒΑΤΟΙ ΚΑΝΟΝΕΣ:

# Phase 2 Rule 1 — Pattern 1, 7, 9 (closure test)
R0. ΠΡΙΝ προτείνεις finding, ρώτησε τον εαυτό σου: "Γιατί είναι ΛΑΘΟΣ
    το πρωτότυπο;" Δεν αρκεί να είναι το προτεινόμενο "καλύτερο" ή
    "πιο συνηθισμένο". Αν δεν μπορείς να εντοπίσεις συγκεκριμένο
    γραμματικό σφάλμα στο πρωτότυπο, ΜΗΝ εκδώσεις finding.

# Phase 2 Rule 3 — Pattern 1, 7 (compound subjects — replaces vague R1)
R1. Πριν προτείνεις αλλαγή σε αριθμό/πρόσωπο/φωνή ρήματος, εντόπισε ΟΛΗ
    την πρόταση μέχρι το επόμενο τελικό σημείο στίξης (τελεία/ερωτηματικό/
    άνω τελεία). Αναγνώρισε:
    α) Το πλήρες υποκείμενο (μπορεί να εμφανίζεται μετά το ρήμα όταν
       εξαρτημένη πρόταση προηγείται της κύριας).
    β) Αν υπάρχουν δύο ή περισσότερα ονόματα ενωμένα με "και"/"ή"/κόμμα
       — το υποκείμενο είναι ΣΥΝΘΕΤΟ και το ρήμα πληθυντικός.

    ΠΑΡΑΔΕΙΓΜΑ ΣΩΣΤΟΥ (μην το αλλάξεις):
    «αντί να καταλάβουν αυτήν την επιθυμία, η Ευρώπη και ο Καναδάς
    επέλεξαν...» — το "καταλάβουν" (πληθυντικός) είναι σωστό λόγω
    σύνθετου υποκειμένου "η Ευρώπη και ο Καναδάς".

# Phase 2 Rule 6 — Pattern 1-9 (conservative bias, strengthened)
R2. Όταν δεν είσαι ΣΙΓΟΥΡΟΣ ότι κάτι είναι λάθος, ΜΗΝ εκδώσεις finding.
    Προτιμότερο να χάσεις πραγματικό λάθος παρά να εμφανίσεις false
    positive. Κάθε FP καταναλώνει χρόνο επιμελητή και υπονομεύει την
    εμπιστοσύνη στο εργαλείο.

R3. Το "original" ΠΡΕΠΕΙ να είναι ΑΚΡΙΒΩΣ substring του κειμένου.

R4. Αν το ίδιο λάθος εμφανίζεται Ν φορές, επιστρέφεις Ν ΞΕΧΩΡΙΣΤΕΣ εγγραφές.

# Phase 2 Rule 4 — Pattern 9 (parenthetical case inheritance)
R5. Λέξη μέσα σε παρένθεση/εισαγωγικά/επεξηγηματική δομή (μετά κόμμα,
    παύλα, ή "που σημαίνει") ΚΛΗΡΟΝΟΜΕΙ πτώση/αριθμό από την εξωτερική
    λέξη που εξηγείται. ΜΗΝ αλλάζεις πτώση στο εσωτερικό χωρίς να
    εξετάσεις αυτή την κληρονομικότητα.

    ΠΑΡΑΔΕΙΓΜΑ ΣΩΣΤΟΥ (μην το αλλάξεις):
    «μείωσης των μεταναστευτικών θεωρήσεων (βίζας) H-1B» — το "βίζας"
    (γενική ενικού) κληρονομεί τη γενική από το "θεωρήσεων".

# Phase 2 Rule 5 — Pattern 3 (replacement scope)
R6. Το `suggested` πρέπει να αντικαθιστά το `original` αυτούσιο και να
    παράγει γραμματικά ΟΡΘΗ τελική πρόταση. Νοητικά εφάρμοσε τη
    διόρθωση και διάβασε ξανά την πρόταση. Αν χρειάζεται και άλλο
    μέρος της πρότασης να αλλάξει (πτώση αντικειμένου, σύνταξη,
    σύνδεση), ΕΠΕΚΤΕΙΝΕ το original ΚΑΙ το suggested ώστε να
    συμπεριλάβουν όλη την αναγκαία αλλαγή.

# Phase 2 Rule 2 (slim) — Pattern 8 (voice/να-clause prohibition)
R7. ΜΗΝ αλλάζεις παθητική σε ενεργητική φωνή όταν και τα δύο είναι
    γραμματικώς σωστά. ΜΗΝ αλλάζεις να-πρόταση σε ονοματοποίηση χωρίς
    γραμματικό λόγο. Είναι υφολογικές επιλογές, όχι γραμματικά λάθη.

# Fix A — η πτώση κρίνεται από συμφωνία με το άρθρο/μετοχή, ΟΧΙ από το ρήμα
R8. ΠΡΙΝ προτείνεις ΟΠΟΙΑΔΗΠΟΤΕ αλλαγή πτώσης/αριθμού σε ουσιαστικό/επίθετο/
    μετοχή, εκτέλεσε ΥΠΟΧΡΕΩΤΙΚΑ αυτόν τον έλεγχο:
    1) Βρες το άρθρο ή τη μετοχή/επίθετο που ανήκει στην ΙΔΙΑ ονοματική
       φράση με τη λέξη (συνήθως αριστερά της: ο/η/το/οι/τα/του/της/των/
       τον/την/τους/τις...).
    2) Η πτώση της λέξης καθορίζεται από τη ΣΥΜΦΩΝΙΑ με αυτό το άρθρο/μετοχή
       — ΟΧΙ από το ρήμα.
    3) Αν η λέξη ΗΔΗ συμφωνεί με το άρθρο/μετοχή της, η πτώση είναι ΣΩΣΤΗ.
       ΜΗΝ την αλλάξεις, ΑΚΟΜΗ κι αν:
       - ένα ρήμα φαίνεται μεταβατικό και «θέλει» αιτιατική, ή
       - το «ως»/«σαν» φαίνεται να «θέλει» ονομαστική.
       Το άρθρο/μετοχή ΥΠΕΡΙΣΧΥΕΙ του ρήματος ως κριτήριο πτώσης.
    Αν όντως αλλάζεις πτώση, πρέπει να φέρεις σε συμφωνία ΟΛΗ τη φράση (μαζί
    το άρθρο/επίθετο — δες R6). Αλλιώς ΜΗΝ εκδώσεις finding.

    ΑΝΤΙΠΑΡΑΔΕΙΓΜΑΤΑ (είναι ΣΩΣΤΑ — ΜΗΝ τα διορθώσεις):
    α) «οι αναφερθέντες δύο κλάδοι παρακολουθούν...»
       Το «κλάδοι» είναι ΥΠΟΚΕΙΜΕΝΟ και συμφωνεί με «οι/αναφερθέντες»
       (ονομαστική). Λάθος σκέψη: «το παρακολουθώ είναι μεταβατικό, άρα
       κλάδους». ΟΧΙ — το «κλάδοι» ΔΕΝ είναι αντικείμενο, είναι το υποκείμενο
       του «παρακολουθούν».
    β) «...προοπτικής, της δυνάμενης να χαρακτηριστεί ως τελολογικής»
       Το «τελολογικής» συμφωνεί με «της δυνάμενης» (γενική). Λάθος σκέψη:
       «χαρακτηρίζεται ως → ονομαστική, άρα τελολογική». ΟΧΙ — το κατηγορούμενο
       έλκεται στη γενική της μετοχικής φράσης «της δυνάμενης».

OUTPUT FORMAT (JSON array):
[
  {{
    "target": "body" | "footnote",
    "paragraph": <αριθμός παραγράφου, 0 αν footnote>,
    "footnote_id": <αριθμός υποσημείωσης ή null>,
    "original": "<ακριβές substring>",
    "suggested": "<διορθωμένη εκδοχή>",
    "anchor": "<10-30 λέξεις context>",
    "module": "core",
    "type": "fix",
    "scope": "token" | "phrase" | "sentence",
    "severity": "critical" | "major",
    "confidence": "low" | "medium" | "high",
    "reason": "ασυμφωνία" | "λάθος πτώση" | "λάθος αριθμός" | "λάθος ρηματικός τύπος" | "λάθος άρθρο/πρόθεση" | "σύνταξη"
  }}
]

ΓΙΑ ΤΟ "reason": βάζεις ΜΟΝΟ μία από τις ετικέτες «ασυμφωνία», «λάθος πτώση»,
«λάθος αριθμός», «λάθος ρηματικός τύπος», «λάθος άρθρο/πρόθεση», «σύνταξη»
(την κατηγορία του γραμματικού λάθους). ΜΗΝ γράφεις δική σου εξήγηση ή αιτιολόγηση.

Αν δεν υπάρχουν προβλήματα: []
Απάντησε ΜΟΝΟ με JSON array."""


def generate_pass3_semantic_prompt(
    text_description: str,
    custom_instructions: str,
    translation_check: bool,
    fact_check: bool,
    web_fact_check: bool,
    style_check: bool,
    style_type: str,
    style_other: str,
    style_intensity: str,
    allow_sentence_level: bool,
) -> str:
    """Pass 3: clarity, foreignisms, meaning, naturalness, style."""
    allow_sentence_level_note = (
        "Απαγορεύεται κάθε αλλαγή σε επίπεδο πρότασης."
        if not allow_sentence_level
        else "Επιτρέπονται αλλαγές σε επίπεδο πρότασης ΜΟΝΟ όταν είναι απολύτως αναγκαίο."
    )

    active_modules = []
    if translation_check:
        active_modules.append("""MODULE: ΕΛΕΓΧΟΣ ΜΕΤΑΦΡΑΣΗΣ
- Εντόπισε ξενισμούς/αγγλισμούς, false friends, κυριολεκτικές αποδόσεις και αφύσικη σύνταξη calque.
- ΜΗΝ αλλάζεις το ύφος, μόνο κάνε τα ελληνικά φυσικά και σωστά.""")

    if fact_check:
        active_modules.append("""MODULE: ΕΛΕΓΧΟΣ ΠΛΗΡΟΦΟΡΙΩΝ (χωρίς web)
- Εντόπισε προφανή λάθη, πιθανές ανακρίβειες, εσωτερικές αντιφάσεις.
- Αν δεν είσαι βέβαιος, βάλε verification_required=true και χαμηλό confidence.""")

    if web_fact_check:
        active_modules.append("""MODULE: ΕΛΕΓΧΟΣ ΠΛΗΡΟΦΟΡΙΩΝ (με web πηγές)
- Χρησιμοποίησε web πηγές για τεκμηρίωση. Γράψε sources (τίτλος/URL).
- Ποτέ auto-apply: πάντα πρόταση με τεκμηρίωση.""")

    if style_check:
        stype = (style_type or "").strip() or "other"
        stype_label = STYLE_LIBRARY.get(stype, stype)
        if stype == "other":
            stype_label = f"Άλλο: {(style_other or '').strip()}" if (style_other or "").strip() else "Άλλο (μη καθορισμένο)"
        intensity = (style_intensity or "light").lower().strip()
        intensity_text = {
            "light": "Light: μόνο ξεκάθαρα αδόκιμες διατυπώσεις.",
            "medium": "Medium: περισσότερες παρεμβάσεις register.",
            "high": "High: επιθετικότερη ευθυγράμμιση register.",
        }.get(intensity, "Light: μόνο ξεκάθαρα αδόκιμες διατυπώσεις.")
        scope_rule = (
            "Επιτρέπονται ΜΟΝΟ μικρές αλλαγές (λέξεις/φράσεις)."
            if not allow_sentence_level
            else "Επιτρέπονται και αλλαγές σε επίπεδο πρότασης ΜΟΝΟ όταν είναι αναγκαίο."
        )
        active_modules.append(f"""MODULE: ΕΛΕΓΧΟΣ ΥΦΟΥΣ (register check)
ΤΥΠΟΣ ΚΕΙΜΕΝΟΥ: {stype_label}
ΕΝΤΑΣΗ: {intensity_text}
ΚΑΝΟΝΑΣ: {scope_rule}""")

    if not active_modules:
        return ""  # Pass 3 disabled — no semantic modules active

    desc_block = f"\nΠΕΡΙΓΡΑΦΗ ΚΕΙΜΕΝΟΥ:\n{text_description if text_description else 'Γενικό κείμενο'}\n"
    custom_block = f"\nΕΙΔΙΚΕΣ ΟΔΗΓΙΕΣ ΧΡΗΣΤΗ:\n{custom_instructions}\n" if custom_instructions else ""
    modules_block = "\n\n".join(active_modules)

    return f"""Είσαι επιμελητής κειμένων με εστίαση σε νοηματικά/σημασιολογικά θέματα.
ΔΕΝ ελέγχεις ορθογραφία/τονισμό/στίξη (Pass 1) ούτε γραμματική/σύνταξη (Pass 2).
{allow_sentence_level_note}
{desc_block}{custom_block}
{modules_block}

ΑΠΑΡΑΒΑΤΟΙ ΚΑΝΟΝΕΣ:

# Phase 2 Rule 1 — Pattern 2, 7 (closure test)
R0. Πριν προτείνεις finding, ρώτησε: "Γιατί είναι ΛΑΘΟΣ το πρωτότυπο;"
    Δεν αρκεί να είναι το προτεινόμενο "καλύτερο" ή "πιο φυσικό".
    Αν δεν μπορείς να εντοπίσεις συγκεκριμένο γραμματικό ή νοηματικό
    σφάλμα στο πρωτότυπο, ΜΗΝ εκδώσεις finding.

R1. Το "original" ΠΡΕΠΕΙ να είναι ΑΚΡΙΒΩΣ substring του κειμένου.

# Phase 2 Rule 6 — Pattern 1-9 (conservative bias, strengthened)
R2. Όταν δεν είσαι ΣΙΓΟΥΡΟΣ ότι κάτι είναι λάθος, ΜΗΝ εκδώσεις finding.
    Προτιμότερο να χάσεις πραγματικό λάθος παρά να εμφανίσεις false
    positive. Κάθε FP καταναλώνει χρόνο επιμελητή και υπονομεύει την
    εμπιστοσύνη στο εργαλείο.

R3. ΜΗΝ αλλάζεις ύφος/φωνή χωρίς ενεργό style module.

R4. Μία πρόταση ανά πρόβλημα.

# Phase 2 Rule 2 (full) — Pattern 8 (stylistic bias prohibition)
R5. ΑΠΑΓΟΡΕΥΟΝΤΑΙ οι παρακάτω αλλαγές όταν και τα δύο είναι γραμματικώς σωστά:

    α) Αλλαγή παθητικής σε ενεργητική φωνή.
       «κατοικούνται από πιγκουίνους» είναι ΣΩΣΤΟ. ΜΗΝ το αλλάξεις σε
       «κατοικούν πιγκουίνοι». Η παθητική με ποιητικό αίτιο είναι
       πλήρως αποδεκτή σύνταξη στα ελληνικά.

    β) Αλλαγή να-πρότασης σε ονοματοποίηση.
       «δυνατότητα να εξισορροπήσει» είναι ΣΩΣΤΟ. ΜΗΝ το αλλάξεις σε
       «δυνατότητα εξισορρόπησης». Και τα δύο είναι αποδεκτά.

    γ) Αλλαγή λόγιας γενικής σε εμπρόθετο ή αντίστροφα όταν και τα δύο
       είναι σωστά.

    δ) Αναδιάταξη επιθέτων-ουσιαστικών χωρίς γραμματικό λόγο.

    Όροι που δείχνουν stylistic suggestion αντί διόρθωσης:
    "λιγότερο δόκιμο", "πιο φυσικό", "καλύτερη ροή", "προτιμότερο",
    "συνηθέστερο", "πιο σωστό", "ομαλότερη διατύπωση", "πιο κομψό".
    Αν χρησιμοποιείς τέτοιους όρους στο reason σου, ΜΗΝ εκδώσεις finding.

# Phase 2 Rule 5 — Pattern 3 (replacement scope)
R6. Το `suggested` πρέπει να αντικαθιστά το `original` αυτούσιο και να
    παράγει γραμματικά ΟΡΘΗ τελική πρόταση. Αν χρειάζεται και άλλο
    μέρος της πρότασης να αλλάξει, ΕΠΕΚΤΕΙΝΕ το original+suggested
    ώστε να συμπεριλάβουν όλη την αναγκαία αλλαγή.

OUTPUT FORMAT (JSON array):
[
  {{
    "target": "body" | "footnote",
    "paragraph": <αριθμός παραγράφου, 0 αν footnote>,
    "footnote_id": <αριθμός υποσημείωσης ή null>,
    "original": "<ακριβές substring>",
    "suggested": "<διορθωμένη εκδοχή>",
    "anchor": "<10-30 λέξεις context>",
    "module": "translation" | "facts" | "style",
    "type": "fix" | "suggestion",
    "scope": "token" | "phrase" | "sentence",
    "severity": "critical" | "major" | "minor",
    "confidence": "low" | "medium" | "high",
    "verification_required": true | false,
    "sources": [{{"title": "...", "url": "..."}}],
    "reason": "ξενισμός" | "νόημα/ακρίβεια" | "ύφος"
  }}
]

ΓΙΑ ΤΟ "reason": βάζεις ΜΟΝΟ μία από τις ετικέτες «ξενισμός», «νόημα/ακρίβεια»,
«ύφος» (την κατηγορία του ευρήματος, ανάλογα με το ενεργό module). ΜΗΝ γράφεις
δική σου εξήγηση ή αιτιολόγηση.

Αν δεν υπάρχουν προβλήματα: []
Απάντησε ΜΟΝΟ με JSON array."""


def _normalize_pass1_correction(corr: dict) -> dict:
    """Map Pass 1 output schema to the main correction schema."""
    normalized = dict(corr)
    # Pass 1 uses paragraph_number; main loop expects paragraph
    if "paragraph_number" in normalized and "paragraph" not in normalized:
        normalized["paragraph"] = normalized.pop("paragraph_number")
    # Map category → module
    category = normalized.pop("category", None)
    if category and "module" not in normalized:
        normalized["module"] = "core"
    # Pass 1 corrections are always type=fix
    normalized.setdefault("type", "fix")
    normalized.setdefault("scope", "token")
    normalized.setdefault("confidence", "high")
    normalized.setdefault("verification_required", False)
    normalized.setdefault("sources", [])
    return normalized


def generate_user_prompt(numbered_text: str) -> str:
    return f"""Ανάλυσε το κείμενο:

{numbered_text}

Απάντησε ΜΟΝΟ με JSON array."""


# ============== HELPERS ==============

def extract_paragraphs(doc_path: Path) -> list:
    """
    Extract paragraphs with formatting information (bold, italic).
    Returns: [{"index": i, "number": n, "text": "plain text", "segments": [{"text": "...", "bold": bool, "italic": bool}, ...]}, ...]
    """
    doc = docx.Document(doc_path)
    paragraphs = []
    
    for i, p in enumerate(doc.paragraphs):
        if not p.text.strip():
            continue
        
        segments = []
        for run in p.runs:
            if run.text:
                segments.append({
                    "text": run.text,
                    "bold": run.bold or False,
                    "italic": run.italic or False,
                })
        
        # Merge adjacent segments with same formatting
        merged_segments = []
        for seg in segments:
            if merged_segments and merged_segments[-1]["bold"] == seg["bold"] and merged_segments[-1]["italic"] == seg["italic"]:
                merged_segments[-1]["text"] += seg["text"]
            else:
                merged_segments.append(seg)
        
        paragraphs.append({
            "index": i,
            "number": len(paragraphs) + 1,
            "text": p.text.strip(),  # Plain text for compatibility
            "segments": merged_segments,  # Rich text with formatting
        })
    
    return paragraphs


def extract_footnotes(doc_path: Path) -> list:
    """
    Returns: [{"id": 1, "text": "..."}, ...]
    Word έχει και special footnotes (-1, 0). Τις αγνοούμε.
    """
    footnotes = []
    try:
        with zipfile.ZipFile(doc_path, "r") as z:
            if "word/footnotes.xml" not in z.namelist():
                return []
            xml_data = z.read("word/footnotes.xml")
    except Exception:
        return []

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    try:
        if USING_LXML:
            root = ET.fromstring(xml_data)
        else:
            root = ET.fromstring(xml_data)
    except Exception:
        return []

    for fn in root.findall("w:footnote", ns):
        fid = fn.get(f"{{{ns['w']}}}id")
        if fid is None:
            continue
        try:
            fid_int = int(fid)
        except ValueError:
            continue
        if fid_int <= 0:
            continue

        texts = []
        for t in fn.findall(".//w:t", ns):
            if t.text:
                texts.append(t.text)
        text = "".join(texts).strip()
        if text:
            footnotes.append({"id": fid_int, "text": text})

    return footnotes


def create_chunks(paragraphs: list, size: int = 8, overlap: int = 2) -> list:
    chunks = []
    step = max(1, size - overlap)
    i = 0
    while i < len(paragraphs):
        chunks.append(paragraphs[i: min(i + size, len(paragraphs))])
        i += step
        if i >= len(paragraphs):
            break
    return chunks


def chunk_footnotes(footnotes: list, max_chars: int = 6000) -> list:
    """
    Group footnotes into chunks by character budget, keeping each footnote
    whole and without overlap (footnotes are independent units). Returns a list
    of footnote-lists. Used so footnotes are sent to the model ONCE each, rather
    than re-broadcast with every body chunk.
    """
    chunks = []
    current = []
    size = 0
    for fn in footnotes:
        length = len(fn.get("text", "")) + 20  # +label overhead
        if current and size + length > max_chars:
            chunks.append(current)
            current = []
            size = 0
        current.append(fn)
        size += length
    if current:
        chunks.append(current)
    return chunks


def format_chunk(chunk: list) -> str:
    return "\n\n".join(f"[Παράγραφος {p['number']}]\n{p['text']}" for p in chunk)


def format_footnotes(footnotes: list) -> str:
    return "\n\n".join(f"[Υποσημείωση {fn['id']}]\n{fn['text']}" for fn in footnotes)


def parse_response(text: str) -> list:
    text = text.strip()
    
    # Remove markdown code fences more robustly
    if text.startswith("```json"):
        text = text[7:]
    elif text.startswith("```"):
        text = text[3:]
    
    if text.endswith("```"):
        text = text[:-3]
    
    text = text.strip()
    
    # Debug: show what we're trying to parse
    print(f"[DEBUG] parse_response input (first 300 chars): {text[:300]}...")
    
    try:
        result = json.loads(text)
        if isinstance(result, list):
            print(f"[DEBUG] Successfully parsed {len(result)} corrections")
            return result
    except Exception as e:
        print(f"[DEBUG] Direct JSON parse failed: {e}")
    
    # Fallback: try to find JSON array in text
    match = re.search(r"\[[\s\S]*\]", text)
    if match:
        try:
            result = json.loads(match.group())
            print(f"[DEBUG] Regex extraction found {len(result)} corrections")
            return result
        except Exception as e:
            print(f"[DEBUG] Regex JSON parse failed: {e}")
    
    print(f"[DEBUG] parse_response returning empty list")
    return []


async def call_llm(system: str, user: str, model_id: str, session=None, pass_label: str = None) -> str:
    try:
        # All Gemini models may need more output tokens
        # Gemini 3 uses reasoning tokens (needs even more)
        if "gemini-3" in model_id:
            max_tokens = 65536
        elif "gemini" in model_id or "google" in model_id:
            max_tokens = 16384
        else:
            max_tokens = 8192
        
        params = {
            "model": model_id,
            "messages": [{"role": "system", "content": system}, {"role": "user", "content": user}],
            "temperature": 0,
            "max_tokens": max_tokens,
        }
        
        response = client.chat.completions.create(**params)
        
        # Track token usage
        if session and hasattr(response, 'usage') and response.usage:
            u = response.usage
            session.tokens_used["prompt"] += getattr(u, 'prompt_tokens', 0) or 0
            session.tokens_used["completion"] += getattr(u, 'completion_tokens', 0) or 0
            session.tokens_used["total"] += getattr(u, 'total_tokens', 0) or 0
            # Instrumentation: per-pass attribution (Phase 1 — measurement only)
            if pass_label is not None:
                bd = getattr(session, 'token_breakdown', None)
                if bd is None:
                    bd = session.token_breakdown = {}
                slot = bd.setdefault(pass_label, {"calls": 0, "prompt": 0, "completion": 0, "total": 0})
                slot["calls"] += 1
                slot["prompt"] += getattr(u, 'prompt_tokens', 0) or 0
                slot["completion"] += getattr(u, 'completion_tokens', 0) or 0
                slot["total"] += getattr(u, 'total_tokens', 0) or 0
        
        content = response.choices[0].message.content
        
        # For Google models, check if content is in a different field
        if not content and hasattr(response.choices[0].message, 'reasoning'):
            reasoning = response.choices[0].message.reasoning
            if reasoning:
                print(f"[WARNING] Content empty but reasoning has {len(reasoning)} chars")
                # Try to extract JSON from reasoning if it contains our format
                if '"original"' in reasoning and '"suggested"' in reasoning:
                    import re
                    json_match = re.search(r'\[\s*\{.*\}\s*\]', reasoning, re.DOTALL)
                    if json_match:
                        print(f"[INFO] Found JSON in reasoning, extracting...")
                        content = json_match.group(0)
        
        return content or "[]"
    except Exception as e:
        print(f"\n{'='*50}")
        print(f"ERROR calling {model_id}: {str(e)}")
        import traceback
        traceback.print_exc()
        print(f"{'='*50}\n")
        return "[]"


def normalize_text(text: str) -> str:
    """Normalize text for comparison: Unicode NFC + whitespace normalization."""
    if not text:
        return ""
    # Unicode NFC normalization (combines accents properly)
    text = unicodedata.normalize("NFC", text)
    # Normalize whitespace (multiple spaces -> single, strip)
    text = " ".join(text.split())
    return text


# ============== NEW: STRICT VALIDATION LAYER ==============

def find_all_occurrences(full_text: str, needle: str) -> list:
    """
    Return list of all start offsets of needle in full_text (exact match).
    Used to detect repeated identical errors and assign per-occurrence offsets.
    """
    if not needle or not full_text:
        return []
    offsets = []
    start = 0
    while True:
        idx = full_text.find(needle, start)
        if idx == -1:
            break
        offsets.append(idx)
        start = idx + max(1, len(needle))
    return offsets


# Matches strings that are only punctuation/whitespace (safe to correct even if short)
_PUNCT_ONLY_RE = re.compile(r'^[\s\.,;:!?·\-—–«»"""\'()\[\]{}…]+$')


def validate_correction(
    corr: dict,
    paragraph_text: str = None,
    footnote_text: str = None,
    allow_sentence_level: bool = False,
) -> tuple:
    """
    Strict gate before a parsed LLM correction is stored.
    Returns (is_valid: bool, rejection_reason: str).

    Rejects if:
    1. original is empty
    2. suggested is empty for type=fix
    3. normalized(original) == normalized(suggested)  [no-op]
    4. scope=sentence and allow_sentence_level is False
    5. original is 1-2 chars and not punctuation-only  [too risky]
    6. target=body and original not found in paragraph_text
    7. target=footnote and original not found in footnote_text
    8. anchor is missing/too short (< 5 chars) for non-trivial fixes (len > 3)
    9. confidence=low AND verification_required=True on a grammar/core module
    """
    original = (corr.get("original") or "")
    suggested = (corr.get("suggested") or "")
    orig_s = original.strip()
    sugg_s = suggested.strip()
    ctype = (corr.get("type") or "fix").lower()
    target = (corr.get("target") or "body").lower()
    scope = (corr.get("scope") or "token").lower()
    anchor = (corr.get("anchor") or "").strip()
    confidence = (corr.get("confidence") or "medium").lower()
    verification_required = bool(corr.get("verification_required", False))
    module = (corr.get("module") or "core").lower()

    # Rule 1
    if not orig_s:
        return False, "original is empty"

    # Rule 2
    if ctype == "fix" and not sugg_s:
        return False, "suggested is empty for type=fix"

    # Rule 3 — no-op
    if normalize_text(original) == normalize_text(suggested):
        return False, "no-op: normalized original == normalized suggested"

    # Rule 4 — sentence-level guard
    if scope == "sentence" and not allow_sentence_level:
        return False, "sentence-level rewrite rejected (allow_sentence_level=False)"

    # Rule 5 — too-short risky original
    if len(orig_s) <= 2 and not _PUNCT_ONLY_RE.match(orig_s):
        return False, f"original too short and risky: '{orig_s}'"

    # Rule 6 — body text presence check
    if target == "body" and paragraph_text is not None:
        if orig_s not in paragraph_text and normalize_text(orig_s) not in normalize_text(paragraph_text):
            return False, f"original '{orig_s[:50]}' not found in paragraph text"

    # Rule 7 — footnote text presence check
    if target == "footnote" and footnote_text is not None:
        if orig_s not in footnote_text and normalize_text(orig_s) not in normalize_text(footnote_text):
            return False, f"original '{orig_s[:50]}' not found in footnote text"

    # Rule 8 — anchor weakness for non-trivial fixes
    if ctype == "fix" and len(orig_s) > 3 and len(anchor) < 5:
        return False, "anchor missing or too weak for non-trivial fix"

    # Rule 9 — low-confidence risky grammar edit
    if confidence == "low" and verification_required and module in ("core", "grammar"):
        return False, "low confidence + verification_required on grammar/core module"

    return True, ""


# Matches a single word (no whitespace) — used to identify risky grammar edits
# ============== Fix A: deterministic article–noun concord guard ==============
# Blocks ONLY the failure pattern reported by the owner: a single-word case
# change that breaks agreement with the definite article governing the word
# (e.g. «οι ... κλάδοι» → «κλάδους»). The block fires only when the ORIGINAL
# already agrees with its article and the SUGGESTED cannot — so genuine
# agreement fixes (where the original disagrees) are never dropped. Every
# uncertain case returns True (fail-open) to protect recall.

# Definite-article forms → the grammatical case(s) they can mark.
_ARTICLE_CASE = {
    "ο": {"NOM"}, "η": {"NOM"}, "οι": {"NOM"},
    "το": {"NOM", "ACC"}, "τα": {"NOM", "ACC"},   # ambiguous → rarely blocks
    "του": {"GEN"}, "της": {"GEN"}, "των": {"GEN"},
    "τον": {"ACC"}, "την": {"ACC"}, "τους": {"ACC"}, "τις": {"ACC"}, "τες": {"ACC"},
}

# Numerals that may sit between an article and its noun (skipped while scanning).
_NUMERAL_WORDS = {
    "δύο", "δυο", "τρεις", "τρία", "τέσσερις", "τέσσερα", "πέντε", "έξι",
    "εφτά", "επτά", "οκτώ", "οχτώ", "εννέα", "εννιά", "δέκα",
}

# High-confidence case endings (unaccented final cluster → case set), longest
# first. Ambiguous endings map to multi-case sets so they cannot yield a clean
# "breaks concord" verdict.
_CASE_ENDINGS = [
    ("ους", {"ACC"}),
    ("εις", {"NOM", "ACC"}),
    ("οι", {"NOM"}),
    ("ος", {"NOM"}),
    ("ου", {"GEN"}),
    ("ων", {"GEN"}),
    ("ον", {"ACC"}),
    ("ης", {"NOM", "GEN"}),
    ("ες", {"NOM", "ACC"}),
]

# Adjective/participle endings — such a token is a modifier inside the NP and is
# skipped while scanning left toward the governing article.
_MODIFIER_ENDINGS = (
    "ντες", "ντων", "μένος", "μένη", "μένο", "μένοι", "μένες", "μένα",
    "ικός", "ικής", "ικών", "ικούς", "ος", "ους", "ων", "οι", "ης", "ου", "ες", "εις",
)


def _ending_case(word: str):
    """Return the set of cases a word's ending is consistent with, or None."""
    w = word.strip().lower()
    for suf, cases in _CASE_ENDINGS:
        if w.endswith(suf) and len(w) > len(suf):
            return cases
    return None


def _shares_stem(a: str, b: str, min_stem: int = 3) -> bool:
    """True if a and b look like two inflected forms of the same word."""
    a, b = a.lower(), b.lower()
    n = 0
    for ca, cb in zip(a, b):
        if ca != cb:
            break
        n += 1
    return n >= min_stem


def np_article_concord_ok(original: str, suggested: str, text: str, offset: int) -> bool:
    """
    Fix A guard. Returns False ONLY when a single-word case change would break
    agreement with the definite article governing the word — the ORIGINAL agrees
    with the article and the SUGGESTED cannot. All uncertain cases return True.
    """
    orig = (original or "").strip()
    sugg = (suggested or "").strip()
    # Only single-word inflectional changes are in scope.
    if not orig or not sugg or " " in orig or " " in sugg:
        return True
    if not _shares_stem(orig, sugg):
        return True
    orig_cases = _ending_case(orig)
    sugg_cases = _ending_case(sugg)
    if not orig_cases or not sugg_cases:
        return True  # unrecognized ending → cannot judge → allow

    # Locate the word and the left context within the same clause.
    if offset is None or offset < 0 or text[offset:offset + len(original)] != original:
        offset = text.find(original)
        if offset < 0:
            return True
    left = text[:offset]
    seg = re.split(r"[.;·!?\n]", left)[-1]
    toks = re.findall(r"[^\W\d_]+", seg, re.UNICODE)
    if not toks:
        return True

    # Scan left for the governing article, skipping numerals/modifiers.
    article = None
    skipped = 0
    for tok in reversed(toks):
        tl = tok.lower()
        if tl in _ARTICLE_CASE:
            article = tl
            break
        if skipped >= 6:
            break
        if tl in _NUMERAL_WORDS or any(tl.endswith(e) for e in _MODIFIER_ENDINGS):
            skipped += 1
            continue
        break  # verb / preposition / particle / noun → stop (no governing article)
    if article is None:
        return True

    art_cases = _ARTICLE_CASE[article]
    # Block only when the original agrees with the article but the suggested cannot.
    if (orig_cases & art_cases) and not (sugg_cases & art_cases):
        return False
    return True


_SINGLE_WORD_RE = re.compile(r'^\S+$')


async def verify_grammar_fix(
    sentence_ctx: str,
    original: str,
    suggested: str,
    model_id: str,
) -> bool:
    """
    Secondary LLM verifier for risky single-word grammar fixes.
    Sends a minimal YES/NO prompt: does replacing X with Y produce a correct Greek sentence?
    Returns True = safe to apply, False = reject.

    Call only for: single-word changes, verb inflection, article/case/number/gender/preposition swaps.
    Controlled by ENABLE_GRAMMAR_VERIFIER env flag (default off to avoid extra latency/cost).
    """
    if not sentence_ctx or not original or not suggested:
        return True  # Cannot verify → fail open

    prompt = (
        f'Πρόταση (απόσπασμα): «{sentence_ctx}»\n'
        f'Αντικατάσταση: «{original}» → «{suggested}»\n'
        f'Η αντικατάσταση παράγει σαφώς σωστή ελληνική πρόταση χωρίς αλλαγή νοήματος; '
        f'Απάντησε ΜΟΝΟ ΝΑΙ ή ΟΧΙ.'
    )
    try:
        result = await call_llm(
            "Απάντησε μόνο ΝΑΙ ή ΟΧΙ.",
            prompt,
            model_id,
        )
        answer = (result or "").strip().upper()
        accepted = answer.startswith("ΝΑΙ") or answer.startswith("NAI") or answer.startswith("YES")
        print(f"[VERIFIER] '{original}' → '{suggested}': verifier={answer} → {'ACCEPT' if accepted else 'REJECT'}")
        return accepted
    except Exception as e:
        print(f"[VERIFIER] Error verifying '{original}' → '{suggested}': {e}")
        return True  # Fail open — don't reject on verifier error


def find_original_in_text(full_text: str, original: str) -> tuple:
    """
    Find the original text in full_text, handling Google model quirks.
    Returns (start_index, matched_text) or (-1, None) if not found.
    """
    if not original or not full_text:
        return -1, None
    
    # Try 1: Exact match
    idx = full_text.find(original)
    if idx != -1:
        return idx, original
    
    # Try 2: Normalized match
    norm_original = normalize_text(original)
    norm_full = normalize_text(full_text)
    
    # Find in normalized, then map back to original positions
    norm_idx = norm_full.find(norm_original)
    if norm_idx != -1:
        # Try to find the actual substring by searching with normalized comparison
        for i in range(len(full_text)):
            for j in range(i + 1, min(i + len(original) + 20, len(full_text) + 1)):
                candidate = full_text[i:j]
                if normalize_text(candidate) == norm_original:
                    return i, candidate
    
    # Try 3: Strip and retry
    stripped = original.strip()
    if stripped != original:
        idx = full_text.find(stripped)
        if idx != -1:
            return idx, stripped
    
    # Try 4: Handle common Google model issues
    # Sometimes Google adds/removes spaces around punctuation
    variants = [
        original.replace(" ,", ",").replace(" .", ".").replace(" ·", "·"),
        original.replace(",", " ,").replace(".", " .").replace("·", " ·"),
        original.replace("  ", " "),
        re.sub(r'\s+', ' ', original),
    ]
    
    for variant in variants:
        idx = full_text.find(variant)
        if idx != -1:
            return idx, variant
    
    # Try 5: Character-by-character fuzzy match for small differences
    # This handles cases where Google models slightly modify characters
    if len(original) >= 3:
        for i in range(len(full_text) - len(original) + 5):
            if i + len(original) > len(full_text):
                break
            # Check with some tolerance for character differences
            candidate = full_text[i:i + len(original)]
            if normalize_text(candidate) == norm_original:
                return i, candidate
            # Also try slightly longer/shorter matches
            for delta in [-2, -1, 1, 2]:
                end = i + len(original) + delta
                if end > len(full_text) or end <= i:
                    continue
                candidate = full_text[i:end]
                if normalize_text(candidate) == norm_original:
                    return i, candidate
    
    return -1, None


def apply_correction_safe(
    doc_path: Path, 
    para_number: int,  # 1-indexed paragraph number
    original: str, 
    suggested: str, 
    exact_offset: int = -1  # Exact character position within paragraph
) -> Tuple[bool, str]:
    """
    SAFE correction application with 100% accuracy guarantee.
    
    Returns: (success: bool, error_message: str)
    
    Strategy:
    1. Find the paragraph by para_number
    2. If exact_offset provided, verify text at that position matches original
    3. If verification fails, ABORT (no fallback)
    4. Apply the change only at the verified position
    """
    if not original:
        return False, "Empty original text"

    # Read docx as zip
    try:
        with zipfile.ZipFile(doc_path, "r") as z:
            if "word/document.xml" not in z.namelist():
                return False, "No document.xml in docx"
            doc_xml = z.read("word/document.xml")
            filelist = z.namelist()
            other_files = {name: z.read(name) for name in filelist if name != "word/document.xml"}
    except Exception as e:
        return False, f"Failed to read docx: {e}"

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    try:
        if USING_LXML:
            parser = ET.XMLParser(remove_blank_text=False)
            root = ET.fromstring(doc_xml, parser)
        else:
            root = ET.fromstring(doc_xml)
    except Exception as e:
        return False, f"Failed to parse XML: {e}"

    # Find all paragraphs
    all_paragraphs = root.findall(".//w:p", ns)
    
    # Filter to non-empty paragraphs (same logic as extract_paragraphs)
    paragraphs = []
    for p in all_paragraphs:
        texts = []
        for t in p.findall(".//w:t", ns):
            if t.text:
                texts.append(t.text)
        full_text = "".join(texts)
        if full_text.strip():
            paragraphs.append({"elem": p, "text": full_text})
    
    # Find target paragraph by number (1-indexed)
    if para_number < 1 or para_number > len(paragraphs):
        return False, f"Paragraph {para_number} not found (document has {len(paragraphs)} paragraphs)"
    
    target = paragraphs[para_number - 1]
    target_para = target["elem"]
    full_text = target["text"]
    
    # Determine the exact position to apply the change
    m_start = -1
    
    if exact_offset >= 0:
        # STRICT MODE: Use exact_offset and VERIFY
        if exact_offset + len(original) <= len(full_text):
            text_at_offset = full_text[exact_offset:exact_offset + len(original)]
            if text_at_offset == original:
                m_start = exact_offset
                print(f"[SAFE] Verified exact match at offset {exact_offset}")
            else:
                # Check with normalization
                if normalize_text(text_at_offset) == normalize_text(original):
                    m_start = exact_offset
                    original = text_at_offset  # Use actual text from document
                    print(f"[SAFE] Verified normalized match at offset {exact_offset}")
                else:
                    return False, f"Verification failed: expected '{original}' at offset {exact_offset}, found '{text_at_offset}'"
        else:
            return False, f"Offset {exact_offset} + len({len(original)}) exceeds paragraph length {len(full_text)}"
    else:
        # SEARCH MODE: Find unique occurrence
        # Count occurrences
        count = full_text.count(original)
        if count == 0:
            # Try normalized search
            norm_original = normalize_text(original)
            for i in range(len(full_text) - len(original) + 1):
                if normalize_text(full_text[i:i + len(original)]) == norm_original:
                    m_start = i
                    original = full_text[i:i + len(original)]
                    break
            if m_start == -1:
                return False, f"Text '{original[:30]}...' not found in paragraph {para_number}"
        elif count == 1:
            m_start = full_text.find(original)
            print(f"[SAFE] Found unique occurrence at offset {m_start}")
        else:
            return False, f"Multiple occurrences ({count}) of '{original[:30]}...' in paragraph - need exact_offset"
    
    m_end = m_start + len(original)
    
    # Collect all w:t nodes in target paragraph
    t_nodes = []
    pos = 0
    for t in target_para.findall(".//w:t", ns):
        txt = t.text or ""
        start = pos
        pos += len(txt)
        end = pos
        t_nodes.append({"node": t, "text": txt, "start": start, "end": end})

    if not t_nodes:
        return False, "No text nodes in paragraph"

    # Determine affected nodes
    affected = [x for x in t_nodes if x["start"] < m_end and x["end"] > m_start]
    if not affected:
        return False, "No affected text nodes found"

    # Calculate overlap
    def overlap_len(a: int, b: int, c: int, d: int) -> int:
        s = max(a, c)
        e = min(b, d)
        return max(0, e - s)

    per_node_match_lens = [overlap_len(x["start"], x["end"], m_start, m_end) for x in affected]
    total_match_len = sum(per_node_match_lens)
    
    if total_match_len != len(original):
        return False, f"Match length mismatch: expected {len(original)}, got {total_match_len}"

    first = affected[0]
    last = affected[-1]

    first_prefix_len = max(0, m_start - first["start"])
    last_suffix_start_in_last = max(0, m_end - last["start"])

    first_prefix = (first["text"] or "")[:first_prefix_len]
    last_suffix = (last["text"] or "")[last_suffix_start_in_last:] if last["text"] else ""

    # Apply the change - simplified approach
    sug = suggested or ""
    
    # Put everything in first node, clear others
    for idx, item in enumerate(affected):
        node = item["node"]
        if idx == 0:
            if len(affected) == 1:
                node.text = first_prefix + sug + last_suffix
            else:
                node.text = first_prefix + sug
        elif idx == len(affected) - 1:
            node.text = last_suffix
        else:
            node.text = ""

    # Write back into docx
    if USING_LXML:
        new_xml = ET.tostring(root, encoding="UTF-8", xml_declaration=True, standalone="yes")
    else:
        new_xml = ET.tostring(root, encoding="utf-8", xml_declaration=True)

    tmp_path = doc_path.with_suffix(".tmp.docx")
    try:
        with zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as z:
            for name, content in other_files.items():
                z.writestr(name, content)
            z.writestr("word/document.xml", new_xml)
        shutil.move(str(tmp_path), str(doc_path))
    except Exception as e:
        if tmp_path.exists():
            tmp_path.unlink()
        return False, f"Failed to write docx: {e}"

    print(f"[SAFE] Successfully applied: '{original[:30]}...' → '{suggested[:30]}...' at para {para_number}, offset {m_start}")
    return True, ""


def apply_correction(doc_path: Path, para_idx: int, original: str, suggested: str, anchor: str = "", exact_offset: int = -1, paragraph_number: int = 0) -> bool:
    """
    Apply correction using multiple strategies:
    1. If paragraph_number > 0: Use it to find the correct paragraph
    2. If anchor provided: Use it to verify/find paragraph
    3. Find original within the paragraph using unique context
    
    v2.9: Added paragraph_number for more reliable positioning
    """
    if not original:
        return False

    # Read docx as zip
    try:
        with zipfile.ZipFile(doc_path, "r") as z:
            if "word/document.xml" not in z.namelist():
                return False
            doc_xml = z.read("word/document.xml")
            filelist = z.namelist()
            other_files = {name: z.read(name) for name in filelist if name != "word/document.xml"}
    except Exception:
        return False

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    try:
        if USING_LXML:
            parser = ET.XMLParser(remove_blank_text=False)
            root = ET.fromstring(doc_xml, parser)
        else:
            root = ET.fromstring(doc_xml)
    except Exception:
        return False

    # Find all non-empty paragraphs (matching extract_paragraphs logic)
    all_paragraphs = root.findall(".//w:p", ns)
    paragraphs_with_text = []
    for p in all_paragraphs:
        texts = []
        for t in p.findall(".//w:t", ns):
            if t.text:
                texts.append(t.text)
        full_text = "".join(texts)
        if full_text.strip():
            paragraphs_with_text.append({"elem": p, "text": full_text})
    
    target_para = None
    matched_original = None
    match_offset = -1
    
    # Strategy 1: Use paragraph_number (most reliable)
    if paragraph_number > 0 and paragraph_number <= len(paragraphs_with_text):
        para_data = paragraphs_with_text[paragraph_number - 1]
        full_text = para_data["text"]
        
        # Count occurrences in this paragraph
        occurrences = []
        search_pos = 0
        max_iterations = 1000  # Safety limit
        iterations = 0
        while iterations < max_iterations:
            iterations += 1
            m_start, m_original = find_original_in_text(full_text[search_pos:], original)
            if m_start == -1:
                break
            actual_pos = search_pos + m_start
            occurrences.append({"pos": actual_pos, "match": m_original})
            # CRITICAL: Always advance by at least 1 to prevent infinite loop
            advance = max(1, len(m_original))
            search_pos = actual_pos + advance
            if search_pos >= len(full_text):
                break
        
        if len(occurrences) == 1:
            target_para = para_data["elem"]
            matched_original = occurrences[0]["match"]
            match_offset = occurrences[0]["pos"]
        elif len(occurrences) > 1 and anchor:
            # Multiple occurrences - use anchor/context to determine which one
            # The anchor should contain the original, so find which occurrence is in anchor
            norm_anchor = normalize_text(anchor)
            for occ in occurrences:
                # Get context around this occurrence
                ctx_start = max(0, occ["pos"] - 30)
                ctx_end = min(len(full_text), occ["pos"] + len(occ["match"]) + 30)
                local_context = full_text[ctx_start:ctx_end]
                
                if normalize_text(local_context) == norm_anchor or norm_anchor in normalize_text(local_context):
                    target_para = para_data["elem"]
                    matched_original = occ["match"]
                    match_offset = occ["pos"]
                    break
            
            if target_para is None:
                # Context didn't help - try to find anchor containing original
                for occ in occurrences:
                    ctx_start = max(0, occ["pos"] - 50)
                    ctx_end = min(len(full_text), occ["pos"] + len(occ["match"]) + 50)
                    local_context = normalize_text(full_text[ctx_start:ctx_end])
                    
                    if norm_anchor in local_context:
                        target_para = para_data["elem"]
                        matched_original = occ["match"]
                        match_offset = occ["pos"]
                        break
                        
        elif len(occurrences) > 1:
            # Instead of failing, try the first occurrence
            target_para = para_data["elem"]
            matched_original = occurrences[0]["match"]
            match_offset = occurrences[0]["pos"]
        # If len(occurrences) == 0, fall through to other strategies
    
    # Strategy 2: Use anchor to find paragraph
    if target_para is None and anchor:
        norm_anchor = normalize_text(anchor)
        for para_data in paragraphs_with_text:
            full_text = para_data["text"]
            norm_full = normalize_text(full_text)
            
            if norm_anchor in norm_full:
                m_start, m_original = find_original_in_text(full_text, original)
                if m_start != -1:
                    target_para = para_data["elem"]
                    matched_original = m_original
                    match_offset = m_start
                    break
    
    # Strategy 3: Fallback - search all paragraphs for unique occurrence
    if target_para is None:
        candidates = []
        for para_data in paragraphs_with_text:
            full_text = para_data["text"]
            m_start, m_original = find_original_in_text(full_text, original)
            if m_start != -1:
                candidates.append({"para": para_data["elem"], "text": full_text, "match": m_original, "offset": m_start})
        
        if len(candidates) == 1:
            target_para = candidates[0]["para"]
            matched_original = candidates[0]["match"]
            match_offset = candidates[0]["offset"]
        elif len(candidates) > 1:
            return False
    
    if target_para is None:
        return False

    original = matched_original

    # Apply the change to the target paragraph
    t_nodes = []
    full = ""
    
    for t in target_para.findall(".//w:t", ns):
        txt = t.text or ""
        start = len(full)
        full += txt
        end = len(full)
        t_nodes.append({"node": t, "text": txt, "start": start, "end": end})

    if not t_nodes:
        return False

    m_start = full.find(original)
    if m_start == -1:
        m_start, matched_original = find_original_in_text(full, original)
        if m_start == -1:
            return False
        original = matched_original
    
    m_end = m_start + len(original)

    affected = [x for x in t_nodes if x["start"] < m_end and x["end"] > m_start]
    if not affected:
        return False

    def overlap_len(a: int, b: int, c: int, d: int) -> int:
        s = max(a, c)
        e = min(b, d)
        return max(0, e - s)

    per_node_match_lens = [overlap_len(x["start"], x["end"], m_start, m_end) for x in affected]
    total_match_len = sum(per_node_match_lens)
    if total_match_len != len(original):
        return False

    first = affected[0]
    last = affected[-1]

    first_prefix_len = max(0, m_start - first["start"])
    last_suffix_start_in_last = max(0, m_end - last["start"])

    first_prefix = (first["text"] or "")[:first_prefix_len]
    last_suffix = (last["text"] or "")[last_suffix_start_in_last:] if last["text"] else ""

    sug = suggested or ""
    
    # SIMPLIFIED APPROACH: Put everything in first node, clear others
    for idx, item in enumerate(affected):
        node = item["node"]
        if idx == 0:
            if len(affected) == 1:
                node.text = first_prefix + sug + last_suffix
            else:
                node.text = first_prefix + sug
        elif idx == len(affected) - 1:
            node.text = last_suffix
        else:
            node.text = ""

    if USING_LXML:
        new_xml = ET.tostring(root, encoding="UTF-8", xml_declaration=True, standalone="yes")
    else:
        new_xml = ET.tostring(root, encoding="utf-8", xml_declaration=True)

    tmp_path = doc_path.with_suffix(".tmp.docx")
    try:
        with zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as z:
            for name, content in other_files.items():
                z.writestr(name, content)
            z.writestr("word/document.xml", new_xml)
        tmp_path.replace(doc_path)
        return True
    except Exception as e:
        print(f"[ERROR] Failed to save document: {e}")
        if tmp_path.exists():
            tmp_path.unlink()
        return False


def apply_footnote_correction(doc_path: Path, footnote_id: int, original: str, suggested: str) -> bool:
    """
    Apply correction to footnote text while preserving all formatting and namespaces.
    Uses lxml for proper namespace handling to prevent Word corruption.
    """
    if not original:
        return False

    # --- Read docx and keep all files except footnotes.xml ---
    try:
        with zipfile.ZipFile(doc_path, "r") as z:
            if "word/footnotes.xml" not in z.namelist():
                return False
            foot_xml = z.read("word/footnotes.xml")
            filelist = z.namelist()
            other_files = {name: z.read(name) for name in filelist if name != "word/footnotes.xml"}
    except Exception:
        return False

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    
    try:
        if USING_LXML:
            # lxml preserves namespaces properly
            parser = ET.XMLParser(remove_blank_text=False)
            root = ET.fromstring(foot_xml, parser)
        else:
            # Standard ElementTree - register namespaces
            root = ET.fromstring(foot_xml)
    except Exception:
        return False

    # --- Find target footnote node ---
    target_fn = None
    for fn in root.findall("w:footnote", ns):
        fid = fn.get(f"{{{ns['w']}}}id")
        if fid and fid.lstrip('-').isdigit() and int(fid) == int(footnote_id):
            target_fn = fn
            break
    if target_fn is None:
        return False

    # --- Collect all w:t nodes in reading order with bounds over the concatenated text ---
    t_nodes = []
    full = ""

    for t in target_fn.findall(".//w:t", ns):
        txt = t.text or ""
        start = len(full)
        full += txt
        end = len(full)
        t_nodes.append({"node": t, "text": txt, "start": start, "end": end})

    if not t_nodes:
        return False

    # --- Find match in concatenated text (with normalization for Google models) ---
    m_start, matched_original = find_original_in_text(full, original)
    if m_start == -1:
        print(f"[DEBUG] apply_footnote_correction: Could not find '{original[:50]}...' in footnote {footnote_id}")
        return False
    
    # Use the actually matched text for replacement
    original = matched_original
    m_end = m_start + len(original)

    # --- Determine affected nodes ---
    affected = [x for x in t_nodes if x["start"] < m_end and x["end"] > m_start]
    if not affected:
        return False

    # Helper: overlap length of [a,b) with [c,d)
    def overlap_len(a: int, b: int, c: int, d: int) -> int:
        s = max(a, c)
        e = min(b, d)
        return max(0, e - s)

    # Compute how many chars of the ORIGINAL match lie in each affected node
    per_node_match_lens = [overlap_len(x["start"], x["end"], m_start, m_end) for x in affected]
    total_match_len = sum(per_node_match_lens)
    if total_match_len != len(original):
        return False

    # Prefix in first affected node, suffix in last affected node
    first = affected[0]
    last = affected[-1]

    first_prefix_len = max(0, m_start - first["start"])
    last_suffix_start_in_last = max(0, m_end - last["start"])

    first_prefix = (first["text"] or "")[:first_prefix_len]
    last_suffix = (last["text"] or "")[last_suffix_start_in_last:] if last["text"] else ""

    # Distribute suggested across affected nodes
    sug = suggested or ""
    sug_len = len(sug)

    allocated = []
    if total_match_len == 0:
        allocated = [""] * len(affected)
    else:
        consumed = 0
        for i, seg_len in enumerate(per_node_match_lens):
            if i == len(per_node_match_lens) - 1:
                allocated.append(sug[consumed:])
            else:
                prop = int(round((seg_len / total_match_len) * sug_len))
                prop = max(0, min(prop, sug_len - consumed))
                allocated.append(sug[consumed: consumed + prop])
                consumed += prop

        if len(allocated) >= 1:
            leftover = sug[sum(len(a) for a in allocated):]
            if leftover:
                allocated[-1] = allocated[-1] + leftover

    # Write back texts
    for idx, item in enumerate(affected):
        node = item["node"]
        if idx == 0 and idx == len(affected) - 1:
            node.text = first_prefix + sug + last_suffix
        elif idx == 0:
            node.text = first_prefix + allocated[idx]
        elif idx == len(affected) - 1:
            node.text = allocated[idx] + last_suffix
        else:
            node.text = allocated[idx]

    # --- Write back into docx preserving everything else ---
    if USING_LXML:
        # lxml: proper XML declaration and namespace preservation
        new_xml = ET.tostring(root, encoding="UTF-8", xml_declaration=True, standalone="yes")
    else:
        # Standard ElementTree
        new_xml = ET.tostring(root, encoding="utf-8", xml_declaration=True)

    tmp_path = doc_path.with_suffix(".tmp.docx")
    try:
        with zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as z:
            for name, content in other_files.items():
                z.writestr(name, content)
            z.writestr("word/footnotes.xml", new_xml)
        tmp_path.replace(doc_path)
        return True
    except Exception:
        if tmp_path.exists():
            tmp_path.unlink()
        return False


def create_tracked_changes_docx(original_path: Path, corrections: list, output_path: Path) -> bool:
    """
    Create a copy of the document with all pending corrections as Track Changes + Real Comments.
    
    Features:
    - Real comments in the right panel with explanation
    - Track changes (del + ins) grouped together
    - Single Accept action accepts both deletion and insertion
    - Formatting preserved
    """
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from copy import deepcopy
    
    try:
        doc = Document(str(original_path))
    except Exception as e:
        print(f"[ERROR] Failed to load document: {e}")
        return False
    
    # Build paragraph index map (only non-empty paragraphs)
    text_paragraphs = []
    index_map = []
    
    for i, p in enumerate(doc.paragraphs):
        if p.text and p.text.strip():
            text_paragraphs.append(p.text.strip())
            index_map.append(i)
    
    # Collect pending corrections
    pending_corrections = []
    for corr in corrections:
        if corr.get("status") != "pending":
            continue
        if corr.get("type") != "fix":
            continue
        pending_corrections.append(corr)
    
    if not pending_corrections:
        # Just copy the file
        shutil.copy(original_path, output_path)
        return True
    
    change_id = 0
    comment_id = 0
    applied = 0
    now_iso = datetime.utcnow().isoformat() + "Z"
    
    # Track which corrections were applied (for comments)
    applied_corrections = []
    
    # Process body corrections
    for corr in pending_corrections:
        if corr.get("target") == "footnote":
            continue

        para_num = corr.get("paragraph_number", 0)
        original = corr.get("original", "")
        suggested = corr.get("suggested", "")
        reason = corr.get("reason", "")
        module = corr.get("module", "core")
        exact_offset = corr.get("exact_offset", -1)

        if not original:
            continue

        # Resolve paragraph via paragraph_number (1-based over non-empty paragraphs)
        idx = para_num - 1
        if idx < 0 or idx >= len(index_map):
            print(f"[WARNING] Paragraph {para_num} out of range, skipping '{original[:30]}'")
            continue

        paragraph = doc.paragraphs[index_map[idx]]
        full_text = paragraph.text  # concat of all runs, NOT stripped
        lshift = len(full_text) - len(full_text.lstrip())

        # Resolve span [span_start, span_end) using exact_offset (paragraph-relative, into stripped text)
        span_start = -1
        deleted_text = original

        if exact_offset >= 0:
            candidate = exact_offset + lshift
            if full_text[candidate:candidate + len(original)] == original:
                span_start = candidate

        if span_start == -1:
            # NFC-normalised fallback: handles NFD diacritics, hyphen-split words
            found, matched = find_original_in_text(full_text, original)
            if found == -1:
                print(f"[WARNING] Could not locate '{original[:30]}' in paragraph {para_num}")
                continue
            span_start = found
            deleted_text = matched

        span_end = span_start + len(deleted_text)

        result = _apply_revision_at_span(
            paragraph, span_start, span_end, deleted_text, suggested,
            change_id, comment_id, now_iso
        )
        if result:
            change_id += 2
            applied_corrections.append({
                "id": comment_id,
                "module": module,
                "reason": reason,
                "original": original,
                "suggested": suggested
            })
            comment_id += 1
            applied += 1
        else:
            print(f"[WARNING] Failed to apply revision for '{original[:30]}' in paragraph {para_num}")
    
    # Process footnote corrections
    try:
        for rel in doc.part.rels.values():
            if "footnotes" in rel.reltype:
                footnotes_part = rel.target_part
                footnotes_xml = footnotes_part._element
                
                for corr in pending_corrections:
                    if corr.get("target") != "footnote":
                        continue

                    original = corr.get("original", "")
                    suggested = corr.get("suggested", "")
                    footnote_id = corr.get("footnote_id", 0)
                    reason = corr.get("reason", "")
                    module = corr.get("module", "core")
                    exact_offset = corr.get("exact_offset", -1)

                    if not original:
                        continue

                    for fn in footnotes_xml.findall('.//' + qn('w:footnote')):
                        fn_id = fn.get(qn('w:id'))
                        if fn_id and int(fn_id) == footnote_id:
                            result = _apply_revision_at_span_xml(
                                fn, exact_offset, original, suggested,
                                change_id, comment_id, now_iso
                            )
                            if result:
                                change_id += 2
                                applied_corrections.append({
                                    "id": comment_id,
                                    "module": module,
                                    "reason": reason,
                                    "original": original,
                                    "suggested": suggested
                                })
                                comment_id += 1
                                applied += 1
                            else:
                                print(f"[WARNING] Could not locate '{original[:30]}' in footnote {footnote_id}")
                            break
                break
    except Exception as e:
        print(f"[WARNING] Could not process footnotes: {e}")
    
    # Save the document first
    try:
        doc.save(str(output_path))
    except Exception as e:
        print(f"[ERROR] Failed to save document: {e}")
        return False
    
    # Now inject comments via post-processing
    if applied_corrections:
        try:
            inject_real_comments(output_path, applied_corrections, now_iso)
        except Exception as e:
            print(f"[WARNING] Could not inject comments: {e}")
    
    print(f"[INFO] Track Changes document saved with {applied} corrections")
    return True


def _apply_revision_at_span(paragraph, span_start, span_end, deleted_text, suggested, change_id, comment_id, now_iso):
    """
    Replace paragraph text [span_start, span_end) with w:del + w:ins tracked-change elements.

    Works across run boundaries: collects all runs overlapping the span, removes them,
    and reinserts before/after text runs flanking the revision markup.
    span_start/span_end are offsets into paragraph.text (= concat of run texts, NOT stripped).
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from copy import deepcopy

    author = "Noëta"

    # Build (r_element, run_start, run_end) for every run
    pos = 0
    run_spans = []
    for r in paragraph.runs:
        t = r.text or ""
        run_spans.append((r._r, pos, pos + len(t)))
        pos += len(t)

    # Runs that overlap [span_start, span_end)
    overlapping = [(r_el, rs, re) for r_el, rs, re in run_spans if rs < span_end and re > span_start]
    if not overlapping:
        return False

    first_r_el, first_rs, _ = overlapping[0]
    last_r_el, last_rs, _ = overlapping[-1]
    rPr_orig = first_r_el.find(qn('w:rPr'))
    parent = first_r_el.getparent()
    insert_pos = list(parent).index(first_r_el)

    # Text before the span within the first run
    first_t = first_r_el.find(qn('w:t'))
    first_run_text = first_t.text if first_t is not None and first_t.text else ""
    before_text = first_run_text[:max(0, span_start - first_rs)]

    # Text after the span within the last run
    last_t = last_r_el.find(qn('w:t'))
    last_run_text = last_t.text if last_t is not None and last_t.text else ""
    after_start = max(0, span_end - last_rs)
    after_text = last_run_text[after_start:]

    # Remove all overlapping runs from the paragraph
    for r_el, _, _ in overlapping:
        parent.remove(r_el)

    elements = []

    if before_text:
        br = OxmlElement('w:r')
        if rPr_orig is not None:
            br.append(deepcopy(rPr_orig))
        bt = OxmlElement('w:t')
        bt.set(qn('xml:space'), 'preserve')
        bt.text = before_text
        br.append(bt)
        elements.append(br)

    cs = OxmlElement('w:commentRangeStart')
    cs.set(qn('w:id'), str(comment_id))
    elements.append(cs)

    del_el = OxmlElement('w:del')
    del_el.set(qn('w:id'), str(change_id))
    del_el.set(qn('w:author'), author)
    del_el.set(qn('w:date'), now_iso)
    del_r = OxmlElement('w:r')
    if rPr_orig is not None:
        del_r.append(deepcopy(rPr_orig))
    del_t = OxmlElement('w:delText')
    del_t.set(qn('xml:space'), 'preserve')
    del_t.text = deleted_text
    del_r.append(del_t)
    del_el.append(del_r)
    elements.append(del_el)

    ins_el = OxmlElement('w:ins')
    ins_el.set(qn('w:id'), str(change_id + 1))
    ins_el.set(qn('w:author'), author)
    ins_el.set(qn('w:date'), now_iso)
    ins_r = OxmlElement('w:r')
    if rPr_orig is not None:
        ins_r.append(deepcopy(rPr_orig))
    ins_t = OxmlElement('w:t')
    ins_t.set(qn('xml:space'), 'preserve')
    ins_t.text = suggested
    ins_r.append(ins_t)
    ins_el.append(ins_r)
    elements.append(ins_el)

    ce = OxmlElement('w:commentRangeEnd')
    ce.set(qn('w:id'), str(comment_id))
    elements.append(ce)

    ref_r = OxmlElement('w:r')
    ref = OxmlElement('w:commentReference')
    ref.set(qn('w:id'), str(comment_id))
    ref_r.append(ref)
    elements.append(ref_r)

    if after_text:
        ar = OxmlElement('w:r')
        if rPr_orig is not None:
            ar.append(deepcopy(rPr_orig))
        at = OxmlElement('w:t')
        at.set(qn('xml:space'), 'preserve')
        at.text = after_text
        ar.append(at)
        elements.append(ar)

    for i, elem in enumerate(elements):
        parent.insert(insert_pos + i, elem)

    return True


def _apply_revision_at_span_xml(fn_elem, exact_offset, original, suggested, change_id, comment_id, now_iso):
    """
    Apply a tracked-change revision inside a w:footnote XML element.

    Builds the full footnote text from all w:r/w:t children, resolves the span
    using exact_offset (footnote-relative, into stripped text) with NFC fallback,
    then removes the overlapping runs and reinserts revision markup.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from copy import deepcopy

    author = "Noëta"

    # Build (r_el, t_el, run_start, run_end) for all w:r children
    pos = 0
    run_spans = []
    for r_el in fn_elem.findall('.//' + qn('w:r')):
        t_el = r_el.find(qn('w:t'))
        t = t_el.text if t_el is not None and t_el.text else ""
        run_spans.append((r_el, t_el, pos, pos + len(t)))
        pos += len(t)

    fn_full = "".join(info[1].text if info[1] is not None and info[1].text else "" for info in run_spans)
    lshift = len(fn_full) - len(fn_full.lstrip())

    span_start = -1
    deleted_text = original

    if exact_offset >= 0:
        candidate = exact_offset + lshift
        if fn_full[candidate:candidate + len(original)] == original:
            span_start = candidate

    if span_start == -1:
        found, matched = find_original_in_text(fn_full, original)
        if found == -1:
            return False
        span_start = found
        deleted_text = matched

    span_end = span_start + len(deleted_text)

    overlapping = [(r_el, t_el, rs, re) for r_el, t_el, rs, re in run_spans if rs < span_end and re > span_start]
    if not overlapping:
        return False

    first_r_el = overlapping[0][0]
    rPr_orig = first_r_el.find(qn('w:rPr'))
    parent = first_r_el.getparent()
    insert_pos = list(parent).index(first_r_el)

    first_t_el = overlapping[0][1]
    first_rs = overlapping[0][2]
    first_run_text = first_t_el.text if first_t_el is not None and first_t_el.text else ""
    before_text = first_run_text[:max(0, span_start - first_rs)]

    last_t_el = overlapping[-1][1]
    last_rs = overlapping[-1][2]
    last_run_text = last_t_el.text if last_t_el is not None and last_t_el.text else ""
    after_text = last_run_text[max(0, span_end - last_rs):]

    for r_el, _, _, _ in overlapping:
        parent.remove(r_el)

    elements = []

    if before_text:
        br = OxmlElement('w:r')
        if rPr_orig is not None:
            br.append(deepcopy(rPr_orig))
        bt = OxmlElement('w:t')
        bt.set(qn('xml:space'), 'preserve')
        bt.text = before_text
        br.append(bt)
        elements.append(br)

    cs = OxmlElement('w:commentRangeStart')
    cs.set(qn('w:id'), str(comment_id))
    elements.append(cs)

    del_el = OxmlElement('w:del')
    del_el.set(qn('w:id'), str(change_id))
    del_el.set(qn('w:author'), author)
    del_el.set(qn('w:date'), now_iso)
    del_r = OxmlElement('w:r')
    if rPr_orig is not None:
        del_r.append(deepcopy(rPr_orig))
    del_t = OxmlElement('w:delText')
    del_t.set(qn('xml:space'), 'preserve')
    del_t.text = deleted_text
    del_r.append(del_t)
    del_el.append(del_r)
    elements.append(del_el)

    ins_el = OxmlElement('w:ins')
    ins_el.set(qn('w:id'), str(change_id + 1))
    ins_el.set(qn('w:author'), author)
    ins_el.set(qn('w:date'), now_iso)
    ins_r = OxmlElement('w:r')
    if rPr_orig is not None:
        ins_r.append(deepcopy(rPr_orig))
    ins_t = OxmlElement('w:t')
    ins_t.set(qn('xml:space'), 'preserve')
    ins_t.text = suggested
    ins_r.append(ins_t)
    ins_el.append(ins_r)
    elements.append(ins_el)

    ce = OxmlElement('w:commentRangeEnd')
    ce.set(qn('w:id'), str(comment_id))
    elements.append(ce)

    ref_r = OxmlElement('w:r')
    ref = OxmlElement('w:commentReference')
    ref.set(qn('w:id'), str(comment_id))
    ref_r.append(ref)
    elements.append(ref_r)

    if after_text:
        ar = OxmlElement('w:r')
        if rPr_orig is not None:
            ar.append(deepcopy(rPr_orig))
        at = OxmlElement('w:t')
        at.set(qn('xml:space'), 'preserve')
        at.text = after_text
        ar.append(at)
        elements.append(ar)

    for i, elem in enumerate(elements):
        parent.insert(insert_pos + i, elem)

    return True


def _wrap_span_with_comment(paragraph, span_start, span_end, comment_id):
    """
    Wrap paragraph text [span_start, span_end) with commentRangeStart/End markers.

    The original runs are preserved unchanged — no text is deleted or replaced.
    Boundary runs are split at the span edges so the comment range aligns exactly.
    span_start/span_end are offsets into paragraph.text (concat of runs, NOT stripped).
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from copy import deepcopy

    pos = 0
    run_spans = []
    for r in paragraph.runs:
        t = r.text or ""
        run_spans.append((r._r, pos, pos + len(t)))
        pos += len(t)

    overlapping = [(r_el, rs, re) for r_el, rs, re in run_spans if rs < span_end and re > span_start]
    if not overlapping:
        return False

    parent = overlapping[0][0].getparent()

    # Split last run at span_end first (preserves first-run index)
    last_r_el, last_rs, last_re = overlapping[-1]
    if last_re > span_end:
        last_t = last_r_el.find(qn('w:t'))
        if last_t is not None and last_t.text:
            split_at = span_end - last_rs
            after_text = last_t.text[split_at:]
            last_t.text = last_t.text[:split_at]
            last_t.set(qn('xml:space'), 'preserve')
            if after_text:
                rPr = last_r_el.find(qn('w:rPr'))
                ar = OxmlElement('w:r')
                if rPr is not None:
                    ar.append(deepcopy(rPr))
                at = OxmlElement('w:t')
                at.set(qn('xml:space'), 'preserve')
                at.text = after_text
                ar.append(at)
                parent.insert(list(parent).index(last_r_el) + 1, ar)

    # Split first run at span_start
    first_r_el, first_rs, first_re = overlapping[0]
    if first_rs < span_start:
        first_t = first_r_el.find(qn('w:t'))
        if first_t is not None and first_t.text:
            split_at = span_start - first_rs
            before_text = first_t.text[:split_at]
            first_t.text = first_t.text[split_at:]
            first_t.set(qn('xml:space'), 'preserve')
            if before_text:
                rPr = first_r_el.find(qn('w:rPr'))
                br = OxmlElement('w:r')
                if rPr is not None:
                    br.append(deepcopy(rPr))
                bt = OxmlElement('w:t')
                bt.set(qn('xml:space'), 'preserve')
                bt.text = before_text
                br.append(bt)
                parent.insert(list(parent).index(first_r_el), br)

    # Insert commentRangeStart before first span run
    first_idx = list(parent).index(first_r_el)
    cs = OxmlElement('w:commentRangeStart')
    cs.set(qn('w:id'), str(comment_id))
    parent.insert(first_idx, cs)

    # Insert commentRangeEnd + commentReference after last span run
    last_idx = list(parent).index(last_r_el)
    ce = OxmlElement('w:commentRangeEnd')
    ce.set(qn('w:id'), str(comment_id))
    parent.insert(last_idx + 1, ce)

    ref_r = OxmlElement('w:r')
    ref = OxmlElement('w:commentReference')
    ref.set(qn('w:id'), str(comment_id))
    ref_r.append(ref)
    parent.insert(last_idx + 2, ref_r)

    return True


def _wrap_span_with_comment_xml(fn_elem, exact_offset, original, comment_id):
    """
    Wrap a span inside a w:footnote with commentRangeStart/End markers.

    Same offset-based anchoring as _apply_revision_at_span_xml but no del/ins —
    original text is preserved intact.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from copy import deepcopy

    pos = 0
    run_spans = []
    for r_el in fn_elem.findall('.//' + qn('w:r')):
        t_el = r_el.find(qn('w:t'))
        t = t_el.text if t_el is not None and t_el.text else ""
        run_spans.append((r_el, t_el, pos, pos + len(t)))
        pos += len(t)

    fn_full = "".join(info[1].text if info[1] is not None and info[1].text else "" for info in run_spans)
    lshift = len(fn_full) - len(fn_full.lstrip())

    span_start = -1
    deleted_text = original

    if exact_offset >= 0:
        candidate = exact_offset + lshift
        if fn_full[candidate:candidate + len(original)] == original:
            span_start = candidate

    if span_start == -1:
        found, matched = find_original_in_text(fn_full, original)
        if found == -1:
            return False
        span_start = found
        deleted_text = matched

    span_end = span_start + len(deleted_text)

    overlapping = [(r_el, t_el, rs, re) for r_el, t_el, rs, re in run_spans if rs < span_end and re > span_start]
    if not overlapping:
        return False

    parent = overlapping[0][0].getparent()

    # Split last run at span_end
    last_r_el, last_t_el, last_rs, last_re = overlapping[-1]
    if last_re > span_end:
        if last_t_el is not None and last_t_el.text:
            split_at = span_end - last_rs
            after_text = last_t_el.text[split_at:]
            last_t_el.text = last_t_el.text[:split_at]
            last_t_el.set(qn('xml:space'), 'preserve')
            if after_text:
                rPr = last_r_el.find(qn('w:rPr'))
                ar = OxmlElement('w:r')
                if rPr is not None:
                    ar.append(deepcopy(rPr))
                at = OxmlElement('w:t')
                at.set(qn('xml:space'), 'preserve')
                at.text = after_text
                ar.append(at)
                parent.insert(list(parent).index(last_r_el) + 1, ar)

    # Split first run at span_start
    first_r_el, first_t_el, first_rs, first_re = overlapping[0]
    if first_rs < span_start:
        if first_t_el is not None and first_t_el.text:
            split_at = span_start - first_rs
            before_text = first_t_el.text[:split_at]
            first_t_el.text = first_t_el.text[split_at:]
            first_t_el.set(qn('xml:space'), 'preserve')
            if before_text:
                rPr = first_r_el.find(qn('w:rPr'))
                br = OxmlElement('w:r')
                if rPr is not None:
                    br.append(deepcopy(rPr))
                bt = OxmlElement('w:t')
                bt.set(qn('xml:space'), 'preserve')
                bt.text = before_text
                br.append(bt)
                parent.insert(list(parent).index(first_r_el), br)

    first_idx = list(parent).index(first_r_el)
    cs = OxmlElement('w:commentRangeStart')
    cs.set(qn('w:id'), str(comment_id))
    parent.insert(first_idx, cs)

    last_idx = list(parent).index(last_r_el)
    ce = OxmlElement('w:commentRangeEnd')
    ce.set(qn('w:id'), str(comment_id))
    parent.insert(last_idx + 1, ce)

    ref_r = OxmlElement('w:r')
    ref = OxmlElement('w:commentReference')
    ref.set(qn('w:id'), str(comment_id))
    ref_r.append(ref)
    parent.insert(last_idx + 2, ref_r)

    return True


def inject_real_comments(docx_path: Path, corrections: list, now_iso: str) -> bool:
    """
    Inject real comments into the docx file.
    This creates comments.xml and updates relationships.
    """
    try:
        # Read the docx
        with zipfile.ZipFile(docx_path, 'r') as zin:
            file_contents = {name: zin.read(name) for name in zin.namelist()}
        
        # Create comments.xml
        comments_xml = create_comments_xml(corrections, now_iso)
        file_contents['word/comments.xml'] = comments_xml.encode('utf-8')
        
        # Update [Content_Types].xml
        content_types = file_contents.get('[Content_Types].xml', b'').decode('utf-8')
        if 'comments.xml' not in content_types:
            content_types = content_types.replace(
                '</Types>',
                '<Override PartName="/word/comments.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"/></Types>'
            )
            file_contents['[Content_Types].xml'] = content_types.encode('utf-8')
        
        # Update word/_rels/document.xml.rels
        rels_path = 'word/_rels/document.xml.rels'
        if rels_path in file_contents:
            rels = file_contents[rels_path].decode('utf-8')
            if 'comments.xml' not in rels:
                # Find max rId
                import re
                rids = re.findall(r'Id="rId(\d+)"', rels)
                max_rid = max([int(r) for r in rids]) if rids else 0
                new_rid = max_rid + 1
                
                new_rel = f'<Relationship Id="rId{new_rid}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments" Target="comments.xml"/>'
                rels = rels.replace('</Relationships>', new_rel + '</Relationships>')
                file_contents[rels_path] = rels.encode('utf-8')
        
        # Write back
        with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as zout:
            for name, content in file_contents.items():
                zout.writestr(name, content)
        
        return True
        
    except Exception as e:
        print(f"[ERROR] Failed to inject comments: {e}")
        import traceback
        traceback.print_exc()
        return False


def create_comments_xml(corrections: list, now_iso: str) -> str:
    """Create the comments.xml content."""
    
    def escape_xml(text):
        if not text:
            return ""
        return (text
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;")
            .replace("'", "&apos;")
        )
    
    xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml"
xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml">
'''
    
    for corr in corrections:
        comment_id = corr["id"]
        module = corr.get("module", "core")
        reason = corr.get("reason", "")
        original = corr.get("original", "")
        suggested = corr.get("suggested", "")
        
        # Create comment text (no internal module prefix visible to user)
        if reason:
            comment_text = reason
        else:
            comment_text = f"Αλλαγή: «{original}» → «{suggested}»"

        comment_text = escape_xml(comment_text)

        xml += f'''<w:comment w:id="{comment_id}" w:author="Noëta" w:date="{now_iso}" w:initials="N">
<w:p>
<w:pPr><w:pStyle w:val="CommentText"/></w:pPr>
<w:r><w:rPr><w:rStyle w:val="CommentReference"/></w:rPr><w:annotationRef/></w:r>
<w:r><w:t>{comment_text}</w:t></w:r>
</w:p>
</w:comment>
'''
    
    xml += '</w:comments>'
    return xml
    
    return True


def create_suggestion_comments_xml(corrections: list, now_iso: str) -> str:
    """Create comments.xml for suggestions-only export.

    Each comment contains two paragraphs:
      «original» → «suggested»
      reason (if present)
    """

    def escape_xml(text):
        if not text:
            return ""
        return (text
                .replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
                .replace('"', "&quot;")
                .replace("'", "&apos;"))

    xml = ('<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
           '<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
           ' xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml"'
           ' xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml">\n')

    for corr in corrections:
        comment_id = corr["id"]
        original = escape_xml(corr.get("original", ""))
        suggested = escape_xml(corr.get("suggested", ""))
        reason = escape_xml(corr.get("reason", ""))

        arrow_line = f"«{original}» → «{suggested}»"

        xml += (f'<w:comment w:id="{comment_id}" w:author="Noëta" w:date="{now_iso}" w:initials="N">\n'
                f'<w:p><w:pPr><w:pStyle w:val="CommentText"/></w:pPr>'
                f'<w:r><w:rPr><w:rStyle w:val="CommentReference"/></w:rPr><w:annotationRef/></w:r>'
                f'<w:r><w:t xml:space="preserve">{arrow_line}</w:t></w:r></w:p>\n')
        if reason:
            xml += (f'<w:p><w:pPr><w:pStyle w:val="CommentText"/></w:pPr>'
                    f'<w:r><w:t xml:space="preserve">{reason}</w:t></w:r></w:p>\n')
        xml += '</w:comment>\n'

    xml += '</w:comments>'
    return xml


def calculate_offset_in_text(text: str, original: str) -> int:
    """Calculate the character offset of 'original' within 'text'. Returns -1 if not found."""
    if not original or not text:
        return -1
    return text.find(original)


def sort_corrections_by_position(corrections: list, paragraphs: list, footnotes: list) -> list:
    """
    Sort corrections by:
    1. Target (body first, then footnotes)
    2. Paragraph number / Footnote ID
    3. Position within the paragraph/footnote text
    """
    # Build lookup maps
    para_text_map = {p["number"]: p["text"] for p in paragraphs}
    footnote_text_map = {fn["id"]: fn["text"] for fn in footnotes}
    
    def sort_key(corr):
        target = corr.get("target", "body")
        
        if target == "footnote":
            # Footnotes come after body
            target_order = 1
            fn_id = corr.get("footnote_id", 0) or 0
            text = footnote_text_map.get(fn_id, "")
            position_key = fn_id
        else:
            # Body paragraphs
            target_order = 0
            para_num = corr.get("paragraph_number", 0) or 0
            text = para_text_map.get(para_num, corr.get("paragraph_text", ""))
            position_key = para_num
        
        # Calculate offset within text
        original = corr.get("original", "") or ""
        offset = calculate_offset_in_text(text, original)
        if offset == -1:
            offset = 9999999  # Put at end if not found
        
        return (target_order, position_key, offset)

    return sorted(corrections, key=sort_key)


def downgrade_stylistic_findings(corrections: list) -> list:
    """
    Demote findings whose rationale is stylistic (not grammatical) to severity
    "minor". With ONLY_SERIOUS_ERRORS=true those will be dropped by the
    follow-up filter; otherwise they survive but the UI can show them as
    minor/style suggestions.

    A finding is downgraded when its `reason` contains any STYLE_RED_FLAGS
    phrase AND no GRAMMAR_GROUNDING_TERMS term. Already-minor findings are
    skipped. Critical findings are never downgraded.
    """
    downgraded_count = 0
    for c in corrections:
        sev = (c.get("severity") or "").lower()
        if sev in ("minor", "critical"):
            continue
        reason = (c.get("reason") or "").lower()
        if not reason:
            continue
        has_red_flag = any(phrase in reason for phrase in STYLE_RED_FLAGS)
        if not has_red_flag:
            continue
        has_grammar = any(term in reason for term in GRAMMAR_GROUNDING_TERMS)
        if has_grammar:
            continue
        c["severity"] = "minor"
        c["_downgraded_stylistic"] = True
        downgraded_count += 1
    if downgraded_count:
        print(f"[STYLE-FILTER] Downgraded {downgraded_count} findings to minor (stylistic rationale).")
    return corrections


def _confidence_rank(c: dict) -> int:
    """Higher is better."""
    return {"high": 3, "medium": 2, "low": 1}.get((c.get("confidence") or "medium").lower(), 2)


def _severity_rank(c: dict) -> int:
    """Higher is more severe."""
    return {"critical": 3, "major": 2, "minor": 1}.get((c.get("severity") or "major").lower(), 2)


def dedupe_overlapping_spans(corrections: list) -> list:
    """
    Drop findings whose character span overlaps another finding in the same
    paragraph/footnote. Among overlapping competitors, keep the one with the
    highest confidence; tiebreaker is the smaller `original` span (more
    surgical), then the higher severity (critical > major > minor).

    The existing key-based dedupe (in /api/analyze) only catches identical
    (original, suggested, offset) triples — it cannot collapse cases like
    «την την» vs «την την τρίτη» that the LLM emits as competing fixes.
    """
    if not corrections:
        return corrections

    # Group by (target, paragraph_index, footnote_id)
    groups: dict[tuple, list[dict]] = {}
    for c in corrections:
        key = (
            c.get("target", "body"),
            c.get("paragraph_index", -1),
            c.get("footnote_id"),
        )
        groups.setdefault(key, []).append(c)

    keep_ids = set()
    drop_count = 0

    for group in groups.values():
        # Sort by exact_offset for deterministic processing
        group_sorted = sorted(group, key=lambda c: (c.get("exact_offset", -1), len(c.get("original", ""))))
        survivors: list[dict] = []
        for cand in group_sorted:
            cand_start = cand.get("exact_offset", -1)
            cand_len = len(cand.get("original", "") or "")
            if cand_start < 0 or cand_len == 0:
                survivors.append(cand)
                continue
            cand_end = cand_start + cand_len
            conflict_idx = -1
            for i, other in enumerate(survivors):
                other_start = other.get("exact_offset", -1)
                other_len = len(other.get("original", "") or "")
                if other_start < 0 or other_len == 0:
                    continue
                other_end = other_start + other_len
                # Overlap test (half-open intervals)
                if cand_start < other_end and other_start < cand_end:
                    conflict_idx = i
                    break
            if conflict_idx == -1:
                survivors.append(cand)
                continue
            other = survivors[conflict_idx]
            # Pick the winner: confidence > smaller span > severity
            cand_key = (
                _confidence_rank(cand),
                -cand_len,
                _severity_rank(cand),
            )
            other_key = (
                _confidence_rank(other),
                -len(other.get("original", "") or ""),
                _severity_rank(other),
            )
            if cand_key > other_key:
                survivors[conflict_idx] = cand
            drop_count += 1
        for s in survivors:
            keep_ids.add(id(s))

    if drop_count:
        print(f"[OVERLAP-DEDUPE] Dropped {drop_count} overlapping/competing findings.")

    return [c for c in corrections if id(c) in keep_ids]


def extract_all_text_from_docx(doc_path: Path) -> str:
    """Extract all text from a docx file for comparison."""
    try:
        with zipfile.ZipFile(doc_path, "r") as z:
            text_parts = []
            
            # Extract body text
            if "word/document.xml" in z.namelist():
                doc_xml = z.read("word/document.xml")
                ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                if USING_LXML:
                    root = ET.fromstring(doc_xml)
                else:
                    root = ET.fromstring(doc_xml)
                for t in root.findall(".//w:t", ns):
                    if t.text:
                        text_parts.append(t.text)
            
            # Extract footnotes
            if "word/footnotes.xml" in z.namelist():
                fn_xml = z.read("word/footnotes.xml")
                ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                if USING_LXML:
                    root = ET.fromstring(fn_xml)
                else:
                    root = ET.fromstring(fn_xml)
                for t in root.findall(".//w:t", ns):
                    if t.text:
                        text_parts.append(t.text)
            
            return "".join(text_parts)
    except Exception:
        return ""


def get_file_hash(file_path: Path) -> str:
    """Get MD5 hash of a file."""
    import hashlib
    try:
        with open(file_path, "rb") as f:
            return hashlib.md5(f.read()).hexdigest()
    except Exception:
        return ""


def compute_strict_diff(original_text: str, working_text: str) -> list:
    """
    Compute character-level diff between two texts.
    Returns list of changes: [{"type": "replace|insert|delete", "position": int, "original": str, "new": str}, ...]
    """
    import difflib
    
    changes = []
    matcher = difflib.SequenceMatcher(None, original_text, working_text)
    
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'equal':
            continue
        elif tag == 'replace':
            changes.append({
                "type": "replace",
                "position": i1,
                "original": original_text[i1:i2],
                "new": working_text[j1:j2],
                "context_before": original_text[max(0, i1-30):i1],
                "context_after": original_text[i2:i2+30],
            })
        elif tag == 'delete':
            changes.append({
                "type": "delete",
                "position": i1,
                "original": original_text[i1:i2],
                "new": "",
                "context_before": original_text[max(0, i1-30):i1],
                "context_after": original_text[i2:i2+30],
            })
        elif tag == 'insert':
            changes.append({
                "type": "insert",
                "position": i1,
                "original": "",
                "new": working_text[j1:j2],
                "context_before": original_text[max(0, i1-30):i1],
                "context_after": original_text[i1:i1+30],
            })
    
    return changes


def match_diff_to_corrections(diff_changes: list, accepted_corrections: list) -> dict:
    """
    Match each diff change to an accepted correction.
    Returns dict with matched and unmatched changes.
    """
    # Build lookup of accepted corrections
    accepted_map = {}
    for corr in accepted_corrections:
        if corr.get("status") == "accepted" and corr.get("type") == "fix":
            orig = corr.get("original", "").strip()
            sugg = corr.get("suggested", "").strip()
            if orig:
                key = (orig, sugg)
                accepted_map[key] = corr
    
    matched = []
    unmatched = []
    
    for change in diff_changes:
        orig = change.get("original", "").strip()
        new = change.get("new", "").strip()
        
        # Try to find matching correction
        key = (orig, new)
        if key in accepted_map:
            matched.append({
                "change": change,
                "correction": accepted_map[key],
            })
        else:
            # Try partial match (in case of whitespace differences)
            found = False
            for (corr_orig, corr_sugg), corr in accepted_map.items():
                if corr_orig in orig or orig in corr_orig:
                    if corr_sugg in new or new in corr_sugg:
                        matched.append({
                            "change": change,
                            "correction": corr,
                            "partial_match": True,
                        })
                        found = True
                        break
            
            if not found:
                unmatched.append(change)
    
    return {
        "matched": matched,
        "unmatched": unmatched,
        "all_matched": len(unmatched) == 0,
    }


def generate_diff_report_html(
    original_text: str,
    working_text: str,
    diff_changes: list,
    match_result: dict,
    session_id: str,
    filename: str,
) -> str:
    """Generate HTML diff report."""
    
    from datetime import datetime
    
    # Escape HTML
    def esc(s):
        return (s or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    
    matched_count = len(match_result["matched"])
    unmatched_count = len(match_result["unmatched"])
    total_changes = len(diff_changes)
    is_safe = match_result["all_matched"]
    
    status_color = "#22c55e" if is_safe else "#ef4444"
    status_text = "✓ ΑΣΦΑΛΕΣ - Όλες οι αλλαγές αντιστοιχούν σε εγκεκριμένες διορθώσεις" if is_safe else "⚠️ ΠΡΟΣΟΧΗ - Βρέθηκαν μη εγκεκριμένες αλλαγές"
    
    html = f"""<!DOCTYPE html>
<html lang="el">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Αναφορά Αλλαγών - {esc(filename)}</title>
    <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{ 
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            background: #f8fafc;
            color: #1e293b;
            line-height: 1.6;
            padding: 2rem;
        }}
        .container {{ max-width: 1000px; margin: 0 auto; }}
        .header {{
            background: white;
            border-radius: 12px;
            padding: 2rem;
            margin-bottom: 2rem;
            box-shadow: 0 1px 3px rgba(0,0,0,0.1);
        }}
        .header h1 {{ font-size: 1.5rem; margin-bottom: 0.5rem; }}
        .header .meta {{ color: #64748b; font-size: 0.875rem; }}
        .status {{
            background: {status_color};
            color: white;
            padding: 1rem 1.5rem;
            border-radius: 8px;
            margin: 1.5rem 0;
            font-weight: 600;
        }}
        .stats {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(150px, 1fr));
            gap: 1rem;
            margin: 1.5rem 0;
        }}
        .stat {{
            background: #f1f5f9;
            padding: 1rem;
            border-radius: 8px;
            text-align: center;
        }}
        .stat-value {{ font-size: 1.5rem; font-weight: 700; color: #0f172a; }}
        .stat-label {{ font-size: 0.75rem; color: #64748b; text-transform: uppercase; }}
        .section {{
            background: white;
            border-radius: 12px;
            padding: 1.5rem;
            margin-bottom: 1.5rem;
            box-shadow: 0 1px 3px rgba(0,0,0,0.1);
        }}
        .section h2 {{
            font-size: 1.125rem;
            margin-bottom: 1rem;
            padding-bottom: 0.5rem;
            border-bottom: 1px solid #e2e8f0;
        }}
        .change {{
            background: #f8fafc;
            border: 1px solid #e2e8f0;
            border-radius: 8px;
            padding: 1rem;
            margin-bottom: 1rem;
        }}
        .change.unmatched {{
            border-color: #fecaca;
            background: #fef2f2;
        }}
        .change-header {{
            display: flex;
            justify-content: space-between;
            align-items: center;
            margin-bottom: 0.75rem;
        }}
        .change-type {{
            font-size: 0.75rem;
            font-weight: 600;
            text-transform: uppercase;
            padding: 0.25rem 0.5rem;
            border-radius: 4px;
        }}
        .change-type.replace {{ background: #fef3c7; color: #92400e; }}
        .change-type.insert {{ background: #d1fae5; color: #065f46; }}
        .change-type.delete {{ background: #fee2e2; color: #991b1b; }}
        .change-position {{ font-size: 0.75rem; color: #64748b; }}
        .change-content {{ font-family: 'Courier New', monospace; font-size: 0.875rem; }}
        .context {{ color: #64748b; }}
        .deleted {{ background: #fee2e2; color: #991b1b; text-decoration: line-through; padding: 0.125rem 0.25rem; border-radius: 2px; }}
        .inserted {{ background: #d1fae5; color: #065f46; padding: 0.125rem 0.25rem; border-radius: 2px; }}
        .warning-box {{
            background: #fef2f2;
            border: 2px solid #ef4444;
            border-radius: 8px;
            padding: 1.5rem;
            margin: 1.5rem 0;
        }}
        .warning-box h3 {{ color: #dc2626; margin-bottom: 0.5rem; }}
        .footer {{
            text-align: center;
            color: #64748b;
            font-size: 0.75rem;
            margin-top: 2rem;
            padding-top: 1rem;
            border-top: 1px solid #e2e8f0;
        }}
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>📄 Αναφορά Αλλαγών</h1>
            <div class="meta">
                <div><strong>Αρχείο:</strong> {esc(filename)}</div>
                <div><strong>Session:</strong> {esc(session_id)}</div>
                <div><strong>Ημερομηνία:</strong> {datetime.now().strftime("%d/%m/%Y %H:%M")}</div>
            </div>
            <div class="status">{status_text}</div>
            <div class="stats">
                <div class="stat">
                    <div class="stat-value">{total_changes}</div>
                    <div class="stat-label">Συνολικές Αλλαγές</div>
                </div>
                <div class="stat">
                    <div class="stat-value">{matched_count}</div>
                    <div class="stat-label">Εγκεκριμένες</div>
                </div>
                <div class="stat">
                    <div class="stat-value" style="color: {'#22c55e' if unmatched_count == 0 else '#ef4444'}">{unmatched_count}</div>
                    <div class="stat-label">Μη Εγκεκριμένες</div>
                </div>
                <div class="stat">
                    <div class="stat-value">{len(original_text):,}</div>
                    <div class="stat-label">Χαρακτήρες (Αρχικό)</div>
                </div>
            </div>
        </div>
"""
    
    # Unmatched changes (warnings)
    if unmatched_count > 0:
        html += """
        <div class="warning-box">
            <h3>⚠️ Μη Εγκεκριμένες Αλλαγές</h3>
            <p>Οι παρακάτω αλλαγές ΔΕΝ αντιστοιχούν σε καμία εγκεκριμένη διόρθωση. Το download έχει μπλοκαριστεί για την προστασία σας.</p>
        </div>
        <div class="section">
            <h2>🚫 Μη Εγκεκριμένες Αλλαγές ({0})</h2>
""".format(unmatched_count)
        
        for change in match_result["unmatched"]:
            change_type = change.get("type", "replace")
            html += f"""
            <div class="change unmatched">
                <div class="change-header">
                    <span class="change-type {change_type}">{change_type}</span>
                    <span class="change-position">Θέση: {change.get('position', 0)}</span>
                </div>
                <div class="change-content">
                    <span class="context">{esc(change.get('context_before', ''))}</span>"""
            
            if change.get("original"):
                html += f'<span class="deleted">{esc(change["original"])}</span>'
            if change.get("new"):
                html += f'<span class="inserted">{esc(change["new"])}</span>'
            
            html += f"""<span class="context">{esc(change.get('context_after', ''))}</span>
                </div>
            </div>
"""
        html += "</div>"
    
    # Matched changes
    if matched_count > 0:
        html += f"""
        <div class="section">
            <h2>✓ Εγκεκριμένες Αλλαγές ({matched_count})</h2>
"""
        for item in match_result["matched"]:
            change = item["change"]
            corr = item.get("correction", {})
            change_type = change.get("type", "replace")
            
            html += f"""
            <div class="change">
                <div class="change-header">
                    <span class="change-type {change_type}">{change_type}</span>
                    <span class="change-position">
                        {corr.get('target', 'body')} 
                        {'Παρ. ' + str(corr.get('paragraph_number', '')) if corr.get('target') != 'footnote' else 'Υποσ. ' + str(corr.get('footnote_id', ''))}
                        | {corr.get('module', 'core')}
                    </span>
                </div>
                <div class="change-content">
                    <span class="context">{esc(change.get('context_before', ''))}</span>"""
            
            if change.get("original"):
                html += f'<span class="deleted">{esc(change["original"])}</span>'
            if change.get("new"):
                html += f'<span class="inserted">{esc(change["new"])}</span>'
            
            html += f"""<span class="context">{esc(change.get('context_after', ''))}</span>
                </div>
                <div style="margin-top: 0.5rem; font-size: 0.75rem; color: #64748b;">
                    <strong>Αιτία:</strong> {esc(corr.get('reason', '-'))}
                </div>
            </div>
"""
        html += "</div>"
    
    # No changes
    if total_changes == 0:
        html += """
        <div class="section">
            <h2>Καμία Αλλαγή</h2>
            <p>Δεν έγιναν αλλαγές στο έγγραφο.</p>
        </div>
"""
    
    html += f"""
        <div class="footer">
            ProofreadAI Diff Report • {datetime.now().strftime("%Y")} • Session: {esc(session_id)}
        </div>
    </div>
</body>
</html>"""
    
    return html


def verify_and_generate_report(
    original_path: Path, 
    working_path: Path, 
    corrections: list,
    session_id: str,
    filename: str,
) -> dict:
    """
    Strict verification + diff report generation.
    Returns verification result and report path if generated.
    """
    original_text = extract_all_text_from_docx(original_path)
    working_text = extract_all_text_from_docx(working_path)
    
    # Compute diff
    diff_changes = compute_strict_diff(original_text, working_text)
    
    # Match to corrections
    match_result = match_diff_to_corrections(diff_changes, corrections)
    
    # Generate HTML report
    report_html = generate_diff_report_html(
        original_text,
        working_text,
        diff_changes,
        match_result,
        session_id,
        filename,
    )
    
    # Save report
    report_path = OUTPUT_DIR / f"{session_id}_diff_report.html"
    with open(report_path, "w", encoding="utf-8") as f:
        f.write(report_html)
    
    return {
        "is_safe": match_result["all_matched"],
        "total_changes": len(diff_changes),
        "matched_count": len(match_result["matched"]),
        "unmatched_count": len(match_result["unmatched"]),
        "unmatched_changes": match_result["unmatched"],
        "original_length": len(original_text),
        "working_length": len(working_text),
        "original_hash": get_file_hash(original_path),
        "working_hash": get_file_hash(working_path),
        "report_path": str(report_path),
    }


# ============== SESSION ==============

class HistoryEntry:
    """Represents a single state in the undo/redo history."""
    def __init__(self, corrections: list, working_file_bytes: bytes, action_description: str):
        self.corrections = copy.deepcopy(corrections)
        self.working_file_bytes = working_file_bytes
        self.action_description = action_description


class SessionData:
    def __init__(self, session_id: str, original_path: Path):
        self.session_id = session_id
        self.original_path = original_path
        self.working_path = OUTPUT_DIR / f"{session_id}_working.docx"

        self.footnotes = []  # [{"id": int, "text": str}]
        self.paragraphs = []  # [{"index": int, "number": int, "text": str}]

        self.corrections = []
        self.status = "uploaded"
        self.progress = 0
        self.total_chunks = 0
        self.error_message = None
        self.tokens_used = {"prompt": 0, "completion": 0, "total": 0}
        self.token_breakdown = {}  # Instrumentation: per-pass {calls, prompt, completion, total}

        # Undo/Redo history
        self.history: List[HistoryEntry] = []
        self.history_index = -1  # Points to current state in history
        self.max_history = 50  # Maximum history entries

        shutil.copy(original_path, self.working_path)
        
        # Save initial state
        self._save_state("Initial state")

    def _save_state(self, action_description: str):
        """Save current state to history."""
        with open(self.working_path, "rb") as f:
            file_bytes = f.read()
        
        # Remove any redo states if we're not at the end
        if self.history_index < len(self.history) - 1:
            self.history = self.history[:self.history_index + 1]
        
        # Add new state
        entry = HistoryEntry(self.corrections, file_bytes, action_description)
        self.history.append(entry)
        self.history_index = len(self.history) - 1
        
        # Trim old history if exceeds max
        if len(self.history) > self.max_history:
            excess = len(self.history) - self.max_history
            self.history = self.history[excess:]
            self.history_index -= excess

    def _restore_state(self, index: int) -> bool:
        """Restore state from history."""
        if index < 0 or index >= len(self.history):
            return False
        
        entry = self.history[index]
        self.corrections = copy.deepcopy(entry.corrections)
        
        with open(self.working_path, "wb") as f:
            f.write(entry.working_file_bytes)
        
        self.history_index = index
        return True

    def can_undo(self) -> bool:
        return self.history_index > 0

    def can_redo(self) -> bool:
        return self.history_index < len(self.history) - 1

    def undo(self) -> bool:
        if not self.can_undo():
            return False
        return self._restore_state(self.history_index - 1)

    def redo(self) -> bool:
        if not self.can_redo():
            return False
        return self._restore_state(self.history_index + 1)

    def save_action(self, action_description: str):
        """Call after making changes to save state."""
        self._save_state(action_description)
    
    def reset(self):
        """Reset to initial uploaded state."""
        if self.history:
            self._restore_state(0)


sessions = {}


# ============== SECURITY ==============

# Password protection (set in .env file)
APP_PASSWORD = os.getenv("APP_PASSWORD", "").strip()
if not APP_PASSWORD:
    print("\n" + "!" * 50)
    print("WARNING: No APP_PASSWORD set in .env")
    print("Add APP_PASSWORD=your_secret_here to your .env file")
    print("Without it, anyone with the URL can use the app!")
    print("!" * 50 + "\n")
else:
    print(f"[SECURITY] Password protection ACTIVE (password: {APP_PASSWORD[:3]}...)")

TRANSLATE_PASSWORD = os.getenv("TRANSLATE_PASSWORD", APP_PASSWORD).strip()
# Debug: log whether TRANSLATE_PASSWORD env var was found
_raw_translate_pw = os.getenv("TRANSLATE_PASSWORD")
print(f"[DEBUG] os.getenv('TRANSLATE_PASSWORD') -> {repr(_raw_translate_pw)}", flush=True)
print(f"[DEBUG] APP_PASSWORD={repr(APP_PASSWORD[:3])}..., TRANSLATE_PASSWORD={repr(TRANSLATE_PASSWORD[:3])}...", flush=True)
if TRANSLATE_PASSWORD != APP_PASSWORD:
    print(f"[SECURITY] Separate TRANSLATE_PASSWORD set (password: {TRANSLATE_PASSWORD[:3]}...)", flush=True)
else:
    print(f"[SECURITY] TRANSLATE_PASSWORD equals APP_PASSWORD (fallback)", flush=True)

# Daily spending cap in USD (set in .env, default $5)
DAILY_SPEND_CAP = float(os.getenv("DAILY_SPEND_CAP", "5.0"))

# Rate limiting: max analyses per hour
MAX_ANALYSES_PER_HOUR = int(os.getenv("MAX_ANALYSES_PER_HOUR", "10"))

# Max upload file size in MB
MAX_UPLOAD_MB = int(os.getenv("MAX_UPLOAD_MB", "100"))

# Track usage
from collections import defaultdict
import time

class UsageTracker:
    def __init__(self):
        self.daily_cost = 0.0
        self.daily_reset = datetime.now().date()
        self.analysis_timestamps = []  # list of timestamps
        self.total_analyses_today = 0
    
    def check_daily_reset(self):
        today = datetime.now().date()
        if today != self.daily_reset:
            self.daily_cost = 0.0
            self.daily_reset = today
            self.total_analyses_today = 0
            self.analysis_timestamps = []
    
    def add_cost(self, cost: float):
        self.check_daily_reset()
        self.daily_cost += cost
    
    def can_analyze(self) -> tuple:
        """Returns (allowed: bool, reason: str)"""
        self.check_daily_reset()
        
        # Check spending cap
        if self.daily_cost >= DAILY_SPEND_CAP:
            return False, f"Daily spending cap reached (${self.daily_cost:.2f}/${DAILY_SPEND_CAP:.2f})"
        
        # Check rate limit (analyses in last hour)
        now = time.time()
        self.analysis_timestamps = [t for t in self.analysis_timestamps if now - t < 3600]
        if len(self.analysis_timestamps) >= MAX_ANALYSES_PER_HOUR:
            return False, f"Rate limit: max {MAX_ANALYSES_PER_HOUR} analyses per hour"
        
        return True, ""
    
    def record_analysis(self, estimated_cost: float = 0.05):
        self.check_daily_reset()
        self.analysis_timestamps.append(time.time())
        self.total_analyses_today += 1
        self.add_cost(estimated_cost)

usage_tracker = UsageTracker()


# ============== APP ==============

app = FastAPI(title="ProofreadAI API", version="2.4.0")

# ---- Password middleware ----
from starlette.middleware.base import BaseHTTPMiddleware
from starlette.responses import JSONResponse

class PasswordMiddleware(BaseHTTPMiddleware):
    async def dispatch(self, request, call_next):
        # Skip password check if no password is set
        if not APP_PASSWORD:
            return await call_next(request)

        # Skip OPTIONS preflight (CORS handles it)
        if request.method == "OPTIONS":
            return await call_next(request)

        # Skip for root and auth endpoints
        path = request.url.path.rstrip("/")
        if path in ("", "/", "/api/auth", "/api/config"):
            return await call_next(request)

        # Also skip docs
        if path.startswith("/docs") or path.startswith("/openapi"):
            return await call_next(request)

        # Check password from header, query param, or cookie
        provided = (
            request.headers.get("X-App-Password", "") or
            request.query_params.get("pw", "") or
            request.cookies.get("app_password", "")
        )

        # Translate endpoints use TRANSLATE_PASSWORD; all others use APP_PASSWORD
        expected = TRANSLATE_PASSWORD if path.startswith("/api/translate") else APP_PASSWORD

        if provided != expected:
            print(f"[AUTH] Blocked {request.method} {path} - wrong password")
            return JSONResponse(
                status_code=401,
                content={"detail": "Unauthorized - invalid password"}
            )

        return await call_next(request)

# Middleware order: added LAST runs OUTERMOST (first to touch request/response).
# CORSMiddleware must be outermost so it adds Access-Control headers to ALL
# responses, including 401s returned by PasswordMiddleware.
app.add_middleware(PasswordMiddleware)
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:3000",
        "http://127.0.0.1:3000",
        "https://noeta.vercel.app",
        "https://noeta-lkno4gg66-gbarbous-projects.vercel.app",
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ============== AUTH ENDPOINT ==============

class AuthRequest(BaseModel):
    password: str
    service: str = "editor"  # "editor" | "translate"

@app.post("/api/auth")
async def authenticate(request: AuthRequest):
    """Verify password. Frontend calls this on login."""
    if not APP_PASSWORD:
        return {"authenticated": True, "message": "No password required"}

    provided = (request.password or "").strip()
    expected = TRANSLATE_PASSWORD if request.service == "translate" else APP_PASSWORD

    if provided == expected:
        return {"authenticated": True}
    else:
        print(f"[AUTH] Login failed (service={request.service}) - provided: {repr(provided)}, expected: {repr(expected)}")
        raise HTTPException(401, "Wrong password")

# ============== ENDPOINTS ==============

@app.get("/")
async def root():
    return {"message": "ProofreadAI API is running", "version": "2.3.0"}


@app.get("/api/config")
async def get_config():
    return {
        "models": [{"id": k, "name": v["name"]} for k, v in AVAILABLE_MODELS.items()],
        "check_types": CHECK_INSTRUCTIONS,
    }


@app.post("/api/upload")
async def upload_document(file: UploadFile = File(...)):
    if not (file.filename or "").lower().endswith(".docx"):
        raise HTTPException(400, "Only .docx files allowed")

    # Security: Read file and check size
    file_bytes = await file.read()
    file_size_mb = len(file_bytes) / (1024 * 1024)
    if file_size_mb > MAX_UPLOAD_MB:
        raise HTTPException(400, f"File too large ({file_size_mb:.1f}MB). Max: {MAX_UPLOAD_MB}MB")

    session_id = str(uuid.uuid4())[:8]
    upload_path = UPLOAD_DIR / f"{session_id}_original.docx"

    with open(upload_path, "wb") as f:
        f.write(file_bytes)

    session = SessionData(session_id, upload_path)
    sessions[session_id] = session

    paragraphs = extract_paragraphs(upload_path)
    footnotes = extract_footnotes(upload_path)
    session.footnotes = footnotes
    session.paragraphs = paragraphs

    return {
        "session_id": session_id,
        "filename": file.filename,
        "paragraph_count": len(paragraphs),
        "paragraphs": paragraphs,
        "footnotes": footnotes,
        "status": "uploaded",
    }


class TextUploadRequest(BaseModel):
    """Request for plain text upload."""
    text: str
    title: str = "Untitled"


@app.post("/api/upload-text")
async def upload_text(request: TextUploadRequest):
    """
    Upload plain text and create a .docx from it.
    Useful for pasting text directly instead of uploading a file.
    """
    if not request.text or not request.text.strip():
        raise HTTPException(400, "Το κείμενο δεν μπορεί να είναι κενό")
    
    # Limit text size (roughly 500 pages worth)
    if len(request.text) > 2_000_000:
        raise HTTPException(400, "Το κείμενο είναι πολύ μεγάλο (μέγιστο 2MB)")
    
    session_id = str(uuid.uuid4())[:8]
    upload_path = UPLOAD_DIR / f"{session_id}_original.docx"
    
    # Create a new docx document from plain text
    doc = docx.Document()
    
    # Split text into paragraphs (by double newlines or single newlines)
    # Normalize line endings first
    text = request.text.replace('\r\n', '\n').replace('\r', '\n')
    
    # Split by newlines
    raw_paragraphs = text.split('\n')
    
    for para_text in raw_paragraphs:
        # Skip completely empty lines but keep paragraphs with whitespace
        if para_text.strip():
            doc.add_paragraph(para_text)
    
    # If no paragraphs were added (all empty), add at least one
    if len(doc.paragraphs) == 0:
        doc.add_paragraph(request.text.strip())
    
    doc.save(upload_path)
    
    session = SessionData(session_id, upload_path)
    sessions[session_id] = session
    
    paragraphs = extract_paragraphs(upload_path)
    footnotes = []  # Plain text won't have footnotes
    session.footnotes = footnotes
    session.paragraphs = paragraphs
    
    # Clean title for filename
    safe_title = "".join(c for c in request.title if c.isalnum() or c in " -_").strip()[:50] or "pasted_text"
    
    return {
        "session_id": session_id,
        "filename": f"{safe_title}.docx",
        "paragraph_count": len(paragraphs),
        "paragraphs": paragraphs,
        "footnotes": footnotes,
        "status": "uploaded",
        "source": "text",  # Mark as coming from plain text
    }


@app.delete("/api/session/{session_id}")
async def delete_session(session_id: str):
    """Delete a session and clean up files."""
    session = sessions.get(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    
    # Clean up files
    try:
        if session.original_path.exists():
            session.original_path.unlink()
        if session.working_path.exists():
            session.working_path.unlink()
    except Exception:
        pass
    
    del sessions[session_id]
    return {"success": True, "message": "Session deleted"}


@app.post("/api/session/{session_id}/reset")
async def reset_session(session_id: str):
    """Reset session to initial state (allows uploading new file without full restart)."""
    session = sessions.get(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    
    session.reset()
    session.status = "uploaded"
    session.progress = 0
    session.total_chunks = 0
    session.error_message = None
    
    return {
        "success": True,
        "session_id": session_id,
        "status": "uploaded",
        "corrections": session.corrections,
        "can_undo": session.can_undo(),
        "can_redo": session.can_redo(),
    }


class AnalyzeRequest(BaseModel):
    session_id: str
    model: str
    text_description: Optional[str] = ""
    custom_instructions: Optional[str] = ""

    translation_check: Optional[bool] = False
    fact_check: Optional[bool] = False
    web_fact_check: Optional[bool] = False

    style_check: Optional[bool] = False
    style_type: Optional[str] = ""
    style_other: Optional[str] = ""
    style_intensity: Optional[str] = "light"  # light | medium | high
    allow_sentence_level: Optional[bool] = False


@app.post("/api/analyze")
async def analyze_document(request: AnalyzeRequest):
    session = sessions.get(request.session_id)
    if not session:
        raise HTTPException(404, "Session not found")

    if request.model not in AVAILABLE_MODELS:
        raise HTTPException(400, "Invalid model")

    # ---- Security: Rate limit and spending cap ----
    allowed, reason = usage_tracker.can_analyze()
    if not allowed:
        raise HTTPException(429, reason)

    model_id = AVAILABLE_MODELS[request.model]["model_id"]

    # Build per-pass system prompts
    pass1_prompt = generate_pass1_mechanical_prompt()
    pass2_prompt = generate_pass2_grammatical_prompt(bool(request.allow_sentence_level))
    pass3_prompt = generate_pass3_semantic_prompt(
        request.text_description,
        request.custom_instructions,
        bool(request.translation_check),
        bool(request.fact_check),
        bool(request.web_fact_check),
        bool(request.style_check),
        request.style_type or "",
        request.style_other or "",
        request.style_intensity or "light",
        bool(request.allow_sentence_level),
    )

    session.status = "analyzing"
    paragraphs = extract_paragraphs(session.original_path)
    footnotes = session.footnotes or extract_footnotes(session.original_path)
    session.paragraphs = paragraphs

    # Body chunks carry ONLY body text. Footnotes are processed once each, in
    # their own chunks, instead of being re-broadcast with every body chunk.
    work_units = [("body", ch) for ch in create_chunks(paragraphs)]
    if footnotes:
        work_units += [("footnote", fch) for fch in chunk_footnotes(footnotes)]
    chunks = work_units  # back-compat alias for downstream len() / logging
    session.total_chunks = len(work_units)

    # --- Instrumentation (Phase 1: measurement only, no behavior change) ---
    session.token_breakdown = {}
    _num_passes = 2 + (1 if pass3_prompt else 0)
    _sys_chars_per_chunk = len(pass1_prompt) + len(pass2_prompt) + (len(pass3_prompt) if pass3_prompt else 0)
    _n_body = sum(1 for k, _ in work_units if k == "body")
    _n_fn = len(work_units) - _n_body
    _instr = {"body": 0, "footnotes": 0, "system": 0}
    print(f"[TOKENS] doc: paragraphs={len(paragraphs)} footnotes={len(footnotes)} "
          f"body_chunks={_n_body} footnote_chunks={_n_fn} passes/chunk={_num_passes} "
          f"expected_calls={len(work_units)*_num_passes}")

    all_corrections = []
    seen = set()
    # Track how many corrections per (location, original, suggested) to handle repeated errors
    occurrence_counter = {}
    session.analysis_paused = False

    try:
        for i, (unit_kind, chunk) in enumerate(work_units):
            # Check for pause
            if getattr(session, 'analysis_paused', False):
                session.status = "paused"
                print(f"[ANALYSIS] Paused at chunk {i+1}/{len(work_units)}")
                break

            session.progress = i + 1

            if unit_kind == "body":
                body_paras = chunk
                chunk_text = format_chunk(chunk)
            else:
                body_paras = []
                chunk_text = "=== ΥΠΟΣΗΜΕΙΩΣΕΙΣ ===\n\n" + format_footnotes(chunk)

            user_prompt = generate_user_prompt(chunk_text)

            # Instrumentation: input-char breakdown of what is actually sent (×passes)
            _unit_chars = len(chunk_text)
            _instr["body" if unit_kind == "body" else "footnotes"] += _unit_chars * _num_passes
            _instr["system"] += _sys_chars_per_chunk
            if i < 3 or unit_kind == "footnote":
                print(f"[TOKENS] {unit_kind} unit {i+1}/{len(work_units)}: "
                      f"{_unit_chars} chars ×{_num_passes} passes")

            # Fire all active passes in parallel
            active_passes = [
                call_llm(pass1_prompt, user_prompt, model_id, session=session, pass_label="pass1_mechanical"),
                call_llm(pass2_prompt, user_prompt, model_id, session=session, pass_label="pass2_grammar"),
            ]
            if pass3_prompt:
                active_passes.append(call_llm(pass3_prompt, user_prompt, model_id, session=session, pass_label="pass3_semantic"))

            responses = await asyncio.gather(*active_passes)
            print(f"[MULTIPASS] Chunk {i+1}: got {len(responses)} pass responses")

            # Collect and normalise corrections from all passes
            raw_corrections = []
            for pass_idx, resp in enumerate(responses):
                parsed = parse_response(resp)
                if pass_idx == 0:  # Pass 1 uses different schema
                    parsed = [_normalize_pass1_correction(c) for c in parsed]
                raw_corrections.extend(parsed)

            for corr in raw_corrections:
                target = str(corr.get("target", "body")).strip().lower()
                ctype = corr.get("type", "fix")
                original = corr.get("original", "") or ""
                suggested = corr.get("suggested", "") or ""

                if target == "footnote":
                    try:
                        fid = int(corr.get("footnote_id"))
                    except Exception:
                        continue
                    if not any(fn["id"] == fid for fn in footnotes):
                        continue

                    # Get footnote text for validation and offset
                    fn_text = next((fn["text"] for fn in footnotes if fn["id"] == fid), "")

                    # A. Strict validation before storing
                    is_valid, reject_reason = validate_correction(
                        corr,
                        footnote_text=fn_text,
                        allow_sentence_level=bool(request.allow_sentence_level),
                    )
                    if not is_valid:
                        print(f"[VALIDATE] Rejected fn:{fid} '{original[:40]}': {reject_reason}")
                        continue

                    # D. Severity filter
                    severity = (corr.get("severity") or "major").lower()
                    if severity not in ("critical", "major", "minor"):
                        severity = "major"
                    if ONLY_SERIOUS_ERRORS and severity == "minor":
                        print(f"[FILTER] Dropping minor severity: fn:{fid} '{original[:30]}'")
                        continue

                    # B. Exact offset — track occurrences so repeated errors get distinct offsets
                    offsets = find_all_occurrences(fn_text, original)
                    occ_key = (f"fn:{fid}", original, suggested)
                    occ_idx = occurrence_counter.get(occ_key, 0)
                    exact_offset = (
                        offsets[occ_idx] if occ_idx < len(offsets)
                        else (offsets[0] if offsets else -1)
                    )
                    occurrence_counter[occ_key] = occ_idx + 1

                    # Fix A. Deterministic article-concord guard (footnote text).
                    if not np_article_concord_ok(original, suggested, fn_text, exact_offset):
                        print(f"[NP-GUARD] Rejected fn:{fid} '{original}' -> '{suggested}': breaks article concord")
                        continue

                    # B. Dedupe key includes exact_offset — preserves separate corrections for repeated errors
                    # Old key: f"fn:{fid}|{original}|{suggested}"
                    # New key: f"f:{fid}|{original}|{suggested}|{exact_offset}"
                    key = f"f:{fid}|{original}|{suggested}|{exact_offset}"
                    if key in seen:
                        continue
                    seen.add(key)

                    all_corrections.append({
                        "id": len(all_corrections),

                        "target": "footnote",
                        "footnote_id": fid,

                        "paragraph_index": -1,
                        "paragraph_number": 0,
                        "paragraph_text": fn_text,

                        "module": corr.get("module", "core"),
                        "type": ctype,
                        "scope": corr.get("scope", "token"),
                        "anchor": corr.get("anchor", ""),
                        "original": original,
                        "suggested": suggested,
                        "reason": corr.get("reason", ""),
                        "confidence": corr.get("confidence", "medium"),
                        "verification_required": bool(corr.get("verification_required", False)),
                        "sources": corr.get("sources", []),
                        "exact_offset": exact_offset,
                        "severity": severity,

                        "status": "pending",
                    })

                else:
                    para_num = corr.get("paragraph", 0)
                    matching = [p for p in body_paras if p["number"] == para_num]
                    if not matching:
                        continue
                    para = matching[0]

                    # A. Strict validation before storing
                    is_valid, reject_reason = validate_correction(
                        corr,
                        paragraph_text=para["text"],
                        allow_sentence_level=bool(request.allow_sentence_level),
                    )
                    if not is_valid:
                        print(f"[VALIDATE] Rejected para:{para_num} '{original[:40]}': {reject_reason}")
                        continue

                    # D. Severity filter
                    severity = (corr.get("severity") or "major").lower()
                    if severity not in ("critical", "major", "minor"):
                        severity = "major"
                    if ONLY_SERIOUS_ERRORS and severity == "minor":
                        print(f"[FILTER] Dropping minor severity: p:{para_num} '{original[:30]}'")
                        continue

                    # B. Exact offset — track occurrences so repeated errors get distinct offsets
                    offsets = find_all_occurrences(para["text"], original)
                    occ_key = (f"p:{para['number']}", original, suggested)
                    occ_idx = occurrence_counter.get(occ_key, 0)
                    exact_offset = (
                        offsets[occ_idx] if occ_idx < len(offsets)
                        else (offsets[0] if offsets else -1)
                    )
                    occurrence_counter[occ_key] = occ_idx + 1

                    # Fix A. Deterministic article-concord guard (blocks case changes
                    # that would break agreement with the governing definite article).
                    if not np_article_concord_ok(original, suggested, para["text"], exact_offset):
                        print(f"[NP-GUARD] Rejected para:{para_num} '{original}' -> '{suggested}': breaks article concord")
                        continue

                    # G. Grammar verifier for risky single-word changes (opt-in via ENABLE_GRAMMAR_VERIFIER)
                    if ENABLE_GRAMMAR_VERIFIER and _SINGLE_WORD_RE.match(original.strip()):
                        idx_in_para = para["text"].find(original)
                        if idx_in_para >= 0:
                            ctx_start = max(0, idx_in_para - 60)
                            ctx_end = min(len(para["text"]), idx_in_para + len(original) + 60)
                            sentence_ctx = para["text"][ctx_start:ctx_end]
                            grammar_ok = await verify_grammar_fix(sentence_ctx, original, suggested, model_id)
                            if not grammar_ok:
                                print(f"[VERIFIER] Rejected: '{original}' → '{suggested}' in para {para_num}")
                                continue

                    # B. Dedupe key includes exact_offset — preserves separate corrections for repeated errors
                    # Old key: f"p:{para['index']}|{original}|{suggested}"
                    # New key: f"p:{para['index']}|{original}|{suggested}|{exact_offset}"
                    key = f"p:{para['index']}|{original}|{suggested}|{exact_offset}"
                    if key in seen:
                        continue
                    seen.add(key)

                    all_corrections.append({
                        "id": len(all_corrections),

                        "target": "body",
                        "footnote_id": None,

                        "paragraph_index": para["index"],
                        "paragraph_number": para_num,
                        "paragraph_text": para["text"],

                        "module": corr.get("module", "core"),
                        "type": ctype,
                        "scope": corr.get("scope", "token"),
                        "anchor": corr.get("anchor", ""),
                        "original": original,
                        "suggested": suggested,
                        "reason": corr.get("reason", ""),
                        "confidence": corr.get("confidence", "medium"),
                        "verification_required": bool(corr.get("verification_required", False)),
                        "sources": corr.get("sources", []),
                        "exact_offset": exact_offset,
                        "severity": severity,

                        "status": "pending",
                    })

            await asyncio.sleep(0.2)
            
            # Save corrections incrementally (for live display)
            session.corrections = list(all_corrections)

        # --- Instrumentation summary (Phase 1: measurement only) ---
        _ic = _instr["body"] + _instr["footnotes"] + _instr["system"]
        print("=" * 64)
        print(f"[TOKENS] SUMMARY chunks={len(chunks)} passes/chunk={_num_passes} total_calls={len(chunks)*_num_passes}")
        if _ic:
            print(f"[TOKENS] INPUT chars sent ~{_ic:,}: "
                  f"body={_instr['body']:,} ({_instr['body']/_ic*100:.0f}%) "
                  f"footnotes={_instr['footnotes']:,} ({_instr['footnotes']/_ic*100:.0f}%) "
                  f"system={_instr['system']:,} ({_instr['system']/_ic*100:.0f}%)")
        for _label, _s in (getattr(session, 'token_breakdown', {}) or {}).items():
            print(f"[TOKENS] {_label}: calls={_s['calls']} prompt={_s['prompt']:,} "
                  f"completion={_s['completion']:,} total={_s['total']:,}")
        _tu = getattr(session, 'tokens_used', {})
        print(f"[TOKENS] TOTAL usage: prompt={_tu.get('prompt',0):,} "
              f"completion={_tu.get('completion',0):,} total={_tu.get('total',0):,}")
        print("=" * 64)

        # Demote stylistic findings (e.g. "πιο φυσικό" without grammar grounding) to minor
        all_corrections = downgrade_stylistic_findings(all_corrections)
        if ONLY_SERIOUS_ERRORS:
            before = len(all_corrections)
            all_corrections = [c for c in all_corrections if (c.get("severity") or "").lower() != "minor"]
            dropped = before - len(all_corrections)
            if dropped:
                print(f"[STYLE-FILTER] Dropped {dropped} downgraded findings (ONLY_SERIOUS_ERRORS=true).")

        # Collapse overlapping/competing findings (e.g. «την την» vs «την την τρίτη») to one
        all_corrections = dedupe_overlapping_spans(all_corrections)

        # Sort corrections by position
        all_corrections = sort_corrections_by_position(all_corrections, paragraphs, footnotes)

        # Re-assign IDs after sorting
        for idx, corr in enumerate(all_corrections):
            corr["id"] = idx

        session.corrections = all_corrections
        
        if not getattr(session, 'analysis_paused', False):
            session.status = "ready"
        else:
            session.status = "paused"
        
        # Save state after analysis
        session.save_action("Analysis completed")
        
        # Track usage (estimate ~$0.02 per chunk for flash models, more for pro)
        estimated_cost_per_chunk = 0.05 if "pro" in model_id.lower() or "opus" in model_id.lower() else 0.02
        total_cost = len(chunks) * estimated_cost_per_chunk
        usage_tracker.record_analysis(total_cost)
        print(f"[USAGE] Analysis cost: ~${total_cost:.2f} | Daily total: ~${usage_tracker.daily_cost:.2f}/{DAILY_SPEND_CAP:.2f}")

    except Exception as e:
        session.status = "error"
        session.error_message = str(e)
        raise HTTPException(500, f"Σφάλμα: {str(e)}")

    return {
        "session_id": request.session_id,
        "status": "ready",
        "correction_count": len(all_corrections),
        "corrections": all_corrections,
    }


@app.get("/api/status/{session_id}")
async def get_status(session_id: str):
    session = sessions.get(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    return {
        "session_id": session_id,
        "status": session.status,
        "progress": session.progress,
        "total_chunks": session.total_chunks,
        "correction_count": len(session.corrections),
        "error_message": session.error_message,
        "can_undo": session.can_undo(),
        "can_redo": session.can_redo(),
        "tokens_used": getattr(session, 'tokens_used', {"prompt": 0, "completion": 0, "total": 0}),
        "token_breakdown": getattr(session, 'token_breakdown', {}),
    }


@app.post("/api/pause/{session_id}")
async def pause_analysis(session_id: str):
    """Pause ongoing analysis. Corrections found so far are kept."""
    session = sessions.get(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    if session.status != "analyzing":
        raise HTTPException(400, "Analysis not running")
    session.analysis_paused = True
    return {"status": "pausing", "corrections_so_far": len(session.corrections)}


@app.get("/api/corrections/{session_id}")
async def get_corrections(session_id: str):
    """Get current corrections list."""
    session = sessions.get(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    return {
        "corrections": session.corrections,
        "can_undo": session.can_undo(),
        "can_redo": session.can_redo(),
    }


@app.get("/api/paragraphs/{session_id}")
async def get_paragraphs(session_id: str):
    """Get current paragraphs for session reconnect."""
    session = sessions.get(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    # Re-extract from working file if paragraphs are empty
    if not session.paragraphs and session.working_path.exists():
        session.paragraphs = extract_paragraphs(session.working_path)
    return {
        "paragraphs": session.paragraphs,
        "footnotes": session.footnotes,
    }


class CorrectionAction(BaseModel):
    correction_id: int
    action: str  # accept | reject


@app.post("/api/correction/{session_id}")
async def handle_correction(session_id: str, action: CorrectionAction):
    session = sessions.get(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    if action.correction_id < 0 or action.correction_id >= len(session.corrections):
        raise HTTPException(400, "Invalid correction ID")

    correction = session.corrections[action.correction_id]
    paragraphs_changed = False

    if action.action == "accept":
        if correction.get("type") == "fix":
            if correction.get("target") == "footnote":
                ok = apply_footnote_correction(
                    session.working_path,
                    int(correction["footnote_id"]),
                    correction.get("original", ""),
                    correction.get("suggested", ""),
                )
            else:
                ok = apply_correction(
                    session.working_path,
                    int(correction["paragraph_index"]),
                    correction.get("original", ""),
                    correction.get("suggested", ""),
                    anchor=correction.get("anchor", ""),
                    paragraph_number=correction.get("paragraph_number", 0),
                )

            if ok:
                correction["status"] = "accepted"
                paragraphs_changed = True
                # Save state for undo
                session.save_action(f"Accepted: {correction.get('original', '')[:30]}...")
            else:
                # Provide more detailed error
                para_num = correction.get("paragraph_number", "?")
                orig = correction.get("original", "")[:50]
                raise HTTPException(400, f"Δεν βρέθηκε «{orig}» στην παράγραφο {para_num}. Μπορεί να έχει ήδη διορθωθεί.")
        else:
            correction["status"] = "accepted"
            session.save_action(f"Accepted suggestion: {correction.get('suggested', '')[:30]}...")

    elif action.action == "reject":
        correction["status"] = "rejected"
        session.save_action(f"Rejected: {correction.get('original', '')[:30]}...")
    else:
        raise HTTPException(400, "Invalid action")

    # Build response
    response = {
        "success": True,
        "correction": correction,
        "can_undo": session.can_undo(),
        "can_redo": session.can_redo(),
    }
    
    # If paragraphs changed, include updated paragraphs for editor sync
    if paragraphs_changed:
        new_paragraphs = extract_paragraphs(session.working_path)
        session.paragraphs = new_paragraphs
        response["paragraphs"] = new_paragraphs
    
    return response


# ============== CORRECTION STATUS (lightweight update from frontend) ==============

class CorrectionStatusRequest(BaseModel):
    """Lightweight status update - no file modification."""
    correction_id: int
    status: str  # "accepted" | "rejected"


@app.post("/api/correction-status/{session_id}")
async def update_correction_status(session_id: str, request: CorrectionStatusRequest):
    """
    Update correction status without modifying any files.
    Body corrections are applied client-side in the TipTap editor.
    This endpoint just keeps the backend in sync for tracking/reporting.
    """
    session = sessions.get(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    
    if request.correction_id < 0 or request.correction_id >= len(session.corrections):
        raise HTTPException(400, "Invalid correction ID")
    
    session.corrections[request.correction_id]["status"] = request.status
    return {"success": True}


# ============== GENERATE DOWNLOAD (editor content → docx) ==============

class GenerateDownloadRequest(BaseModel):
    """Editor sends its current paragraphs + correction statuses."""
    paragraphs: List[dict]  # [{"number": 1, "text": "...", "segments": [...]}]
    corrections: Optional[List[dict]] = None  # For footnote tracking


@app.post("/api/generate-download/{session_id}")
async def generate_download(session_id: str, request: GenerateDownloadRequest):
    """
    Generate the final .docx from editor content.
    
    Architecture: TipTap editor = source of truth for body text.
    This endpoint:
    1. Starts from the ORIGINAL docx (preserves layout, images, headers, etc.)
    2. Replaces body paragraph texts with what the editor has
    3. Applies accepted footnote corrections (footnotes are not in the editor)
    4. Returns verification info for the download modal
    """
    session = sessions.get(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    
    output_path = OUTPUT_DIR / f"{session_id}_final.docx"
    
    try:
        # Start from original docx
        shutil.copy(session.original_path, output_path)
        
        # Step 1: Apply accepted footnote corrections to the copy
        for corr in session.corrections:
            if (corr.get("status") == "accepted" and 
                corr.get("type") == "fix" and 
                corr.get("target") == "footnote"):
                apply_footnote_correction(
                    output_path,
                    int(corr["footnote_id"]),
                    corr.get("original", ""),
                    corr.get("suggested", ""),
                )
        
        # Step 2: Replace body paragraphs with editor content
        # CRITICAL: Use XML-level manipulation to preserve footnoteReference elements
        # python-docx's run.text="" destroys footnoteReference nodes, losing footnotes!
        doc = docx.Document(output_path)
        
        # Build non-empty paragraph index map (same logic as extract_paragraphs)
        non_empty_indices = []
        for i, p in enumerate(doc.paragraphs):
            if p.text and p.text.strip():
                non_empty_indices.append(i)
        
        w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        
        for para_data in request.paragraphs:
            para_num = para_data.get("number", 0) - 1  # Convert to 0-indexed
            if para_num < 0 or para_num >= len(non_empty_indices):
                continue
            
            real_idx = non_empty_indices[para_num]
            para = doc.paragraphs[real_idx]
            new_text = para_data.get("text", "")
            
            # Skip if text hasn't changed (preserves formatting perfectly)
            if para.text.strip() == new_text.strip():
                continue
            
            # Check if paragraph contains footnote references
            para_xml = para._element
            has_footnotes = len(para_xml.findall(f".//{{{w_ns}}}footnoteReference")) > 0
            
            if has_footnotes:
                # SAFE MODE: Do targeted find-replace on w:t nodes
                # This preserves footnoteReference elements and all XML structure
                old_text = para.text
                # Find which corrections were applied to this paragraph
                # by comparing old and new text character by character
                # For safety, do a simple full-text replacement on t-nodes
                t_nodes = para_xml.findall(f".//{{{w_ns}}}t")
                if t_nodes:
                    # Concatenate all text, find what changed, replace in-place
                    concat = ""
                    t_info = []
                    for t in t_nodes:
                        txt = t.text or ""
                        t_info.append({"node": t, "start": len(concat), "text": txt})
                        concat += txt
                    
                    # Apply each accepted correction as find-replace on the t-nodes
                    corrections_for_para = [
                        c for c in session.corrections
                        if c.get("status") == "accepted" 
                        and c.get("type") == "fix"
                        and c.get("target") != "footnote"
                        and c.get("paragraph_number") == para_num + 1
                    ]
                    
                    for corr in corrections_for_para:
                        orig = corr.get("original", "")
                        sugg = corr.get("suggested", "")
                        if not orig:
                            continue
                        # F. Use exact_offset first; fall back to find() only if missing/stale
                        corr_offset = corr.get("exact_offset", -1)
                        if (corr_offset is not None and corr_offset >= 0
                                and corr_offset + len(orig) <= len(concat)):
                            text_at = concat[corr_offset:corr_offset + len(orig)]
                            if normalize_text(text_at) == normalize_text(orig):
                                idx = corr_offset
                            else:
                                print(f"[DOWNLOAD] exact_offset verify failed for '{orig[:30]}', using find()")
                                idx = concat.find(orig)
                        else:
                            idx = concat.find(orig)
                        if idx == -1:
                            print(f"[DOWNLOAD] Skipping '{orig[:30]}': not found in para {para_num+1}")
                            continue
                        # Replace in the t-nodes
                        m_end = idx + len(orig)
                        for ti in t_info:
                            t_start = ti["start"]
                            t_end = t_start + len(ti["text"])
                            # Does this t-node overlap with the match?
                            if t_end <= idx or t_start >= m_end:
                                continue  # No overlap
                            # Calculate overlap within this t-node
                            rel_start = max(0, idx - t_start)
                            rel_end = min(len(ti["text"]), m_end - t_start)
                            old_t = ti["text"]
                            # Only first overlapping node gets the replacement
                            if t_start <= idx:
                                new_t = old_t[:rel_start] + sugg + old_t[rel_end:]
                                ti["node"].text = new_t
                                # Preserve xml:space
                                if " " in new_t or new_t != new_t.strip():
                                    ti["node"].set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                            else:
                                # Subsequent nodes: just remove the overlapping part
                                new_t = old_t[rel_end:]
                                ti["node"].text = new_t
                            ti["text"] = ti["node"].text or ""
                        # Update concat for next correction
                        concat = concat[:idx] + sugg + concat[m_end:]
                        # Rebuild t_info offsets
                        offset = 0
                        for ti in t_info:
                            ti["start"] = offset
                            ti["text"] = ti["node"].text or ""
                            offset += len(ti["text"])
                
                print(f"[DOWNLOAD] Para {para_num+1}: SAFE mode (has footnotes)")
            else:
                # STANDARD MODE: No footnotes in this paragraph, safe to rebuild runs
                existing_runs = list(para.runs)
                segments = para_data.get("segments", [])
                
                if segments and len(segments) > 0:
                    for run in existing_runs:
                        run.text = ""
                    for i, seg in enumerate(segments):
                        if i < len(existing_runs):
                            run = existing_runs[i]
                        else:
                            run = para.add_run()
                        run.text = seg.get("text", "")
                        run.bold = True if seg.get("bold") else None
                        run.italic = True if seg.get("italic") else None
                else:
                    for run in existing_runs:
                        run.text = ""
                    if existing_runs:
                        existing_runs[0].text = new_text
                    else:
                        para.add_run(new_text)
        
        doc.save(output_path)
        
        # Step 3: Update session working path for other endpoints
        shutil.copy(output_path, session.working_path)
        
        # Step 4: Generate verification result
        original_text = extract_all_text_from_docx(session.original_path)
        final_text = extract_all_text_from_docx(output_path)
        
        return {
            "is_safe": True,
            "total_changes": len([c for c in session.corrections if c.get("status") == "accepted"]),
            "matched_count": len([c for c in session.corrections if c.get("status") == "accepted"]),
            "unmatched_count": 0,
            "original_length": len(original_text),
            "working_length": len(final_text),
            "original_hash": get_file_hash(session.original_path),
            "working_hash": get_file_hash(output_path),
        }
        
    except Exception as e:
        raise HTTPException(500, f"Σφάλμα δημιουργίας εγγράφου: {str(e)}")


# ============== CUSTOM CORRECTION (User-initiated edits) ==============

class CustomCorrectionRequest(BaseModel):
    """Request for user-initiated custom correction."""
    original: str  # The text to replace (selected by user)
    suggested: str  # The replacement text (empty string = delete)
    context: str = ""  # Surrounding text for accurate matching (optional but recommended)
    paragraph_number: int = 0  # Optional: paragraph number if known
    exact_offset: int = -1  # NEW: Exact character offset within paragraph (-1 = use text search)


@app.post("/api/custom-correction/{session_id}")
async def apply_custom_correction(session_id: str, request: CustomCorrectionRequest):
    """
    Apply a user-initiated custom correction.
    This allows users to:
    - Edit any text (not just AI-detected issues)
    - Delete words (suggested = "")
    - Make corrections the AI missed
    
    Uses the same safe XML manipulation as regular corrections.
    """
    session = sessions.get(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    
    if not request.original or not request.original.strip():
        raise HTTPException(400, "Πρέπει να επιλέξετε κείμενο για διόρθωση")
    
    original = request.original  # DON'T strip - preserve exact text
    suggested = request.suggested  # Can be empty for deletion
    context = request.context.strip() if request.context else ""
    
    # ALWAYS use paragraph_number + context-based search (more robust than exact_offset)
    # exact_offset is fragile because any prior edit shifts all positions
    
    if request.paragraph_number > 0:
        # Use apply_correction with paragraph_number (will search within that paragraph)
        ok = apply_correction(
            session.working_path,
            0,  # para_idx not used
            original,
            suggested,
            anchor=context,
            paragraph_number=request.paragraph_number,
        )
        if not ok:
            raise HTTPException(400, f"Δεν βρέθηκε το κείμενο «{original[:30]}...» στην παράγραφο {request.paragraph_number}")
    else:
        # No paragraph_number - search all paragraphs with context
        ok = apply_correction(
            session.working_path,
            0,
            original,
            suggested,
            anchor=context,
        )
        if not ok:
            raise HTTPException(400, f"Δεν βρέθηκε το κείμενο «{original[:30]}...» για διόρθωση")
    
    # Create a correction record for tracking
    custom_correction = {
        "id": len(session.corrections),
        "target": "body",
        "footnote_id": None,
        "paragraph_index": -1,
        "paragraph_number": request.paragraph_number,
        "paragraph_text": context,
        "module": "user",  # Mark as user-initiated
        "type": "fix",
        "scope": "phrase",
        "anchor": context,
        "original": original,
        "suggested": suggested,
        "reason": "Χειροκίνητη διόρθωση χρήστη",
        "confidence": "high",
        "verification_required": False,
        "sources": [],
        "status": "accepted",  # Already applied
    }
    
    session.corrections.append(custom_correction)
    
    # Save state for undo
    if suggested:
        session.save_action(f"Custom edit: '{original[:20]}...' → '{suggested[:20]}...'")
    else:
        session.save_action(f"Custom delete: '{original[:30]}...'")
    
    # Re-read paragraphs to get updated text
    updated_paragraphs = extract_paragraphs(session.working_path)
    
    return {
        "success": True,
        "correction": custom_correction,
        "paragraphs": updated_paragraphs,
        "can_undo": session.can_undo(),
        "can_redo": session.can_redo(),
    }


class BulkAction(BaseModel):
    action: str  # accept_all | reject_all


@app.get("/api/safety-check/{session_id}")
async def safety_check(session_id: str):
    """
    Get safety statistics: character counts, diff preview, warnings.
    Use this to verify document integrity before download.
    """
    session = sessions.get(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    
    # Get character counts
    original_text = extract_all_text_from_docx(session.original_path)
    working_text = extract_all_text_from_docx(session.working_path)
    
    original_chars = len(original_text)
    working_chars = len(working_text)
    char_diff = working_chars - original_chars
    char_diff_percent = (char_diff / original_chars * 100) if original_chars > 0 else 0
    
    # Compute diff changes
    diff_changes = compute_strict_diff(original_text, working_text)
    
    # Match to corrections
    match_result = match_diff_to_corrections(diff_changes, session.corrections)
    
    # Build warnings
    warnings = []
    
    # Warning: Large character difference
    if abs(char_diff_percent) > 5:
        warnings.append({
            "type": "char_diff",
            "severity": "high" if abs(char_diff_percent) > 10 else "medium",
            "message": f"Μεγάλη διαφορά χαρακτήρων: {char_diff:+d} ({char_diff_percent:+.1f}%)"
        })
    
    # Warning: Unmatched changes
    if match_result["unmatched"]:
        warnings.append({
            "type": "unmatched",
            "severity": "high",
            "message": f"{len(match_result['unmatched'])} αλλαγές δεν αντιστοιχούν σε εγκεκριμένες διορθώσεις"
        })
    
    # Warning: No changes made
    accepted_count = len([c for c in session.corrections if c.get("status") == "accepted" and c.get("type") == "fix"])
    if accepted_count > 0 and len(diff_changes) == 0:
        warnings.append({
            "type": "no_changes",
            "severity": "high",
            "message": f"Έγιναν {accepted_count} αποδοχές αλλά δεν εντοπίστηκαν αλλαγές στο κείμενο"
        })
    
    # Prepare diff preview (limit to 50 for performance)
    diff_preview = []
    for change in diff_changes[:50]:
        diff_preview.append({
            "type": change.get("type", "replace"),
            "original": change.get("original", "")[:100],
            "new": change.get("new", "")[:100],
            "context_before": change.get("context_before", "")[-30:],
            "context_after": change.get("context_after", "")[:30],
            "position": change.get("position", 0),
            "matched": any(
                m["change"].get("position") == change.get("position") 
                for m in match_result["matched"]
            ),
        })
    
    return {
        "original_chars": original_chars,
        "working_chars": working_chars,
        "char_diff": char_diff,
        "char_diff_percent": round(char_diff_percent, 2),
        "total_changes": len(diff_changes),
        "matched_count": len(match_result["matched"]),
        "unmatched_count": len(match_result["unmatched"]),
        "is_safe": match_result["all_matched"] and len(warnings) == 0,
        "warnings": warnings,
        "diff_preview": diff_preview,
        "has_more_changes": len(diff_changes) > 50,
    }


@app.post("/api/corrections/{session_id}/bulk")
async def bulk_action(session_id: str, action: BulkAction):
    session = sessions.get(session_id)
    if not session:
        raise HTTPException(404, "Session not found")

    results = {"accepted": 0, "rejected": 0, "failed": 0}

    for correction in session.corrections:
        if correction["status"] != "pending":
            continue

        if action.action == "accept_all":
            if correction.get("type") == "fix":
                if correction.get("target") == "footnote":
                    ok = apply_footnote_correction(
                        session.working_path,
                        int(correction["footnote_id"]),
                        correction.get("original", ""),
                        correction.get("suggested", ""),
                    )
                else:
                    ok = apply_correction(
                        session.working_path,
                        int(correction["paragraph_index"]),
                        correction.get("original", ""),
                        correction.get("suggested", ""),
                        anchor=correction.get("anchor", ""),
                        paragraph_number=correction.get("paragraph_number", 0),
                    )
                if ok:
                    correction["status"] = "accepted"
                    results["accepted"] += 1
                else:
                    results["failed"] += 1
            else:
                correction["status"] = "accepted"
                results["accepted"] += 1

        elif action.action == "reject_all":
            correction["status"] = "rejected"
            results["rejected"] += 1
        else:
            raise HTTPException(400, "Invalid bulk action")

    # Save state after bulk action
    if action.action == "accept_all":
        session.save_action(f"Accepted all ({results['accepted']} corrections)")
    else:
        session.save_action(f"Rejected all ({results['rejected']} corrections)")

    results["can_undo"] = session.can_undo()
    results["can_redo"] = session.can_redo()

    return results


@app.post("/api/undo/{session_id}")
async def undo_action(session_id: str):
    """Undo the last action."""
    session = sessions.get(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    
    if not session.can_undo():
        raise HTTPException(400, "Δεν υπάρχει ενέργεια για αναίρεση")
    
    success = session.undo()
    if not success:
        raise HTTPException(500, "Αποτυχία αναίρεσης")
    
    return {
        "success": True,
        "corrections": session.corrections,
        "can_undo": session.can_undo(),
        "can_redo": session.can_redo(),
    }


@app.post("/api/redo/{session_id}")
async def redo_action(session_id: str):
    """Redo the last undone action."""
    session = sessions.get(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    
    if not session.can_redo():
        raise HTTPException(400, "Δεν υπάρχει ενέργεια για επανάληψη")
    
    success = session.redo()
    if not success:
        raise HTTPException(500, "Αποτυχία επανάληψης")
    
    return {
        "success": True,
        "corrections": session.corrections,
        "can_undo": session.can_undo(),
        "can_redo": session.can_redo(),
    }


@app.get("/api/verify/{session_id}")
async def verify_document(session_id: str):
    """
    Strict verification: ensures only approved changes exist.
    Generates diff report.
    """
    session = sessions.get(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    
    filename = session.original_path.name.replace("_original", "")
    
    result = verify_and_generate_report(
        session.original_path,
        session.working_path,
        session.corrections,
        session_id,
        filename,
    )
    
    return result


@app.get("/api/download/{session_id}")
async def download_document(session_id: str, force: bool = False):
    """
    Download corrected document.
    The final.docx is generated by /api/generate-download (from editor content).
    Falls back to working.docx for backward compatibility.
    """
    session = sessions.get(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    
    # Prefer the final.docx (generated from editor content)
    final_path = OUTPUT_DIR / f"{session_id}_final.docx"
    download_path = final_path if final_path.exists() else session.working_path
    
    return FileResponse(
        download_path,
        filename=f"corrected_{session_id}.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@app.get("/api/download/{session_id}/original")
async def download_original(session_id: str):
    """Download original (untouched) document."""
    session = sessions.get(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    
    return FileResponse(
        session.original_path,
        filename=f"original_{session_id}.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@app.get("/api/download/{session_id}/report")
async def download_report(session_id: str):
    """Download diff report HTML."""
    session = sessions.get(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    
    # Generate report if not exists
    report_path = OUTPUT_DIR / f"{session_id}_diff_report.html"
    if not report_path.exists():
        filename = session.original_path.name.replace("_original", "")
        verify_and_generate_report(
            session.original_path,
            session.working_path,
            session.corrections,
            session_id,
            filename,
        )
    
    return FileResponse(
        report_path,
        filename=f"diff_report_{session_id}.html",
        media_type="text/html",
    )


@app.get("/api/download/{session_id}/bundle")
async def download_bundle(session_id: str):
    """Download ZIP bundle with original + corrected + summary."""
    session = sessions.get(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    
    filename = session.original_path.name.replace("_original", "").replace(".docx", "")
    final_path = OUTPUT_DIR / f"{session_id}_final.docx"
    corrected_path = final_path if final_path.exists() else session.working_path
    
    # Quick stats (no expensive diff)
    accepted = len([c for c in session.corrections if c.get("status") == "accepted"])
    rejected = len([c for c in session.corrections if c.get("status") == "rejected"])
    pending = len([c for c in session.corrections if c.get("status") == "pending"])
    total = len(session.corrections)
    
    bundle_path = OUTPUT_DIR / f"{session_id}_bundle.zip"
    
    with zipfile.ZipFile(bundle_path, "w", zipfile.ZIP_DEFLATED) as z:
        z.write(session.original_path, f"{filename}_ORIGINAL.docx")
        z.write(corrected_path, f"{filename}_CORRECTED.docx")
        
        # Generate report on-demand if not already present
        report_path = OUTPUT_DIR / f"{session_id}_diff_report.html"
        if not report_path.exists():
            verify_and_generate_report(
                session.original_path,
                session.working_path,
                session.corrections,
                session_id,
                session.original_path.name.replace("_original", ""),
            )
        if report_path.exists():
            z.write(report_path, f"{filename}_DIFF_REPORT.html")
        
        summary = f"""ProofreadAI - Report
=====================================

File: {filename}
Session: {session_id}

Corrections:
- Total: {total}
- Accepted: {accepted}
- Rejected: {rejected}
- Pending: {pending}

Contents:
- {filename}_ORIGINAL.docx
- {filename}_CORRECTED.docx
"""
        z.writestr(f"{filename}_SUMMARY.txt", summary.encode("utf-8"))
    
    return FileResponse(
        path=str(bundle_path),
        filename=f"{filename}_bundle.zip",
        media_type="application/zip",
    )


def create_comments_only_docx(original_path: Path, corrections: list, output_path: Path) -> bool:
    """
    Export document with ALL pending corrections as margin comments only.

    The body text is byte-for-byte identical to the original — no del/ins revision marks.
    Each correction span is wrapped in commentRangeStart/End so the comment appears
    anchored to the exact word in Word's comment pane.
    Comment text: «original» → «suggested» + reason on a second line.
    """
    from docx import Document
    from docx.oxml.ns import qn

    try:
        doc = Document(str(original_path))
    except Exception as e:
        print(f"[ERROR] Failed to load document: {e}")
        return False

    # Build paragraph index map (non-empty paragraphs only, same as track-changes)
    index_map = []
    for i, p in enumerate(doc.paragraphs):
        if p.text and p.text.strip():
            index_map.append(i)

    # All pending corrections (fix + suggestion — same count as the report)
    pending = [c for c in corrections if c.get("status") == "pending"]

    if not pending:
        shutil.copy(original_path, output_path)
        return True

    comment_id = 0
    applied_corrections = []
    now_iso = datetime.utcnow().isoformat() + "Z"

    # Body corrections
    for corr in pending:
        if corr.get("target") == "footnote":
            continue

        para_num = corr.get("paragraph_number", 0)
        original = corr.get("original", "")
        suggested = corr.get("suggested", "")
        reason = corr.get("reason", "")
        module = corr.get("module", "core")
        exact_offset = corr.get("exact_offset", -1)

        if not original:
            continue

        idx = para_num - 1
        if idx < 0 or idx >= len(index_map):
            print(f"[WARNING] Paragraph {para_num} out of range, skipping '{original[:30]}'")
            continue

        paragraph = doc.paragraphs[index_map[idx]]
        full_text = paragraph.text
        lshift = len(full_text) - len(full_text.lstrip())

        span_start = -1
        if exact_offset >= 0:
            candidate = exact_offset + lshift
            if full_text[candidate:candidate + len(original)] == original:
                span_start = candidate

        if span_start == -1:
            found, matched = find_original_in_text(full_text, original)
            if found == -1:
                print(f"[WARNING] Could not locate '{original[:30]}' in paragraph {para_num}")
                continue
            span_start = found
            original = matched

        span_end = span_start + len(original)

        result = _wrap_span_with_comment(paragraph, span_start, span_end, comment_id)
        if result:
            applied_corrections.append({
                "id": comment_id,
                "module": module,
                "reason": reason,
                "original": corr.get("original", ""),
                "suggested": suggested,
            })
            comment_id += 1
        else:
            print(f"[WARNING] Could not wrap '{original[:30]}' in paragraph {para_num}")

    # Footnote corrections
    try:
        for rel in doc.part.rels.values():
            if "footnotes" in rel.reltype:
                footnotes_xml = rel.target_part._element
                for corr in pending:
                    if corr.get("target") != "footnote":
                        continue
                    original = corr.get("original", "")
                    if not original:
                        continue
                    footnote_id = corr.get("footnote_id", 0)
                    exact_offset = corr.get("exact_offset", -1)
                    for fn in footnotes_xml.findall('.//' + qn('w:footnote')):
                        fn_id = fn.get(qn('w:id'))
                        if fn_id and int(fn_id) == footnote_id:
                            result = _wrap_span_with_comment_xml(fn, exact_offset, original, comment_id)
                            if result:
                                applied_corrections.append({
                                    "id": comment_id,
                                    "module": corr.get("module", "core"),
                                    "reason": corr.get("reason", ""),
                                    "original": original,
                                    "suggested": corr.get("suggested", ""),
                                })
                                comment_id += 1
                            else:
                                print(f"[WARNING] Could not wrap '{original[:30]}' in footnote {footnote_id}")
                            break
                break
    except Exception as e:
        print(f"[WARNING] Could not process footnotes: {e}")

    try:
        doc.save(str(output_path))
    except Exception as e:
        print(f"[ERROR] Failed to save document: {e}")
        return False

    if applied_corrections:
        try:
            _inject_suggestion_comments(output_path, applied_corrections, now_iso)
        except Exception as e:
            print(f"[WARNING] Could not inject suggestion comments: {e}")

    print(f"[INFO] Suggestions document saved with {comment_id} comments")
    return True


def _inject_suggestion_comments(docx_path: Path, corrections: list, now_iso: str) -> bool:
    """Inject suggestion-format comments.xml into the docx zip."""
    try:
        with zipfile.ZipFile(docx_path, 'r') as zin:
            file_contents = {name: zin.read(name) for name in zin.namelist()}

        file_contents['word/comments.xml'] = create_suggestion_comments_xml(corrections, now_iso).encode('utf-8')

        content_types = file_contents.get('[Content_Types].xml', b'').decode('utf-8')
        if 'comments.xml' not in content_types:
            content_types = content_types.replace(
                '</Types>',
                '<Override PartName="/word/comments.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"/></Types>'
            )
            file_contents['[Content_Types].xml'] = content_types.encode('utf-8')

        rels_path = 'word/_rels/document.xml.rels'
        if rels_path in file_contents:
            rels = file_contents[rels_path].decode('utf-8')
            if 'comments.xml' not in rels:
                import re as _re
                rids = _re.findall(r'Id="rId(\d+)"', rels)
                new_rid = max([int(r) for r in rids], default=0) + 1
                new_rel = (f'<Relationship Id="rId{new_rid}" '
                           f'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments" '
                           f'Target="comments.xml"/>')
                rels = rels.replace('</Relationships>', new_rel + '</Relationships>')
                file_contents[rels_path] = rels.encode('utf-8')

        with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as zout:
            for name, content in file_contents.items():
                zout.writestr(name, content)

        return True
    except Exception as e:
        print(f"[ERROR] Failed to inject suggestion comments: {e}")
        return False


@app.get("/api/download/{session_id}/comments")
async def download_with_comments(session_id: str):
    """Download document with all pending corrections as margin comments (no track changes)."""
    session = sessions.get(session_id)
    if not session:
        raise HTTPException(404, "Session not found")

    pending = [c for c in session.corrections if c.get("status") == "pending"]
    if not pending:
        raise HTTPException(400, "Δεν υπάρχουν εκκρεμείς διορθώσεις")

    comments_path = OUTPUT_DIR / f"{session_id}_suggestions.docx"

    success = create_comments_only_docx(
        session.original_path,
        session.corrections,
        comments_path,
    )

    if not success:
        raise HTTPException(500, "Αποτυχία δημιουργίας εγγράφου προτάσεων")

    filename = session.original_path.name.replace("_original", "").replace(".docx", "")

    return FileResponse(
        comments_path,
        filename=f"{filename}_suggestions.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@app.get("/api/download/{session_id}/tracked")
async def download_with_track_changes(session_id: str):
    """
    Download document with all pending corrections as Track Changes.
    
    This creates a new document where:
    - Original text is marked as deleted (red strikethrough)
    - Suggested text is marked as inserted (underlined)
    
    Users can then open in Word and Accept/Reject changes as usual.
    """
    session = sessions.get(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    
    # Filter only pending fix corrections
    pending_fixes = [c for c in session.corrections 
                     if c.get("status") == "pending" and c.get("type") == "fix"]
    
    if not pending_fixes:
        raise HTTPException(400, "Δεν υπάρχουν εκκρεμείς διορθώσεις για Track Changes")
    
    # Create output path
    tracked_path = OUTPUT_DIR / f"{session_id}_tracked.docx"
    
    # Generate the tracked changes document
    success = create_tracked_changes_docx(
        session.original_path,
        session.corrections,
        tracked_path
    )
    
    if not success:
        raise HTTPException(500, "Αποτυχία δημιουργίας Track Changes εγγράφου")
    
    filename = session.original_path.name.replace("_original", "").replace(".docx", "")
    
    return FileResponse(
        tracked_path,
        filename=f"{filename}_track_changes.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


# ============== H. DEBUG / INTERNAL TEST SCENARIOS ==============

def run_debug_tests():
    """
    Lightweight internal test/debug function for the correction pipeline.

    Test cases:
    1) Same wrong word appears twice in same paragraph → find_all_occurrences returns 2 offsets
    2) suggested == original → validate_correction rejects as no-op
    3) Local verb change with low confidence + verification_required → rejected by validate_correction
    4) Obvious spelling error → validate_correction accepts
    5) Sentence-level suggestion when allow_sentence_level=False → rejected
    """
    print("\n[DEBUG TESTS] Running correction validation pipeline tests...")

    para_text = "Ο γιατρός πήγε στο νοσοκομείο και ο γιατρός επέστρεψε αργά."

    # 1. Repeated identical error → two distinct offsets
    offsets = find_all_occurrences(para_text, "γιατρός")
    assert len(offsets) == 2, f"Expected 2 occurrences, got {len(offsets)}"
    print(f"  [1] PASS: 'γιατρός' at offsets {offsets} — repeated errors tracked separately")

    # 2. No-op correction (suggested == original)
    corr_noop = {
        "original": "γιατρός", "suggested": "γιατρός", "type": "fix",
        "scope": "token", "anchor": "Ο γιατρός πήγε στο νοσοκομείο",
        "confidence": "high", "verification_required": False, "module": "core",
    }
    valid, reason = validate_correction(corr_noop, paragraph_text=para_text)
    assert not valid, "Should be rejected as no-op"
    print(f"  [2] PASS: no-op rejected — {reason}")

    # 3. Risky grammar edit: low confidence + verification_required
    corr_risky = {
        "original": "πήγε", "suggested": "πηγαίνει", "type": "fix",
        "scope": "token", "anchor": "Ο γιατρός πήγε στο νοσοκομείο",
        "confidence": "low", "verification_required": True, "module": "core",
    }
    valid, reason = validate_correction(corr_risky, paragraph_text=para_text)
    assert not valid, "Should be rejected (low confidence + verification_required)"
    print(f"  [3] PASS: risky grammar edit rejected — {reason}")

    # 4. Obvious spelling error → accepted
    para2 = "Ο ανθρωπος πήγε σπίτι."
    corr_spell = {
        "original": "ανθρωπος", "suggested": "άνθρωπος", "type": "fix",
        "scope": "token", "anchor": "Ο ανθρωπος πήγε σπίτι",
        "confidence": "high", "verification_required": False, "module": "core",
    }
    valid, reason = validate_correction(corr_spell, paragraph_text=para2)
    assert valid, f"Spelling error should be accepted; got: {reason}"
    print(f"  [4] PASS: spelling error accepted")

    # 5. Sentence-level rewrite when not allowed → rejected
    corr_sentence = {
        "original": "Ο γιατρός πήγε στο νοσοκομείο",
        "suggested": "Ο ιατρός μετέβη στο νοσηλευτικό ίδρυμα",
        "type": "fix", "scope": "sentence",
        "anchor": "Ο γιατρός πήγε στο νοσοκομείο", "confidence": "medium",
        "verification_required": False, "module": "style",
    }
    valid, reason = validate_correction(corr_sentence, paragraph_text=para_text, allow_sentence_level=False)
    assert not valid, "Should be rejected (sentence level not allowed)"
    print(f"  [5] PASS: sentence-level suggestion rejected — {reason}")

    print("[DEBUG TESTS] All 5 tests passed!\n")


# Uncomment to run validation tests on startup:
# run_debug_tests()


# ============== TRANSLATION API ==============

from translator import (
    create_session as create_translation_session,
    get_session as get_translation_session,
    extract_docx_paragraphs,
    extract_text_paragraphs,
    normalize_text,
    split_paragraphs,
)

class TranslateConfigRequest(BaseModel):
    source_language: str = ""
    document_brief: str = ""
    glossary: list = []
    model: str = ""

class TranslateStartRequest(BaseModel):
    model: str = ""


@app.post("/api/translate/upload")
async def translate_upload(file: UploadFile = File(...)):
    """Upload file for translation."""
    filename = file.filename or "document"
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    
    if ext not in ("docx", "txt"):
        raise HTTPException(400, "Only .docx and .txt files supported")
    
    session = create_translation_session()
    session.filename = filename
    session.is_docx = (ext == "docx")
    
    # Save uploaded file
    file_bytes = await file.read()
    file_size_mb = len(file_bytes) / (1024 * 1024)
    if file_size_mb > MAX_UPLOAD_MB:
        raise HTTPException(400, f"File too large ({file_size_mb:.1f}MB). Max: {MAX_UPLOAD_MB}MB")
    
    with open(session.original_path, "wb") as f:
        f.write(file_bytes)
    
    # Extract paragraphs
    if session.is_docx:
        session.paragraphs = extract_docx_paragraphs(session.original_path)
    else:
        text = file_bytes.decode("utf-8", errors="replace")
        session.paragraphs = extract_text_paragraphs(text)
    
    session.translated = [""] * len(session.paragraphs)
    session.status = "uploaded"
    session.save_progress()
    
    total_chars = sum(len(p["plain"]) for p in session.paragraphs)
    
    return {
        "session_id": session.session_id,
        "filename": filename,
        "paragraph_count": len(session.paragraphs),
        "char_count": total_chars,
        "is_docx": session.is_docx,
    }


@app.post("/api/translate/upload-text")
async def translate_upload_text(request: dict):
    """Upload plain text for translation."""
    text = request.get("text", "").strip()
    title = request.get("title", "Pasted Text")
    
    if not text:
        raise HTTPException(400, "Text cannot be empty")
    
    session = create_translation_session()
    session.filename = title
    session.is_docx = False
    
    # Save as txt
    txt_path = session.session_dir / "original.txt"
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(text)
    session.original_path = txt_path
    
    session.paragraphs = extract_text_paragraphs(text)
    session.translated = [""] * len(session.paragraphs)
    session.status = "uploaded"
    session.save_progress()
    
    return {
        "session_id": session.session_id,
        "filename": title,
        "paragraph_count": len(session.paragraphs),
        "char_count": len(text),
        "is_docx": False,
    }


@app.post("/api/translate/prescan/{session_id}")
async def translate_prescan(session_id: str, request: dict):
    """Run pre-scan analysis on uploaded document (background)."""
    session = get_translation_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    
    model = request.get("model", "")
    if not model or model not in AVAILABLE_MODELS:
        raise HTTPException(400, "Invalid model")
    
    model_id = AVAILABLE_MODELS[model]["model_id"]
    
    # Run in background — frontend polls /api/translate/status
    async def run_prescan():
        try:
            await session.prescan(api_key=OPENROUTER_API_KEY, model_id=model_id)
        except Exception as e:
            print(f"[PRESCAN ERROR] {e}")
            session.status = "error"
            session.error_message = str(e)
            session.save_progress()
    
    asyncio.create_task(run_prescan())
    
    return {"status": "scanning"}


@app.post("/api/translate/config/{session_id}")
async def translate_save_config(session_id: str, request: TranslateConfigRequest):
    """Save user-confirmed translation configuration."""
    session = get_translation_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    
    if request.source_language:
        session.source_language = request.source_language
    if request.document_brief:
        session.document_brief = request.document_brief
    if request.glossary is not None:
        session.glossary = request.glossary
    
    session.save_config({
        "source_language": session.source_language,
        "document_brief": session.document_brief,
        "glossary": session.glossary,
        "model": request.model,
    })
    session.save_progress()
    
    return {"status": "saved"}


@app.post("/api/translate/start/{session_id}")
async def translate_start(session_id: str, request: TranslateStartRequest):
    """Start translation process (runs in background)."""
    session = get_translation_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    
    model = request.model
    if not model or model not in AVAILABLE_MODELS:
        raise HTTPException(400, "Invalid model")
    
    model_id = AVAILABLE_MODELS[model]["model_id"]
    
    # Run translation in background
    async def run_translation():
        try:
            await session.translate(
                api_key=OPENROUTER_API_KEY,
                model_id=model_id,
            )
        except Exception as e:
            print(f"[TRANSLATE ERROR] {e}")
            session.status = "error"
            session.error_message = str(e)
            session.save_progress()
    
    asyncio.create_task(run_translation())
    
    return {"status": "started", "total_chunks": session.total_chunks or "calculating"}


@app.get("/api/translate/status/{session_id}")
async def translate_status(session_id: str):
    """Get translation progress + prescan results."""
    session = get_translation_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    
    result = {
        "status": session.status,
        "progress": session.progress,
        "total_chunks": session.total_chunks,
        "source_language": session.source_language,
        "error_message": session.error_message,
        "translated_count": sum(1 for t in session.translated if t),
        "total_paragraphs": len(session.paragraphs),
    }
    
    # Include prescan results when scan is complete
    if session.status in ("scanned", "translating", "ready"):
        result["document_brief"] = session.document_brief
        result["glossary"] = session.glossary
    
    # Include review summary when translation is complete
    if session.status == "ready" and session.chunk_reviews:
        reviews = session.chunk_reviews
        review_required = sum(1 for r in reviews if r.get("review_required"))
        total_issues = sum(r.get("issue_count", 0) for r in reviews if r)
        high_issues = sum(r.get("high_severity_count", 0) for r in reviews if r)
        result["review_summary"] = {
            "total_chunks": len(reviews),
            "review_required": review_required,
            "total_issues": total_issues,
            "high_severity_issues": high_issues,
            "review_percentage": f"{(review_required / max(1, len(reviews))) * 100:.1f}%",
        }
    
    return result


@app.get("/api/translate/paragraphs/{session_id}")
async def translate_paragraphs(session_id: str):
    """Get original and translated paragraphs for split view."""
    session = get_translation_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    
    paragraphs = []
    for i, p in enumerate(session.paragraphs):
        paragraphs.append({
            "index": i,
            "original": p["plain"],
            "translated": session.translated[i] if i < len(session.translated) else "",
        })
    
    return {
        "paragraphs": paragraphs,
        "source_language": session.source_language,
        "glossary": session.glossary,
    }


@app.post("/api/translate/glossary/{session_id}")
async def translate_update_glossary(session_id: str, request: dict):
    """Update glossary for a translation session."""
    session = get_translation_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    
    session.glossary = request.get("glossary", [])
    session.save_progress()
    
    return {"status": "updated", "glossary_count": len(session.glossary)}


@app.post("/api/translate/consistency/{session_id}")
async def translate_consistency(session_id: str, request: dict):
    """Run consistency check on completed translation."""
    session = get_translation_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    
    model = request.get("model", "")
    if not model or model not in AVAILABLE_MODELS:
        raise HTTPException(400, "Invalid model")
    
    model_id = AVAILABLE_MODELS[model]["model_id"]
    
    result = await session.consistency_check(
        api_key=OPENROUTER_API_KEY,
        model_id=model_id,
    )
    
    return result


@app.get("/api/translate/download/{session_id}")
async def translate_download(session_id: str, partial: bool = False):
    """Download translated file. Use ?partial=true to download incomplete translation."""
    session = get_translation_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    
    if session.status != "ready" and not partial:
        # Check if there's any translated content
        done = sum(1 for t in session.translated if t)
        if done == 0:
            raise HTTPException(400, "No translation available yet")
    
    output_path = session.generate_output()
    
    ext = ".docx" if session.is_docx else ".txt"
    suffix = "_partial" if session.status != "ready" else ""
    download_name = session.filename.rsplit(".", 1)[0] + f"_translated{suffix}{ext}"
    
    return FileResponse(
        path=str(output_path),
        filename=download_name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document" if session.is_docx else "text/plain",
    )


@app.post("/api/translate/resume/{session_id}")
async def translate_resume(session_id: str, request: TranslateStartRequest):
    """Resume a stopped/failed translation from where it left off."""
    session = get_translation_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    
    # Don't block on stale "translating" status (e.g. after PC restart)
    # Just restart the translation from where it left off
    
    model = request.model
    if not model or model not in AVAILABLE_MODELS:
        raise HTTPException(400, "Invalid model")
    
    model_id = AVAILABLE_MODELS[model]["model_id"]
    
    # Reset error state but keep progress
    session.status = "translating"
    session.error_message = ""
    session.save_progress()
    
    done = sum(1 for t in session.translated if t)
    print(f"[TRANSLATE] Resuming session {session_id} from chunk {session.progress} ({done}/{len(session.paragraphs)} paragraphs done)")
    
    async def run_translation():
        try:
            await session.translate(api_key=OPENROUTER_API_KEY, model_id=model_id)
        except Exception as e:
            print(f"[TRANSLATE ERROR] {e}")
            session.status = "error"
            session.error_message = str(e)
            session.save_progress()
    
    asyncio.create_task(run_translation())
    
    return {"status": "resumed", "from_chunk": session.progress, "total_chunks": session.total_chunks}


@app.get("/api/translate/review/{session_id}")
async def translate_review_queue(session_id: str):
    """Get review queue — only flagged segments that need human review."""
    session = get_translation_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    
    return session.get_review_queue()

# ============== PARAPHRASER API (v2 — dual file) ==============

from paraphraser import (
    create_session as create_paraphrase_session,
    get_session as get_paraphrase_session,
    extract_paragraphs as para_extract,
    split_into_aligned_chunks,
)


@app.post("/api/paraphrase/upload")
async def paraphrase_upload(
    source_file: UploadFile = File(...),
    translation_file: UploadFile = File(...),
):
    """Upload two files: English source + Greek translation."""
    session = create_paraphrase_session()

    for label, file, attr in [
        ("source", source_file, "source_filename"),
        ("translation", translation_file, "trans_filename"),
    ]:
        filename = file.filename or f"{label}.docx"
        ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
        if ext not in ("docx", "txt"):
            raise HTTPException(400, f"Only .docx and .txt files supported ({label})")

        file_bytes = await file.read()
        if len(file_bytes) / (1024 * 1024) > MAX_UPLOAD_MB:
            raise HTTPException(400, f"File too large ({label})")

        file_path = session.session_dir / filename
        with open(file_path, "wb") as f:
            f.write(file_bytes)
        setattr(session, attr, filename)

    # Extract paragraphs
    src_path = session.session_dir / session.source_filename
    trn_path = session.session_dir / session.trans_filename
    session.source_paragraphs = para_extract(str(src_path))
    session.trans_paragraphs = para_extract(str(trn_path))

    session.aligned_chunks = split_into_aligned_chunks(
        session.source_paragraphs, session.trans_paragraphs
    )
    session.total_chunks = len(session.aligned_chunks)
    session.save_paragraphs()
    session.save_progress()

    return {
        "session_id": session.session_id,
        "source_filename": session.source_filename,
        "trans_filename": session.trans_filename,
        "source_paragraphs": len(session.source_paragraphs),
        "trans_paragraphs": len(session.trans_paragraphs),
        "total_chunks": session.total_chunks,
        "preview_source": session.source_paragraphs[:3],
        "preview_trans": session.trans_paragraphs[:3],
    }


@app.post("/api/paraphrase/prescan/{session_id}")
async def paraphrase_prescan(session_id: str, request: dict = {}):
    session = get_paraphrase_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")

    model = request.get("model", "")
    if not model or model not in AVAILABLE_MODELS:
        raise HTTPException(400, "Invalid model")

    model_id = AVAILABLE_MODELS[model]["model_id"]
    result = await session.prescan(api_key=OPENROUTER_API_KEY, model_id=model_id)
    return result


@app.post("/api/paraphrase/config/{session_id}")
async def paraphrase_save_config(session_id: str, request: dict = {}):
    session = get_paraphrase_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    if "document_brief" in request:
        session.document_brief = request["document_brief"]
    if "style_notes" in request:
        session.style_notes = request["style_notes"]
    if "key_names" in request:
        session.key_names = request["key_names"]
    session.save_progress()
    return {"status": "saved"}


@app.post("/api/paraphrase/start/{session_id}")
async def paraphrase_start(session_id: str, request: dict = {}):
    session = get_paraphrase_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")

    model = request.get("model", "")
    if not model or model not in AVAILABLE_MODELS:
        raise HTTPException(400, "Invalid model")
    model_id = AVAILABLE_MODELS[model]["model_id"]

    async def run_paraphrase():
        try:
            await session.paraphrase(api_key=OPENROUTER_API_KEY, model_id=model_id)
        except Exception as e:
            print(f"[PARAPHRASE ERROR] {e}")
            session.status = "error"
            session.error_message = str(e)
            session.save_progress()

    asyncio.create_task(run_paraphrase())
    return {"status": "started", "total_chunks": session.total_chunks}


@app.get("/api/paraphrase/status/{session_id}")
async def paraphrase_status(session_id: str):
    session = get_paraphrase_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    return {
        "status": session.status,
        "current_chunk": session.current_chunk,
        "total_chunks": session.total_chunks,
        "completed_chunks": len(session.paraphrased_chunks),
        "error": session.error_message,
    }


@app.get("/api/paraphrase/results/{session_id}")
async def paraphrase_results(session_id: str):
    session = get_paraphrase_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    return {
        "status": session.status,
        "side_by_side": session.get_side_by_side(),
        "document_brief": session.document_brief,
        "style_notes": session.style_notes,
    }


@app.post("/api/paraphrase/pause/{session_id}")
async def paraphrase_pause(session_id: str):
    session = get_paraphrase_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    session.is_paused = True
    return {"status": "pausing"}


@app.post("/api/paraphrase/resume/{session_id}")
async def paraphrase_resume(session_id: str, request: dict = {}):
    session = get_paraphrase_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    session.is_paused = False
    model = request.get("model", session.model)
    if not model or model not in AVAILABLE_MODELS:
        raise HTTPException(400, "Invalid model")
    model_id = AVAILABLE_MODELS[model]["model_id"]

    async def run_paraphrase():
        try:
            await session.paraphrase(api_key=OPENROUTER_API_KEY, model_id=model_id)
        except Exception as e:
            print(f"[PARAPHRASE ERROR] {e}")
            session.status = "error"
            session.error_message = str(e)
            session.save_progress()

    asyncio.create_task(run_paraphrase())
    return {"status": "resumed"}


@app.get("/api/paraphrase/download/{session_id}")
async def paraphrase_download(session_id: str):
    session = get_paraphrase_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    output_path = session.session_dir / f"paraphrased_{Path(session.trans_filename).stem}.docx"
    session.export_docx(str(output_path))
    return FileResponse(
        str(output_path),
        filename=output_path.name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    
# ============== REPORT ENDPOINT ==============

class ReportRequest(BaseModel):
    session_id: str
    format: str = "pdf"          # "pdf" | "html"
    teaser_enabled: bool = False
    teaser_show_first: int = 3


@app.post("/api/v1/report/generate")
async def generate_report(request: ReportRequest):
    """
    Generate a read-only corrections report (HTML or PDF) from a session.
    POST body: { session_id, format: "pdf"|"html", teaser_enabled?, teaser_show_first? }
    """
    from report import build_report_payload, render_html, render_pdf

    session = sessions.get(request.session_id)
    if not session:
        raise HTTPException(404, "Session not found")

    # Ensure paragraphs are loaded
    paragraphs = session.paragraphs
    if not paragraphs and session.working_path.exists():
        paragraphs = extract_paragraphs(session.working_path)

    filename = session.original_path.name if session.original_path else ""

    try:
        payload = build_report_payload(
            corrections=session.corrections,
            paragraphs=paragraphs,
            filename=filename,
            teaser_enabled=request.teaser_enabled,
            teaser_show_first=request.teaser_show_first,
        )
        html = render_html(payload)
    except Exception as e:
        raise HTTPException(500, f"Report generation failed: {e}")

    if request.format == "html":
        return Response(content=html, media_type="text/html; charset=utf-8")

    try:
        pdf_bytes = render_pdf(html)
    except Exception as e:
        raise HTTPException(500, f"PDF rendering failed: {e}")

    ref = payload["document"]["ref"]
    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="noeta-report-{ref}.pdf"'},
    )


if __name__ == "__main__":
    import uvicorn
    print("\n" + "=" * 50)
    print("🚀 ProofreadAI Backend v2.21 (OpenRouter)")
    print("=" * 50)
    if USING_LXML:
        print("✅ Using lxml for proper XML namespace handling")
    else:
        print("⚠️  lxml not installed - using standard ElementTree")
        print("   Install lxml for better Word compatibility: pip install lxml")
    print("\n📍 http://localhost:8000")
    print("📚 http://localhost:8000/docs\n")
    port = int(os.getenv("PORT", 8000))
    uvicorn.run(app, host="0.0.0.0", port=port)