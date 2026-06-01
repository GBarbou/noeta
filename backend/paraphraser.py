"""
ProofreadAI Paraphraser Engine v2.0
=====================================
Dual-file: English source + existing Greek translation.
Produces alternative Greek translation faithful to English but distinct from existing.
"""

import os
import re
import json
import time
import uuid
import asyncio
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from datetime import datetime

from openai import OpenAI

try:
    import docx
    from docx.oxml.ns import qn
except ImportError:
    docx = None


# ============== CONFIGURATION ==============

PARAPHRASE_DIR = Path("./paraphrase_sessions")
PARAPHRASE_DIR.mkdir(exist_ok=True)

MAX_CHARS_PER_CHUNK = 4000
CONTEXT_CHUNKS = 2
PAUSE_BETWEEN_CALLS = 0.5
MAX_RETRIES = 3


# ============== PROMPTS ==============

PRESCAN_PROMPT = """Αναλύεις δύο κείμενα: ένα αγγλικό πρωτότυπο και μια υπάρχουσα ελληνική μετάφρασή του.

ΔΕΙΓΜΑ ΑΓΓΛΙΚΟΥ:
{english_sample}

ΔΕΙΓΜΑ ΕΛΛΗΝΙΚΟΥ:
{greek_sample}

Απάντησε σε JSON:

{{
  "genre": "το είδος (π.χ. μυθιστόρημα, βιογραφία, ιστορικό)",
  "register": "το ύφος της ελληνικής μετάφρασης (λογοτεχνικό, ακαδημαϊκό κλπ)",
  "era_style": "εκτίμηση εποχής μετάφρασης (μεσοπολεμική, μεταπολεμική, σύγχρονη)",
  "key_names": ["κύρια ονόματα και η ελληνική απόδοσή τους"],
  "style_notes": "παρατηρήσεις για το μεταφραστικό ύφος, ιδιωματισμούς, χαρακτηριστικές επιλογές",
  "translation_quality": "γενική εκτίμηση ποιότητας μετάφρασης"
}}
"""


def make_paraphrase_prompt(
    english_chunk: str,
    greek_chunk: str,
    document_brief: str,
    style_notes: str,
    context_before: str = "",
) -> str:
    """Generate the paraphrase prompt for a single chunk."""

    context_section = ""
    if context_before:
        context_section = f"""
ΠΡΟΗΓΟΥΜΕΝΟ ΚΕΙΜΕΝΟ ΤΗΣ ΝΕΑΣ ΑΠΟΔΟΣΗΣ ΣΟΥ (για συνέχεια ύφους — ΜΗΝ το επαναλάβεις):
---
{context_before[-1500:]}
---
"""

    return f"""Είσαι έμπειρος Έλληνας λογοτεχνικός επιμελητής. Σου δίνεται:
1. Μια ΥΠΑΡΧΟΥΣΑ ΕΛΛΗΝΙΚΗ ΜΕΤΑΦΡΑΣΗ — αυτή είναι η ΒΑΣΗ σου. Θα αλλάξεις ΕΛΑΧΙΣΤΑ.
2. Το ΑΓΓΛΙΚΟ ΠΡΩΤΟΤΥΠΟ — χρησιμοποίησέ το ΜΟΝΟ για να βεβαιωθείς ότι δεν αλλοιώνεις το νόημα.

ΣΤΟΧΟΣ: Κάνε τις ΕΛΑΧΙΣΤΕΣ ΔΥΝΑΤΕΣ αλλαγές ώστε το κείμενο να μην είναι πια λέξη-προς-λέξη αντίγραφο, αλλά να παραμένει πολύ κοντά στο πρωτότυπο. Σκέψου 1-2 λέξεις ανά πρόταση — ΟΧΙ ΠΕΡΙΣΣΟΤΕΡΕΣ.

ΠΛΗΡΟΦΟΡΙΕΣ ΕΓΓΡΑΦΟΥ:
{document_brief}

ΣΗΜΕΙΩΣΕΙΣ ΥΦΟΥΣ:
{style_notes}
{context_section}
ΤΙ ΑΛΛΑΖΕΙΣ (με φειδώ):
- Αντικαθιστάς 1-2 λέξεις ανά πρόταση με συνώνυμα ή ισοδύναμες εκφράσεις
- Διορθώνεις τυχόν ορθογραφικά/γραμματικά λάθη
- Εκσυγχρονίζεις λέξεις που ακούγονται παρωχημένες
- Ελέγχεις με το αγγλικό αν η μετάφραση έχει λάθη νοήματος — ΑΝ ΝΑΙ, διόρθωσέ τα

ΤΙ ΔΕΝ ΑΛΛΑΖΕΙΣ:
- ΜΗΝ αναδιατυπώνεις ολόκληρες προτάσεις
- ΜΗΝ αλλάζεις τη σειρά λέξεων ή προτάσεων
- ΜΗΝ αλλάζεις τη σύνταξη εκτός αν υπάρχει λάθος
- ΜΗΝ αλλάζεις μεταφορές ή εικόνες που δουλεύουν καλά
- ΜΗΝ ανεβάζεις ή κατεβάζεις το ύφος
- ΜΗΝ αλλάζεις ονόματα, τοπωνύμια, αριθμούς
- ΚΡΑΤΑ ακριβώς τον ίδιο αριθμό παραγράφων
- ΑΝ μια πρόταση είναι σωστή, φυσική και χωρίς πρόβλημα → ΑΦΗΣΕ ΤΗΝ ΩΣ ΕΧΕΙ

ΣΗΜΑΝΤΙΚΟ: Αν δεν βρίσκεις λόγο να αλλάξεις μια πρόταση, ΜΗΝ ΤΗΝ ΑΛΛΑΖΕΙΣ. Είναι ΑΠΟΛΥΤΩΣ αποδεκτό μερικές προτάσεις να μείνουν ακριβώς ίδιες.

ΠΑΡΑΔΕΙΓΜΑ:
Υπάρχουσα: «Ήταν ένα ήσυχο μεσημέρι, αρχή του καλοκαιριού, στα νότια βουνά της Αραβίας.»
Σωστό: «Ήταν ένα ήρεμο μεσημέρι, στις αρχές του καλοκαιριού, στα νότια βουνά της Αραβίας.»
(μόνο 2 μικρές αλλαγές: ήσυχο→ήρεμο, αρχή→στις αρχές)

ΛΑΘΟΣ: «Εκείνο το γαλήνιο μεσημέρι, στις αρχές του καλοκαιριού, στα νότια ορεινά της Αραβίας.»
(πάρα πολλές αλλαγές — ΑΠΑΓΟΡΕΥΕΤΑΙ)

Απάντησε ΜΟΝΟ με το κείμενο. Καμία εξήγηση.

ΥΠΑΡΧΟΥΣΑ ΕΛΛΗΝΙΚΗ ΜΕΤΑΦΡΑΣΗ (Η ΒΑΣΗ ΣΟΥ):
---
{greek_chunk}
---

ΑΓΓΛΙΚΟ ΠΡΩΤΟΤΥΠΟ (ΜΑΡΤΥΡΑΣ ΝΟΗΜΑΤΟΣ):
---
{english_chunk}
---"""


# ============== DOCUMENT EXTRACTION ==============

def extract_docx_paragraphs(filepath: str) -> List[str]:
    if docx is None:
        raise ImportError("python-docx is required")
    doc = docx.Document(filepath)
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    return paragraphs


def extract_text_paragraphs(filepath: str) -> List[str]:
    with open(filepath, "r", encoding="utf-8") as f:
        content = f.read()
    raw_paras = re.split(r'\n\s*\n', content)
    paragraphs = []
    for p in raw_paras:
        text = p.strip()
        if text:
            text = re.sub(r'\s+', ' ', text)
            paragraphs.append(text)
    return paragraphs


def extract_paragraphs(filepath: str) -> List[str]:
    if filepath.lower().endswith('.docx'):
        return extract_docx_paragraphs(filepath)
    else:
        return extract_text_paragraphs(filepath)


# ============== ALIGNED CHUNKING ==============

def split_into_aligned_chunks(
    source_paras: List[str],
    trans_paras: List[str],
    max_chars: int = MAX_CHARS_PER_CHUNK,
) -> List[Tuple[List[str], List[str]]]:
    """Split both paragraph lists into aligned chunks."""
    max_len = max(len(source_paras), len(trans_paras))
    src = source_paras + [""] * (max_len - len(source_paras))
    trn = trans_paras + [""] * (max_len - len(trans_paras))

    chunks = []
    current_src = []
    current_trn = []
    current_size = 0

    for i in range(max_len):
        para_size = len(trn[i]) + len(src[i])

        if para_size > max_chars * 2:
            if current_trn:
                chunks.append((current_src, current_trn))
                current_src = []
                current_trn = []
                current_size = 0
            chunks.append(([src[i]], [trn[i]]))
            continue

        if current_size + para_size > max_chars * 2 and current_trn:
            chunks.append((current_src, current_trn))
            current_src = []
            current_trn = []
            current_size = 0

        current_src.append(src[i])
        current_trn.append(trn[i])
        current_size += para_size

    if current_trn:
        chunks.append((current_src, current_trn))

    return chunks


# ============== API CALL ==============

async def call_llm(
    api_key: str,
    model_id: str,
    system_prompt: str,
    user_prompt: str,
    temperature: float = 0.6,
    max_tokens: int = 8000,
) -> str:
    client = OpenAI(
        base_url="https://openrouter.ai/api/v1",
        api_key=api_key,
    )
    for attempt in range(MAX_RETRIES):
        try:
            response = await asyncio.to_thread(
                client.chat.completions.create,
                model=model_id,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ],
                temperature=temperature,
                max_tokens=max_tokens,
            )
            return response.choices[0].message.content.strip()
        except Exception as e:
            if attempt < MAX_RETRIES - 1:
                wait = 2 ** attempt
                print(f"[PARAPHRASE] API error (attempt {attempt+1}): {e}, retrying in {wait}s...")
                await asyncio.sleep(wait)
            else:
                raise
    return ""


# ============== SESSION ==============

class ParaphraseSession:
    def __init__(self, session_id: str = None):
        self.session_id = session_id or str(uuid.uuid4())[:8]
        self.session_dir = PARAPHRASE_DIR / self.session_id
        self.session_dir.mkdir(exist_ok=True)

        self.source_filename = ""
        self.trans_filename = ""
        self.source_paragraphs: List[str] = []
        self.trans_paragraphs: List[str] = []
        self.aligned_chunks: List[Tuple[List[str], List[str]]] = []
        self.total_chunks = 0

        self.document_brief = ""
        self.style_notes = ""
        self.key_names: List[str] = []

        self.status = "uploaded"
        self.current_chunk = 0
        self.paraphrased_chunks: Dict[int, str] = {}
        self.error_message = ""
        self.is_paused = False
        self.model = ""

    def save_progress(self):
        state = {
            "session_id": self.session_id,
            "source_filename": self.source_filename,
            "trans_filename": self.trans_filename,
            "status": self.status,
            "current_chunk": self.current_chunk,
            "total_chunks": self.total_chunks,
            "document_brief": self.document_brief,
            "style_notes": self.style_notes,
            "key_names": self.key_names,
            "paraphrased_chunks": self.paraphrased_chunks,
            "model": self.model,
            "error_message": self.error_message,
        }
        with open(self.session_dir / "state.json", "w", encoding="utf-8") as f:
            json.dump(state, f, ensure_ascii=False, indent=2)

    def load_progress(self) -> bool:
        path = self.session_dir / "state.json"
        if not path.exists():
            return False
        with open(path, "r", encoding="utf-8") as f:
            state = json.load(f)
        self.source_filename = state.get("source_filename", "")
        self.trans_filename = state.get("trans_filename", "")
        self.status = state.get("status", "uploaded")
        self.current_chunk = state.get("current_chunk", 0)
        self.total_chunks = state.get("total_chunks", 0)
        self.document_brief = state.get("document_brief", "")
        self.style_notes = state.get("style_notes", "")
        self.key_names = state.get("key_names", [])
        self.paraphrased_chunks = {int(k): v for k, v in state.get("paraphrased_chunks", {}).items()}
        self.model = state.get("model", "")
        self.error_message = state.get("error_message", "")
        return True

    def save_paragraphs(self):
        with open(self.session_dir / "paragraphs.json", "w", encoding="utf-8") as f:
            json.dump({
                "source": self.source_paragraphs,
                "translation": self.trans_paragraphs,
            }, f, ensure_ascii=False, indent=2)

    def load_paragraphs(self) -> bool:
        path = self.session_dir / "paragraphs.json"
        if not path.exists():
            return False
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
        self.source_paragraphs = data.get("source", [])
        self.trans_paragraphs = data.get("translation", [])
        self.aligned_chunks = split_into_aligned_chunks(
            self.source_paragraphs, self.trans_paragraphs
        )
        self.total_chunks = len(self.aligned_chunks)
        return True

    async def prescan(self, api_key: str, model_id: str) -> dict:
        if not self.trans_paragraphs:
            self.load_paragraphs()

        prompt = PRESCAN_PROMPT.format(
            english_sample="\n\n".join(self.source_paragraphs[:10]),
            greek_sample="\n\n".join(self.trans_paragraphs[:10]),
        )

        result_text = await call_llm(
            api_key=api_key,
            model_id=model_id,
            system_prompt="Ειδικός φιλολογικής/μεταφραστικής ανάλυσης. Απαντάς ΜΟΝΟ σε JSON.",
            user_prompt=prompt,
            temperature=0.3,
        )

        try:
            cleaned = re.sub(r'^```json\s*', '', result_text)
            cleaned = re.sub(r'\s*```$', '', cleaned)
            result = json.loads(cleaned)
        except json.JSONDecodeError:
            result = {
                "genre": "unknown", "register": "λογοτεχνικό",
                "era_style": "unknown", "key_names": [],
                "style_notes": result_text[:500], "translation_quality": "",
            }

        self.document_brief = (
            f"Είδος: {result.get('genre', 'N/A')}. "
            f"Ύφος: {result.get('register', 'N/A')}. "
            f"Εποχή μετάφρασης: {result.get('era_style', 'N/A')}."
        )
        self.style_notes = result.get("style_notes", "")
        self.key_names = result.get("key_names", [])
        self.status = "prescanned"
        self.save_progress()

        return {
            "status": "prescanned",
            "analysis": result,
            "source_paragraphs": len(self.source_paragraphs),
            "trans_paragraphs": len(self.trans_paragraphs),
            "total_chunks": self.total_chunks,
            "document_brief": self.document_brief,
            "style_notes": self.style_notes,
        }

    async def paraphrase(self, api_key: str, model_id: str):
        if not self.trans_paragraphs:
            self.load_paragraphs()

        self.status = "paraphrasing"
        self.model = model_id
        self.save_progress()

        system_prompt = (
            "Είσαι έμπειρος Έλληνας λογοτεχνικός επιμελητής. Κάνεις ΕΛΑΧΙΣΤΕΣ αλλαγές "
            "σε υπάρχουσες μεταφράσεις — 1-2 λέξεις ανά πρόταση μέγιστο. Αν μια πρόταση "
            "δεν χρειάζεται αλλαγή, την αφήνεις ακριβώς ως έχει."
        )

        for i in range(self.current_chunk, self.total_chunks):
            if self.is_paused:
                self.status = "paused"
                self.save_progress()
                return

            self.current_chunk = i
            src_chunk, trn_chunk = self.aligned_chunks[i]

            context_before = ""
            for j in range(max(0, i - CONTEXT_CHUNKS), i):
                if j in self.paraphrased_chunks:
                    context_before += self.paraphrased_chunks[j] + "\n\n"

            prompt = make_paraphrase_prompt(
                english_chunk="\n\n".join(src_chunk),
                greek_chunk="\n\n".join(trn_chunk),
                document_brief=self.document_brief,
                style_notes=self.style_notes,
                context_before=context_before,
            )

            paraphrased = await call_llm(
                api_key=api_key, model_id=model_id,
                system_prompt=system_prompt,
                user_prompt=prompt, temperature=0.4,
            )

            orig_count = len([p for p in trn_chunk if p.strip()])
            para_count = len([p for p in paraphrased.split("\n\n") if p.strip()])

            if para_count != orig_count:
                print(f"[PARAPHRASE] Chunk {i}: paragraph mismatch ({orig_count} vs {para_count}), retrying...")
                retry_prompt = prompt + f"\n\nΠΡΟΣΟΧΗ: ΑΚΡΙΒΩΣ {orig_count} παραγράφους."
                paraphrased = await call_llm(
                    api_key=api_key, model_id=model_id,
                    system_prompt=system_prompt,
                    user_prompt=retry_prompt, temperature=0.5,
                )

            self.paraphrased_chunks[i] = paraphrased
            self.save_progress()
            print(f"[PARAPHRASE] Chunk {i+1}/{self.total_chunks} done")

            if i < self.total_chunks - 1:
                await asyncio.sleep(PAUSE_BETWEEN_CALLS)

        self.status = "completed"
        self.current_chunk = self.total_chunks
        self.save_progress()

    def get_side_by_side(self) -> List[dict]:
        results = []
        for i in range(self.total_chunks):
            _, trn_chunk = self.aligned_chunks[i]
            results.append({
                "chunk_index": i,
                "original": "\n\n".join(trn_chunk),
                "paraphrased": self.paraphrased_chunks.get(i, ""),
                "status": "done" if i in self.paraphrased_chunks else (
                    "processing" if i == self.current_chunk and self.status == "paraphrasing" else "pending"
                ),
            })
        return results

    def get_full_paraphrased_text(self) -> str:
        parts = []
        for i in range(self.total_chunks):
            if i in self.paraphrased_chunks:
                parts.append(self.paraphrased_chunks[i])
        return "\n\n".join(parts)

    def export_docx(self, output_path: str):
        if docx is None:
            raise ImportError("python-docx is required")
        doc = docx.Document()
        for para_text in self.get_full_paraphrased_text().split("\n\n"):
            para_text = para_text.strip()
            if para_text:
                doc.add_paragraph(para_text)
        doc.save(output_path)
        return output_path


# ============== SESSION MANAGEMENT ==============

_sessions: Dict[str, ParaphraseSession] = {}

def create_session() -> ParaphraseSession:
    session = ParaphraseSession()
    _sessions[session.session_id] = session
    return session

def get_session(session_id: str) -> Optional[ParaphraseSession]:
    if session_id in _sessions:
        return _sessions[session_id]
    session = ParaphraseSession(session_id)
    if session.load_progress():
        session.load_paragraphs()
        _sessions[session_id] = session
        return session
    return None
