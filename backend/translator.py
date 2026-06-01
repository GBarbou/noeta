"""
ProofreadAI Translation Engine v1.1
====================================
Professional-grade translation to Modern Greek for publishing.

Architecture:
  1. Pre-scan: Read full document → auto-detect language, generate document brief + glossary
  2. Translation: Chunk-by-chunk with document brief + glossary + sliding context window
  3. Verification: Each chunk verified against original (no additions/removals)
  4. Post-processing: Automatic glossary term enforcement (Python find-replace)
  5. Consistency check: Final pass across entire translation for term consistency

Design principles:
  - Source of truth: original document (never modified)
  - Every chunk sees: document brief + glossary + previous 2 translated chunks (context)
  - Glossary enforced via post-processing (reliable, not LLM-dependent)
  - Paragraph count strictly preserved
  - Crash-safe: progress saved after each chunk, resumable
"""

import os
import re
import json
import time
import uuid
import shutil
import asyncio
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from datetime import datetime

from openai import OpenAI

try:
    import docx
except ImportError:
    docx = None


# ============== CONFIGURATION ==============

TRANSLATION_DIR = Path("./translation_sessions")
TRANSLATION_DIR.mkdir(exist_ok=True)

MAX_CHARS_PER_CHUNK = 8000  # Balance between consistency and completeness
CONTEXT_CHUNKS = 2

PAUSE_BETWEEN_CALLS = 1.0
MAX_RETRIES = 5


# ============== PROMPTS ==============

def make_prescan_prompt(text_sample: str) -> str:
    return f"""Αναλύεις ένα κείμενο που πρόκειται να μεταφραστεί στα Νέα Ελληνικά.

Διάβασε προσεκτικά το παρακάτω κείμενο και απάντησε σε JSON μορφή:

{{
  "source_language": "η γλώσσα του κειμένου (π.χ. English, French, Ancient Greek, Katharevousa, Russian, Chinese...)",
  "document_type": "τύπος κειμένου (μυθιστόρημα, διήγημα, βιογραφία, ιστορικό, θεολογικό, λογοτεχνικό δοκίμιο, επιστημονικό, νομικό, δημοσιογραφικό...)",
  "subject": "σύντομη περιγραφή θέματος (1-2 προτάσεις)",
  "era_context": "εποχή/πλαίσιο αν είναι σχετικό",
  "style_register": "ύφος (λογοτεχνικό, ακαδημαϊκό, θεολογικό, δημοσιογραφικό, ανεπίσημο, τεχνικό, νομικό)",
  "narrative_tone": "αφηγηματικός τόνος αν υπάρχει (ουδέτερος, ειρωνικός, επικός, οικείος, λυρικός, τυπικός...)",
  "key_entities": ["σημαντικά ονόματα, τόποι, οργανισμοί που εμφανίζονται"],
  "suggested_glossary": [
    {{"source": "αρχικός όρος", "target": "προτεινόμενη ελληνική απόδοση", "note": "person/place/org/title/term"}},
  ],
  "translation_risks": {{
    "idioms_figurative": "yes/no + σύντομη εξήγηση",
    "archaic_register": "yes/no + σύντομη εξήγηση",
    "cultural_references": "yes/no + σύντομη εξήγηση",
    "wordplay_irony": "yes/no + σύντομη εξήγηση"
  }},
  "translation_notes": "ειδικές παρατηρήσεις/δυσκολίες για τη μετάφραση"
}}

Απάντησε ΜΟΝΟ σε JSON, χωρίς markdown backticks ή σχόλια.

ΚΕΙΜΕΝΟ:
{text_sample}"""


def make_system_prompt(
    source_language: str,
    document_brief: str,
    glossary: List[Dict],
    style_register: str = "ακαδημαϊκό",
) -> str:
    glossary_text = ""
    if glossary:
        glossary_text = "\n\nΥΠΟΧΡΕΩΤΙΚΟ ΓΛΩΣΣΑΡΙ (χρησιμοποίησε ΑΚΡΙΒΩΣ αυτές τις αποδόσεις — ΜΙΑ λέξη/φράση ανά όρο, ΟΧΙ εναλλακτικές με «ή»):\n"
        for entry in glossary:
            src = entry.get("source", "")
            tgt = entry.get("target", "")
            note = entry.get("note", "")
            if "/" in tgt:
                alternatives = [a.strip() for a in tgt.split("/")]
                tgt = alternatives[0]
                if not note:
                    note = f"εναλλακτικά: {', '.join(alternatives[1:])}"
            if " ή " in tgt:
                alternatives = [a.strip() for a in tgt.split(" ή ")]
                tgt = alternatives[0]
                if not note:
                    note = f"εναλλακτικά: {', '.join(alternatives[1:])}"
            glossary_text += f"  • {src} → {tgt}"
            if note:
                glossary_text += f"  ({note})"
            glossary_text += "\n"

    # Genre-specific instructions
    style_lower = (style_register or "").lower().strip()
    
    style_instructions = ""
    if any(kw in style_lower for kw in ["λογοτεχν", "μυθιστόρ", "διήγημ", "αφηγημ", "literary"]):
        style_instructions = """
ΕΙΔΙΚΟΙ ΚΑΝΟΝΕΣ ΛΟΓΟΤΕΧΝΙΚΗΣ ΜΕΤΑΦΡΑΣΗΣ:
ΙΕΡΑΡΧΗΣΗ:
  1. Πρώτα αποδίδεις σωστά το ακριβές νόημα.
  2. Μετά το αποδίδεις σε φυσικά ελληνικά.
  3. Τέλος διατηρείς όσο γίνεται τον λογοτεχνικό ρυθμό και τόνο.

- Μην αποδίδεις ιδιώματα, ειρωνεία ή μεταφορές κατά λέξη όταν αυτό παραπλανεί ή ακούγεται αφύσικο.
- Πριν μεταφράσεις δύσκολες φράσεις, σκέψου πρώτα τι σημαίνουν πραγματικά. Ιδιώματα, βιβλικές/πολιτισμικές νύξεις και ειρωνείες χρειάζονται ερμηνεία, όχι κατά λέξη απόδοση.
- Αν υπάρχει αμφισημία, προτίμησε τη σημασιολογικά ασφαλέστερη λύση.
- Τίτλοι κεφαλαίων: προτίμησε πιστή απόδοση· ελεύθερη μόνο αν η κατά λέξη λύση είναι αδέξια ή ακατανόητη.
- Μην εξομαλύνεις ιδιορρυθμίες του συγγραφέα. Μην χρησιμοποιείς ακαδημαϊκές/εξηγηματικές διατυπώσεις.
- Ελεύθερη απόδοση ΜΟΝΟ όταν η κατά λέξη λύση είναι αφύσικη ή παραπλανητική."""
    elif any(kw in style_lower for kw in ["θεολογ", "εκκλησ", "πατερ"]):
        style_instructions = """
ΕΙΔΙΚΟΙ ΚΑΝΟΝΕΣ ΘΕΟΛΟΓΙΚΗΣ ΜΕΤΑΦΡΑΣΗΣ:
- Σεβάσου την καθιερωμένη θεολογική και εκκλησιαστική ορολογία.
- Διατήρησε τον σοβαρό, ύψηλο τόνο χωρίς αρχαΐσματα.
- Βιβλικά παραθέματα: απέδωσέ τα σύμφωνα με τη νεοελληνική μετάφραση της Βίβλου."""
    elif any(kw in style_lower for kw in ["τεχν", "επιστημ", "technical", "scientific"]):
        style_instructions = """
ΕΙΔΙΚΟΙ ΚΑΝΟΝΕΣ ΤΕΧΝΙΚΗΣ/ΕΠΙΣΤΗΜΟΝΙΚΗΣ ΜΕΤΑΦΡΑΣΗΣ:
- Ακρίβεια ορολογίας πάνω από ύφος.
- Χρησιμοποίησε καθιερωμένη ελληνική ορολογία όπου υπάρχει.
- Αν δεν υπάρχει αντίστοιχος ελληνικός όρος, κράτησε τον ξενόγλωσσο με μεταγραφή."""
    elif any(kw in style_lower for kw in ["νομ", "legal"]):
        style_instructions = """
ΕΙΔΙΚΟΙ ΚΑΝΟΝΕΣ ΝΟΜΙΚΗΣ ΜΕΤΑΦΡΑΣΗΣ:
- Χρησιμοποίησε καθιερωμένη νομική ορολογία.
- Τυπικό ύφος χωρίς λογοτεχνικές ελευθερίες.
- Ακρίβεια στους νομικούς όρους και στη σύνταξη."""

    return f"""Είσαι επαγγελματίας μεταφραστής/επιμελητής ({source_language}→Νέα Ελληνικά) για έκδοση βιβλίου.

ΠΛΗΡΟΦΟΡΙΕΣ ΕΓΓΡΑΦΟΥ:
{document_brief}

Ύφος: {style_register}

Κανόνες (απαρέγκλιτοι):
1) ΜΗΝ προσθέτεις, ΜΗΝ αφαιρείς, ΜΗΝ συνοψίζεις, ΜΗΝ αναδιατάσσεις πληροφορίες. 
   Πιστή απόδοση νοήματος και γεγονότων.
2) Διατήρησε ΑΚΡΙΒΩΣ τον αριθμό παραγράφων και τη σειρά τους.
3) Διατήρησε ΟΛΑ τα κύρια ονόματα, ημερομηνίες, ποσά, τίτλους έργων/τραγουδιών/εταιρειών, 
   και νομικούς όρους με ακρίβεια.
4) Κείμενο φυσικό στα ελληνικά, χωρίς αγγλικισμούς και κατά λέξη μεταφορά αγγλικών ιδιωμάτων. 
   Απόδοση ιδιωμάτων σε καθιερωμένη ελληνική χρήση.
5) Διατήρησε ΑΚΡΙΒΩΣ τα tags μορφοποίησης: [B]...[/B], [I]...[/I], [U]...[/U]. 
   Μην τα αλλάξεις, μην τα μετακινήσεις, μην τα σβήσεις.
6) Μην βάζεις εισαγωγικά/σχόλια/σημειώσεις του μεταφραστή. Μόνο τελικό κείμενο.
7) ΕΙΣΑΓΩΓΙΚΑ: Χρησιμοποίησε ΜΟΝΟ ελληνικά εισαγωγικά « » (guillemets) για διαλόγους 
   και παραθέματα. ΠΟΤΕ " " ή " " ή ' '. Παράδειγμα: «Γεια σου» είπε, ΟΧΙ "Γεια σου".

Θέλω αποτέλεσμα έτοιμο για έκδοση (ορθογραφία/στίξη/ροή άρτια) χωρίς να αλλοιώνεται το περιεχόμενο.
{style_instructions}{glossary_text}"""


def make_translation_prompt(content: str, context: str = "") -> str:
    ctx_section = ""
    if context:
        ctx_section = f"""ΠΡΟΗΓΟΥΜΕΝΟ ΠΛΑΙΣΙΟ (για συνέχεια ύφους — ΜΗΝ το μεταφράσεις ξανά):
<<<
{context[-2000:]}
>>>

"""
    return f"""{ctx_section}Μετάφρασε το παρακάτω κείμενο τηρώντας τους κανόνες.
Δώσε ΜΟΝΟ τη μετάφραση, τίποτα άλλο.

ΚΕΙΜΕΝΟ ΠΡΟΣ ΜΕΤΑΦΡΑΣΗ:
<<<
{content}
>>>"""


def make_critique_prompt(source: str, translation: str, source_language: str) -> str:
    """Critique pass: returns structured JSON issues, NOT rewritten text."""
    return f"""Είσαι κριτικός επιμελητής μετάφρασης {source_language}→Ελληνικά για έκδοση.
Σου δίνω ΠΡΩΤΟΤΥΠΟ και ΜΕΤΑΦΡΑΣΗ. Ο ρόλος σου είναι ΜΟΝΟ να εντοπίσεις προβλήματα, ΟΧΙ να ξαναγράψεις.

ΕΛΕΓΞΕ:
1. ΠΙΣΤΟΤΗΤΑ: Λείπει κάτι; Προστέθηκε κάτι; Αλλοιώθηκε νόημα;
2. ΚΑΤΑ ΛΕΞΗ ΑΠΟΔΟΣΕΙΣ: Ιδιώματα/μεταφορές αποδόθηκαν μηχανικά αντί σημασιολογικά;
3. ΦΥΣΙΚΟΤΗΤΑ: Ακούγεται αφύσικο στα ελληνικά;
4. ΑΚΡΙΒΕΙΑ: Ονόματα, ημερομηνίες, γεγονότα σωστά;
5. ΜΟΡΦΟΠΟΙΗΣΗ: Tags [B]/[I]/[U], εισαγωγικά « », αριθμός παραγράφων.

ΚΑΝΟΝΕΣ:
- ΜΗΝ αναφέρεις υφολογικές "βελτιώσεις" σε σωστές φράσεις. ΜΟΝΟ πραγματικά προβλήματα.
- Αν η μετάφραση είναι καλή, επέστρεψε κενή λίστα issues.
- severity: "low" (μικρό), "medium" (αξιοσημείωτο), "high" (σοβαρό λάθος νοήματος)
- type: "mistranslation" | "omission" | "addition" | "literalism" | "register" | "formatting" | "factual"
- risk_type: "none" | "idiom" | "ambiguity" | "irony" | "allusion" | "cultural" | "factual" | "register"

Απάντησε ΜΟΝΟ σε JSON, χωρίς markdown backticks:
{{
  "issues": [
    {{
      "paragraph_index": <αριθμός παραγράφου στο chunk, ξεκινώντας από 1>,
      "severity": "low|medium|high",
      "type": "mistranslation|omission|addition|literalism|register|formatting|factual",
      "risk_type": "none|idiom|ambiguity|irony|allusion|cultural|factual|register",
      "source_span": "<η φράση στο πρωτότυπο>",
      "translated_span": "<πώς αποδόθηκε>",
      "reason": "<σύντομη εξήγηση>",
      "suggested_fix": "<προτεινόμενη διόρθωση>"
    }}
  ],
  "chunk_confidence": "high|medium|low"
}}

Αν δεν υπάρχουν προβλήματα: {{"issues": [], "chunk_confidence": "high"}}

ΠΡΩΤΟΤΥΠΟ ({source_language}):
<<<
{source}
>>>

ΜΕΤΑΦΡΑΣΗ (Ελληνικά):
<<<
{translation}
>>>"""


def make_apply_fixes_prompt(source: str, translation: str, issues_json: str, source_language: str) -> str:
    """Apply-fixes pass: fix ONLY the listed issues, nothing else."""
    return f"""Είσαι επιμελητής μετάφρασης {source_language}→Ελληνικά.
Σου δίνω ΠΡΩΤΟΤΥΠΟ, ΜΕΤΑΦΡΑΣΗ και ΛΙΣΤΑ ΠΡΟΒΛΗΜΑΤΩΝ σε JSON.

ΚΑΝΟΝΕΣ:
- Διόρθωσε ΜΟΝΟ τα προβλήματα που αναφέρονται στη λίστα.
- ΜΗΝ αλλάξεις ΤΙΠΟΤΑ άλλο. Αν μια πρόταση δεν αναφέρεται στα issues, κράτησέ την ακριβώς.
- Διατήρησε ΑΚΡΙΒΩΣ τον αριθμό και τη σειρά παραγράφων.
- Διατήρησε ΑΚΡΙΒΩΣ τα tags [B]/[I]/[U] και τα εισαγωγικά « ».

Δώσε ΜΟΝΟ τη διορθωμένη μετάφραση, χωρίς σχόλια.

ΠΡΟΒΛΗΜΑΤΑ ΠΡΟΣ ΔΙΟΡΘΩΣΗ:
{issues_json}

ΠΡΩΤΟΤΥΠΟ ({source_language}):
<<<
{source}
>>>

ΜΕΤΑΦΡΑΣΗ (Ελληνικά) ΠΡΟΣ ΔΙΟΡΘΩΣΗ:
<<<
{translation}
>>>"""


# Keep old verify as fallback (not used in main flow)
def make_verify_prompt(source: str, translation: str, source_language: str) -> str:
    """Legacy verify prompt — kept for backward compatibility."""
    return f"""Είσαι κριτικός επιμελητής μετάφρασης {source_language}→Ελληνικά για έκδοση.
Σου δίνω ΠΡΩΤΟΤΥΠΟ και ΜΕΤΑΦΡΑΣΗ.
Διόρθωσε ΜΟΝΟ πραγματικά προβλήματα (παραλείψεις, κατά λέξη αποδόσεις, λάθη). ΜΗΝ κάνεις rewrite.
Διατήρησε αριθμό παραγράφων, tags [B]/[I]/[U], εισαγωγικά « ».
Δώσε ΜΟΝΟ τη διορθωμένη μετάφραση.

ΠΡΩΤΟΤΥΠΟ ({source_language}):
<<<
{source}
>>>

ΜΕΤΑΦΡΑΣΗ:
<<<
{translation}
>>>"""


# ============== TEXT PROCESSING ==============

def normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)
    return text

def split_paragraphs(text: str) -> List[str]:
    text = normalize_text(text).strip()
    # First try splitting by blank lines (standard)
    parts = re.split(r"\n\s*\n", text)
    parts = [p.strip() for p in parts if p.strip()]
    
    # If we got only 1 huge paragraph, try splitting by single newlines
    # (common with pasted text or certain document formats)
    if len(parts) == 1 and len(parts[0]) > 1000:
        lines = [l.strip() for l in parts[0].split("\n") if l.strip()]
        if len(lines) > 1:
            parts = lines
    
    return parts

def chunk_paragraphs(paragraphs: List[str], max_chars: int = MAX_CHARS_PER_CHUNK) -> List[List[int]]:
    chunks: List[List[int]] = []
    current: List[int] = []
    cur_len = 0
    for i, p in enumerate(paragraphs):
        add_len = len(p) + 2
        if current and (cur_len + add_len > max_chars):
            chunks.append(current)
            current = [i]
            cur_len = len(p)
        else:
            current.append(i)
            cur_len += add_len
    if current:
        chunks.append(current)
    return chunks


# ============== POST-PROCESSING: GLOSSARY ENFORCEMENT ==============

def enforce_glossary(translated_texts: List[str], glossary: List[Dict]) -> List[str]:
    """
    Post-processing: enforce glossary terms via Python find-replace.
    This is 100% reliable unlike LLM-based enforcement.
    
    Also fixes the "X ή Y" problem where the LLM writes both alternatives
    from a glossary entry like "X/Y" instead of picking one.
    """
    if not glossary:
        return translated_texts
    
    result = list(translated_texts)
    replacements_made = 0
    
    for entry in glossary:
        target_raw = entry.get("target", "").strip()
        source = entry.get("source", "").strip()
        if not target_raw or not source:
            continue
        
        # Determine the canonical target (first alternative if "/" present)
        if "/" in target_raw:
            alternatives = [a.strip() for a in target_raw.split("/")]
            target = alternatives[0]
        elif " ή " in target_raw:
            alternatives = [a.strip() for a in target_raw.split(" ή ")]
            target = alternatives[0]
        else:
            alternatives = []
            target = target_raw
        
        for i, text in enumerate(result):
            if not text:
                continue
            
            original_text = text
            
            # Fix 1: Replace "X ή Y" patterns generated by LLM from glossary alternatives
            if alternatives and len(alternatives) >= 2:
                for j in range(len(alternatives)):
                    for k in range(len(alternatives)):
                        if j != k:
                            wrong_pattern = f"{alternatives[j]} ή {alternatives[k]}"
                            if wrong_pattern in text:
                                text = text.replace(wrong_pattern, target)
                                replacements_made += 1
            
            # Fix 2: Case-insensitive enforcement of the correct term
            pattern = re.compile(re.escape(target), re.IGNORECASE)
            matches = list(pattern.finditer(text))
            
            for match in reversed(matches):
                found = match.group()
                if found != target:
                    text = text[:match.start()] + target + text[match.end():]
                    replacements_made += 1
            
            result[i] = text
    
    if replacements_made > 0:
        print(f"[POST-PROCESS] Glossary enforcement: {replacements_made} replacements made")
    
    return result


def fix_quotation_marks(translated_texts: List[str]) -> List[str]:
    """Post-processing: normalize ALL quotation marks to Greek « ».
    
    Handles:
    - Smart/curly double quotes: " " → « »
    - Smart/curly single quotes used as dialogue: ' ' → « »  
    - Straight double quotes: "text" → «text»
    - Escaped or leftover ASCII quotes
    - Nested quotes: outer « », inner " " (standard Greek typography)
    - Already correct « » are preserved
    """
    result = []
    for text in translated_texts:
        if not text:
            result.append(text)
            continue
        t = text
        
        # Step 1: Convert smart/curly double quotes to « »
        t = t.replace('\u201c', '«').replace('\u201d', '»')  # " "
        
        # Step 2: Convert smart/curly single quotes used as dialogue markers to « »
        # But preserve apostrophes (single char between word chars, e.g. don't)
        # Replace opening ' followed by text and closing '
        t = re.sub(r'\u2018([^\u2019]{2,}?)\u2019', r'«\1»', t)
        
        # Step 3: Convert straight double quotes to « »
        # Match pairs of straight quotes with content between them
        # Use a loop to handle multiple pairs
        iterations = 0
        while '"' in t and iterations < 50:
            t_new = re.sub(r'"([^"]+?)"', r'«\1»', t, count=1)
            if t_new == t:
                break  # No more pairs found
            t = t_new
            iterations += 1
        
        # Step 4: Clean up any remaining lone straight quotes 
        # If a " is right before a word (likely opening), replace with «
        t = re.sub(r'"(\w)', r'«\1', t)
        # If a " is right after a word/punctuation (likely closing), replace with »
        t = re.sub(r'(\w|[.!?;,])"', r'\1»', t)
        
        # Step 5: Fix double guillemets from over-replacement
        t = t.replace('««', '«').replace('»»', '»')
        
        # Step 6: Fix spacing around guillemets (Greek typography)
        # No space after « and no space before »
        t = re.sub(r'«\s+', '«', t)
        t = re.sub(r'\s+»', '»', t)
        
        result.append(t)
    return result


def fix_midword_capitals(translated_texts: List[str]) -> List[str]:
    """Post-processing: fix stray uppercase letters mid-word.
    
    Some LLMs (especially Gemini) produce artifacts like "ανεπηρέαστΟς" 
    where a single Greek letter (usually Ο, Η, Α) appears uppercase 
    inside an otherwise lowercase word. This happens because:
    - Formatting tag boundaries confuse the model
    - The model mistakes single-char runs for articles (ο, η, α)
    - Unicode confusion between Latin O and Greek Ο
    
    Rules:
    - A Greek uppercase letter surrounded by lowercase Greek letters → lowercase
    - Exceptions: all-caps words, acronyms, first letter of sentence
    - Preserves legitimate capitalization (start of word after space/punctuation)
    """
    # Greek lowercase range
    GR_LOWER = 'αβγδεζηθικλμνξοπρσςτυφχψωάέήίόύώϊϋΐΰ'
    # Greek uppercase range
    GR_UPPER = 'ΑΒΓΔΕΖΗΘΙΚΛΜΝΞΟΠΡΣΣΤΥΦΧΨΩΆΈΉΊΌΎΏΪΫ'
    
    result = []
    fixes_made = 0
    
    for text in translated_texts:
        if not text:
            result.append(text)
            continue
        
        # Pattern: lowercase Greek + uppercase Greek + lowercase Greek
        # This catches "ανεπηρέαστΟς", "κάποιΗ", etc.
        def fix_midcap(m):
            nonlocal fixes_made
            fixes_made += 1
            return m.group(1) + m.group(2).lower() + m.group(3)
        
        t = re.sub(
            f'([{GR_LOWER}])([{GR_UPPER}])([{GR_LOWER}])',
            fix_midcap,
            text
        )
        
        # Also catch uppercase at the END of a lowercase word: "ανεπηρέαστοΣ"
        # (lowercase Greek + uppercase Greek at word boundary)
        def fix_endcap(m):
            nonlocal fixes_made
            fixes_made += 1
            return m.group(1) + m.group(2).lower()
        
        t = re.sub(
            f'([{GR_LOWER}])([{GR_UPPER}])(?=\\s|$|[.,;:!?«»\\-–—])',
            fix_endcap,
            t
        )
        
        result.append(t)
    
    if fixes_made > 0:
        print(f"[POST-PROCESS] Fixed {fixes_made} stray mid-word capitals")
    
    return result

TAG_B = ("[B]", "[/B]")
TAG_I = ("[I]", "[/I]")
TAG_U = ("[U]", "[/U]")

def runs_to_tagged_text(paragraph) -> str:
    # First, merge adjacent runs with identical formatting
    # This prevents tags from splitting words (e.g., [B]ανεπηρέαστ[/B][B]Ος[/B])
    merged_runs = []
    for run in paragraph.runs:
        t = run.text
        if not t:
            continue
        fmt = (bool(run.bold), bool(run.italic), bool(run.underline))
        if merged_runs and merged_runs[-1][1] == fmt:
            merged_runs[-1] = (merged_runs[-1][0] + t, fmt)
        else:
            merged_runs.append((t, fmt))
    
    out = []
    for text, (bold, italic, underline) in merged_runs:
        t = text
        if bold:
            t = f"{TAG_B[0]}{t}{TAG_B[1]}"
        if italic:
            t = f"{TAG_I[0]}{t}{TAG_I[1]}"
        if underline:
            t = f"{TAG_U[0]}{t}{TAG_U[1]}"
        out.append(t)
    return "".join(out).strip()

def tagged_text_to_runs(paragraph, tagged: str):
    for r in paragraph.runs:
        r._r.getparent().remove(r._r)
    pattern = re.compile(r"(\[/?[IBU]\])")
    tokens = [tok for tok in pattern.split(tagged) if tok]
    bold = italic = underline = False
    for tok in tokens:
        if tok == TAG_B[0]: bold = True
        elif tok == TAG_B[1]: bold = False
        elif tok == TAG_I[0]: italic = True
        elif tok == TAG_I[1]: italic = False
        elif tok == TAG_U[0]: underline = True
        elif tok == TAG_U[1]: underline = False
        else:
            if tok:
                run = paragraph.add_run(tok)
                run.bold = bold
                run.italic = italic
                run.underline = underline

def extract_docx_paragraphs(doc_path: Path) -> List[Dict]:
    doc = docx.Document(str(doc_path))
    result = []
    for idx, p in enumerate(doc.paragraphs):
        tagged = runs_to_tagged_text(p)
        plain = p.text.strip()
        if plain:
            result.append({"index": idx, "tagged": tagged, "plain": plain})
    return result

def extract_text_paragraphs(text: str) -> List[Dict]:
    paragraphs = split_paragraphs(text)
    return [{"index": i, "tagged": p, "plain": p} for i, p in enumerate(paragraphs)]


# ============== LLM CALLS ==============

async def call_llm(
    system_prompt: str,
    user_prompt: str,
    model_id: str,
    api_key: str,
    temperature: float = 0.2,
) -> str:
    client = OpenAI(base_url="https://openrouter.ai/api/v1", api_key=api_key)
    
    max_tokens = 16384
    if "gemini-3" in model_id:
        max_tokens = 65536
    
    for attempt in range(MAX_RETRIES):
        try:
            response = await asyncio.to_thread(
                client.chat.completions.create,
                model=model_id,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ],
                temperature=temperature,
                max_tokens=max_tokens,
            )
            if not response or not response.choices:
                raise RuntimeError("Empty response from LLM")
            content = response.choices[0].message.content
            if content is None:
                raise RuntimeError("LLM returned None content (possible rate limit or content filter)")
            return content.strip()
        except Exception as e:
            print(f"[LLM] Attempt {attempt+1}/{MAX_RETRIES} failed: {e}")
            if attempt < MAX_RETRIES - 1:
                await asyncio.sleep(3 * (2 ** attempt))
                continue
            raise e


# ============== TRANSLATION SESSION ==============

class TranslationSession:
    def __init__(self, session_id: str):
        self.session_id = session_id
        self.session_dir = TRANSLATION_DIR / session_id
        self.session_dir.mkdir(exist_ok=True)
        
        self.original_path = self.session_dir / "original.docx"
        self.output_path = self.session_dir / "translated.docx"
        self.progress_path = self.session_dir / "progress.json"
        self.config_path = self.session_dir / "config.json"
        
        self.status = "idle"
        self.progress = 0
        self.total_chunks = 0
        self.source_language = ""
        self.document_brief = ""
        self.style_register = ""
        self.glossary: List[Dict] = []
        self.paragraphs: List[Dict] = []
        self.translated: List[str] = []
        self.error_message = ""
        self.filename = ""
        self.is_docx = False
        self.chunk_reviews: List[Dict] = []  # Per-chunk critique results
    
    def save_progress(self):
        data = {
            "status": self.status,
            "progress": self.progress,
            "total_chunks": self.total_chunks,
            "source_language": self.source_language,
            "document_brief": self.document_brief,
            "style_register": self.style_register,
            "glossary": self.glossary,
            "translated": self.translated,
            "filename": self.filename,
            "is_docx": self.is_docx,
            "error_message": self.error_message,
            "chunk_reviews": self.chunk_reviews,
        }
        with open(self.progress_path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    
    def load_progress(self) -> bool:
        if not self.progress_path.exists():
            return False
        try:
            with open(self.progress_path, "r", encoding="utf-8") as f:
                data = json.load(f)
            self.status = data.get("status", "idle")
            self.progress = data.get("progress", 0)
            self.total_chunks = data.get("total_chunks", 0)
            self.source_language = data.get("source_language", "")
            self.document_brief = data.get("document_brief", "")
            self.style_register = data.get("style_register", "")
            self.glossary = data.get("glossary", [])
            self.translated = data.get("translated", [])
            self.filename = data.get("filename", "")
            self.is_docx = data.get("is_docx", False)
            self.error_message = data.get("error_message", "")
            self.chunk_reviews = data.get("chunk_reviews", [])
            
            # Re-extract paragraphs from original file (not saved in progress)
            if self.is_docx and self.original_path.exists():
                self.paragraphs = extract_docx_paragraphs(self.original_path)
            elif not self.is_docx:
                txt_path = self.session_dir / "original.txt"
                if txt_path.exists():
                    text = open(txt_path, "r", encoding="utf-8").read()
                    self.paragraphs = extract_text_paragraphs(text)
                    self.original_path = txt_path
            
            return True
        except Exception:
            return False
    
    def save_config(self, config: Dict):
        with open(self.config_path, "w", encoding="utf-8") as f:
            json.dump(config, f, ensure_ascii=False, indent=2)
    
    def get_review_queue(self) -> Dict:
        """Generate review queue — only flagged segments for human review."""
        flagged = []
        stats = {"total_chunks": len(self.chunk_reviews), "clean": 0, "fixed": 0, "review_required": 0,
                 "issues_by_type": {}, "issues_by_severity": {"low": 0, "medium": 0, "high": 0}}
        
        for review in self.chunk_reviews:
            if not review:
                continue
            issues = review.get("issues", [])
            if review.get("review_required"):
                stats["review_required"] += 1
                for issue in issues:
                    issue_type = issue.get("type", "unknown")
                    severity = issue.get("severity", "low")
                    stats["issues_by_type"][issue_type] = stats["issues_by_type"].get(issue_type, 0) + 1
                    stats["issues_by_severity"][severity] = stats["issues_by_severity"].get(severity, 0) + 1
                
                # Get paragraph texts for context
                para_indices = review.get("para_indices", [])
                source_texts = [self.paragraphs[i]["plain"] for i in para_indices if i < len(self.paragraphs)]
                translated_texts = [self.translated[i] for i in para_indices if i < len(self.translated)]
                
                flagged.append({
                    "chunk_idx": review.get("chunk_idx"),
                    "confidence": review.get("chunk_confidence"),
                    "issues": issues,
                    "source_preview": "\n".join(source_texts)[:500],
                    "translation_preview": "\n".join(translated_texts)[:500],
                })
            elif issues:
                stats["fixed"] += 1
            else:
                stats["clean"] += 1
        
        return {
            "flagged_chunks": flagged,
            "stats": stats,
            "review_percentage": f"{(stats['review_required'] / max(1, stats['total_chunks'])) * 100:.1f}%",
        }
    
    async def prescan(self, api_key: str, model_id: str):
        self.status = "scanning"
        self.save_progress()
        
        all_text = "\n\n".join(p["plain"] for p in self.paragraphs)
        
        if len(all_text) > 200000:
            sample = all_text[:50000] + "\n\n[...]\n\n" + all_text[len(all_text)//2 - 10000:len(all_text)//2 + 10000] + "\n\n[...]\n\n" + all_text[-20000:]
        else:
            sample = all_text
        
        try:
            # ---- Pass 1: Document analysis + initial glossary ----
            print(f"[PRESCAN] Pass 1: Document analysis ({len(sample):,} chars)...")
            result = await call_llm(
                system_prompt="Αναλύεις κείμενα για μεταφραστικό γραφείο. Απαντάς ΜΟΝΟ σε JSON.",
                user_prompt=make_prescan_prompt(sample),
                model_id=model_id,
                api_key=api_key,
                temperature=0.1,
            )
            
            result = result.strip()
            if result.startswith("```"):
                result = re.sub(r"```json?\s*", "", result)
                result = result.rstrip("`").strip()
            
            data = json.loads(result)
            
            self.source_language = data.get("source_language", "Unknown")
            initial_glossary = data.get("suggested_glossary", [])
            self.style_register = data.get("style_register", "ακαδημαϊκό")
            
            self.document_brief = (
                f"Τύπος: {data.get('document_type', 'N/A')}\n"
                f"Θέμα: {data.get('subject', 'N/A')}\n"
                f"Εποχή/Πλαίσιο: {data.get('era_context', 'N/A')}\n"
                f"Ύφος: {data.get('style_register', 'N/A')}\n"
                f"Σημειώσεις: {data.get('translation_notes', 'N/A')}"
            )
            
            print(f"[PRESCAN] Pass 1 done: {self.source_language}, {len(initial_glossary)} initial terms")
            
            # ---- Pass 2: Dedicated terminology extraction ----
            print(f"[PRESCAN] Pass 2: Terminology extraction...")
            await asyncio.sleep(PAUSE_BETWEEN_CALLS)
            
            term_prompt = f"""Εξαγωγή ΟΛΩΝ των ονομάτων και επαναλαμβανόμενων όρων από το παρακάτω κείμενο ({self.source_language}).

Θέλω ΠΛΗΡΗ ΛΙΣΤΑ με:
1. ΟΛΟΥΣ τους ανθρώπους (πλήρη ονόματα)
2. ΟΛΟΥΣ τους τόπους (πόλεις, χώρες, κτίρια)
3. ΟΛΟΥΣ τους οργανισμούς (εταιρείες, ιδρύματα)
4. ΟΛΟΥΣ τους τίτλους έργων (βιβλία, τραγούδια, ταινίες)
5. ΟΛΟΥΣ τους τεχνικούς/ειδικούς όρους που εμφανίζονται 2+ φορές

Για κάθε όρο δώσε την ελληνική απόδοση/μεταγραφή.

Απάντησε ΜΟΝΟ σε JSON array:
[
  {{"source": "...", "target": "...", "note": "person/place/org/title/term"}}
]

Χωρίς backticks, χωρίς σχόλια, ΜΟΝΟ JSON array.

ΚΕΙΜΕΝΟ:
{sample[:80000]}"""

            try:
                terms_result = await call_llm(
                    system_prompt="Είσαι ειδικός στην εξαγωγή ορολογίας. Απαντάς ΜΟΝΟ σε JSON array.",
                    user_prompt=term_prompt,
                    model_id=model_id,
                    api_key=api_key,
                    temperature=0.1,
                )
                
                terms_result = terms_result.strip()
                if terms_result.startswith("```"):
                    terms_result = re.sub(r"```json?\s*", "", terms_result)
                    terms_result = terms_result.rstrip("`").strip()
                
                extracted_terms = json.loads(terms_result)
                if not isinstance(extracted_terms, list):
                    extracted_terms = []
                print(f"[PRESCAN] Pass 2 done: {len(extracted_terms)} terms extracted")
            except Exception as e:
                print(f"[PRESCAN] Pass 2 failed (non-fatal): {e}")
                extracted_terms = []
            
            # ---- Merge: deduplicate ----
            seen = set()
            merged = []
            for entry in initial_glossary:
                key = entry.get("source", "").lower().strip()
                if key and key not in seen:
                    seen.add(key)
                    merged.append(entry)
            for entry in extracted_terms:
                key = entry.get("source", "").lower().strip()
                if key and key not in seen:
                    seen.add(key)
                    merged.append(entry)
            
            self.glossary = merged
            
            # Filter glossary for literary texts: only enforce names/places/orgs/titles
            # General vocabulary terms can cause rigid replacements in literary prose
            style_lower = (self.style_register or "").lower()
            is_literary = any(kw in style_lower for kw in ["λογοτεχν", "μυθιστόρ", "διήγημ", "αφηγημ", "literary"])
            if is_literary:
                SAFE_NOTES = {"person", "place", "org", "title", "organization"}
                filtered = []
                for entry in merged:
                    note = (entry.get("note") or "").strip().lower()
                    if note in SAFE_NOTES:
                        filtered.append(entry)
                print(f"[PRESCAN] Literary mode: glossary filtered {len(merged)} → {len(filtered)} (names/places/titles only)")
                self.glossary = filtered
            
            print(f"[PRESCAN] Final glossary: {len(self.glossary)} terms")
            
            self.status = "scanned"
            self.save_progress()
            
            return {
                "source_language": self.source_language,
                "document_brief": self.document_brief,
                "glossary": self.glossary,
                "key_entities": data.get("key_entities", []),
                "translation_notes": data.get("translation_notes", ""),
                "paragraph_count": len(self.paragraphs),
                "char_count": len(all_text),
            }
            
        except Exception as e:
            self.status = "error"
            self.error_message = str(e)
            self.save_progress()
            raise
    
    async def translate(self, api_key: str, model_id: str):
        self.status = "translating"
        
        para_texts = [p["tagged"] for p in self.paragraphs]
        chunks = chunk_paragraphs(para_texts)
        self.total_chunks = len(chunks)
        
        if len(self.translated) != len(self.paragraphs):
            self.translated = [""] * len(self.paragraphs)
        
        self.save_progress()
        
        style = self.style_register if self.style_register else "ακαδημαϊκό"
        system = make_system_prompt(self.source_language, self.document_brief, self.glossary, style)
        
        try:
            for chunk_idx, para_indices in enumerate(chunks):
                if chunk_idx < self.progress:
                    continue
                
                self.progress = chunk_idx
                self.save_progress()
                
                chunk_content = "\n\n".join(para_texts[i] for i in para_indices)
                
                # Build context window
                context = ""
                if chunk_idx > 0:
                    ctx_start = max(0, chunk_idx - CONTEXT_CHUNKS)
                    ctx_indices = []
                    for ci in range(ctx_start, chunk_idx):
                        ctx_indices.extend(chunks[ci])
                    context = "\n\n".join(
                        self.translated[i] for i in ctx_indices[-10:]
                        if self.translated[i]
                    )
                
                # Per-chunk retry (up to 3 attempts)
                chunk_success = False
                final_text = ""
                chunk_review = {"chunk_idx": chunk_idx, "issues": [], "chunk_confidence": "high", "review_required": False}
                for chunk_attempt in range(3):
                    try:
                        # Pass 1: Translate
                        print(f"[TRANSLATE] Chunk {chunk_idx+1}/{len(chunks)} — translating (attempt {chunk_attempt+1})...")
                        translated = await call_llm(
                            system_prompt=system,
                            user_prompt=make_translation_prompt(chunk_content, context),
                            model_id=model_id,
                            api_key=api_key,
                        )
                        await asyncio.sleep(PAUSE_BETWEEN_CALLS)
                        
                        # Pass 2: Critique (JSON issues)
                        print(f"[TRANSLATE] Chunk {chunk_idx+1}/{len(chunks)} — critiquing...")
                        critique_raw = await call_llm(
                            system_prompt="Είσαι κριτικός μεταφραστικής ποιότητας. Απαντάς ΜΟΝΟ σε JSON.",
                            user_prompt=make_critique_prompt(chunk_content, translated, self.source_language),
                            model_id=model_id,
                            api_key=api_key,
                            temperature=0.1,
                        )
                        await asyncio.sleep(PAUSE_BETWEEN_CALLS)
                        
                        # Parse critique
                        critique_data = {"issues": [], "chunk_confidence": "high"}
                        critique_parse_failed = False
                        try:
                            critique_clean = critique_raw.strip()
                            if critique_clean.startswith("```"):
                                critique_clean = re.sub(r"```json?\s*", "", critique_clean)
                                critique_clean = critique_clean.rstrip("`").strip()
                            critique_data = json.loads(critique_clean)
                        except (json.JSONDecodeError, Exception) as parse_err:
                            print(f"[TRANSLATE] Critique parse failed — marking chunk as untrusted: {parse_err}")
                            critique_data = {"issues": [], "chunk_confidence": "low"}
                            critique_parse_failed = True
                        
                        issues = critique_data.get("issues", [])
                        confidence = critique_data.get("chunk_confidence", "high")
                        high_severity = [i for i in issues if i.get("severity") == "high"]
                        
                        # Store review metadata
                        risky_types = {"allusion", "irony", "ambiguity", "cultural", "factual"}
                        has_risky = any(i.get("risk_type") in risky_types for i in issues)
                        
                        chunk_review = {
                            "chunk_idx": chunk_idx,
                            "para_indices": list(para_indices),
                            "issues": issues,
                            "issue_count": len(issues),
                            "high_severity_count": len(high_severity),
                            "chunk_confidence": confidence,
                            "critique_parse_failed": critique_parse_failed,
                            "review_required": (
                                confidence == "low" or 
                                len(high_severity) > 0 or 
                                critique_parse_failed or
                                has_risky or
                                len([i for i in issues if i.get("severity") in ("medium", "high")]) >= 3
                            ),
                        }
                        
                        if issues:
                            # Pass 3: Apply fixes (only if issues found)
                            print(f"[TRANSLATE] Chunk {chunk_idx+1}/{len(chunks)} — applying {len(issues)} fixes ({len(high_severity)} high)...")
                            issues_json = json.dumps(issues, ensure_ascii=False, indent=2)
                            final_text = await call_llm(
                                system_prompt=system,
                                user_prompt=make_apply_fixes_prompt(chunk_content, translated, issues_json, self.source_language),
                                model_id=model_id,
                                api_key=api_key,
                            )
                            await asyncio.sleep(PAUSE_BETWEEN_CALLS)
                        else:
                            # No issues — keep original translation (save 1 LLM call)
                            print(f"[TRANSLATE] Chunk {chunk_idx+1}/{len(chunks)} — clean, no fixes needed")
                            final_text = translated
                        
                        chunk_success = True
                        break
                    except Exception as chunk_err:
                        print(f"[TRANSLATE] Chunk {chunk_idx+1} attempt {chunk_attempt+1} failed: {chunk_err}")
                        if chunk_attempt < 2:
                            wait_time = 10 * (chunk_attempt + 1)
                            print(f"[TRANSLATE] Retrying in {wait_time}s...")
                            await asyncio.sleep(wait_time)
                        else:
                            raise RuntimeError(f"Chunk {chunk_idx+1} failed after 3 attempts: {chunk_err}")
                
                if not chunk_success:
                    continue
                
                # Save review data
                while len(self.chunk_reviews) <= chunk_idx:
                    self.chunk_reviews.append({})
                self.chunk_reviews[chunk_idx] = chunk_review
                
                # Parse into paragraphs
                out_paras = split_paragraphs(final_text)
                
                # Paragraph count guard
                expected = len(para_indices)
                if len(out_paras) != expected:
                    print(f"[TRANSLATE] Paragraph mismatch: expected {expected}, got {len(out_paras)}. Attempting repair...")
                    repair_system = system + "\nΕΠΙΠΛΕΟΝ: Η έξοδος ΠΡΕΠΕΙ να έχει ΑΚΡΙΒΩΣ τον ίδιο αριθμό παραγράφων με το πρωτότυπο."
                    repair_user = f"""Θέλω ΑΚΡΙΒΩΣ {expected} παραγράφους.
ΠΡΩΤΟΤΥΠΟ ({expected} παράγραφοι):
<<<
{chunk_content}
>>>
ΤΡΕΧΟΥΣΑ ΜΕΤΑΦΡΑΣΗ ({len(out_paras)} παράγραφοι):
<<<
{final_text}
>>>
Δώσε ΜΟΝΟ τη σωστή μετάφραση σε {expected} παραγράφους, χωρίς σχόλια."""
                    repaired = await call_llm(
                        system_prompt=repair_system,
                        user_prompt=repair_user,
                        model_id=model_id,
                        api_key=api_key,
                    )
                    out_paras = split_paragraphs(repaired)
                    await asyncio.sleep(PAUSE_BETWEEN_CALLS)
                
                # Assign to translated array
                if len(out_paras) == expected:
                    for j, pi in enumerate(para_indices):
                        self.translated[pi] = out_paras[j]
                else:
                    print(f"[TRANSLATE] WARNING: Still mismatched ({len(out_paras)} vs {expected})")
                    for j, pi in enumerate(para_indices):
                        if j < len(out_paras):
                            self.translated[pi] = out_paras[j]
                        else:
                            self.translated[pi] = ""
                
                self.progress = chunk_idx + 1
                self.save_progress()
                print(f"[TRANSLATE] Chunk {chunk_idx+1}/{len(chunks)} — done")
            
            # ===== RESCUE PASS: Re-translate any empty paragraphs =====
            untranslated = [i for i, t in enumerate(self.translated) if not t.strip()]
            if untranslated:
                print(f"[RESCUE] {len(untranslated)} paragraphs still empty after main pass. Re-translating individually...")
                
                # Group consecutive untranslated paragraphs into mini-chunks
                rescue_chunks = []
                current_group = [untranslated[0]]
                for idx in untranslated[1:]:
                    if idx == current_group[-1] + 1:
                        current_group.append(idx)
                    else:
                        rescue_chunks.append(current_group)
                        current_group = [idx]
                rescue_chunks.append(current_group)
                
                for rc_idx, rc_indices in enumerate(rescue_chunks):
                    rc_content = "\n\n".join(para_texts[i] for i in rc_indices)
                    
                    # Build context from surrounding translated paragraphs
                    ctx_before = ""
                    for look_back in range(rc_indices[0] - 1, max(rc_indices[0] - 4, -1), -1):
                        if look_back >= 0 and self.translated[look_back]:
                            ctx_before = self.translated[look_back]
                            break
                    
                    rescue_success = False
                    for rescue_attempt in range(3):
                        try:
                            print(f"[RESCUE] Group {rc_idx+1}/{len(rescue_chunks)} ({len(rc_indices)} paras, attempt {rescue_attempt+1})...")
                            ctx_section = (
                                f"\nΠΡΟΗΓΟΥΜΕΝΟ ΠΛΑΙΣΙΟ (για συνέχεια ύφους — ΜΗΝ το μεταφράσεις ξανά):\n<<<\n{ctx_before[-1500:]}\n>>>\n"
                                if ctx_before else ""
                            )
                            rescue_prompt = f"""Μετάφρασε ΑΚΡΙΒΩΣ {len(rc_indices)} παραγράφους.
Δώσε ΜΟΝΟ τη μετάφραση σε ΑΚΡΙΒΩΣ {len(rc_indices)} παραγράφους (χωρισμένες με κενή γραμμή).
{ctx_section}ΚΕΙΜΕΝΟ ΠΡΟΣ ΜΕΤΑΦΡΑΣΗ:
<<<
{rc_content}
>>>"""
                            rescued = await call_llm(
                                system_prompt=system,
                                user_prompt=rescue_prompt,
                                model_id=model_id,
                                api_key=api_key,
                            )
                            await asyncio.sleep(PAUSE_BETWEEN_CALLS)
                            
                            rescued_paras = split_paragraphs(rescued)
                            
                            if len(rescued_paras) == len(rc_indices):
                                for j, pi in enumerate(rc_indices):
                                    self.translated[pi] = rescued_paras[j]
                                rescue_success = True
                                break
                            elif len(rescued_paras) >= 1:
                                # Partial success — assign what we can
                                for j, pi in enumerate(rc_indices):
                                    if j < len(rescued_paras):
                                        self.translated[pi] = rescued_paras[j]
                                # If we got more paras than expected, merge excess into last
                                if len(rescued_paras) > len(rc_indices):
                                    last_pi = rc_indices[-1]
                                    overflow = rescued_paras[len(rc_indices):]
                                    self.translated[last_pi] += "\n" + "\n".join(overflow)
                                rescue_success = True
                                break
                        except Exception as rescue_err:
                            print(f"[RESCUE] Attempt {rescue_attempt+1} failed: {rescue_err}")
                            if rescue_attempt < 2:
                                await asyncio.sleep(5 * (rescue_attempt + 1))
                    
                    if rescue_success:
                        self.save_progress()
                        print(f"[RESCUE] Group {rc_idx+1} rescued successfully")
                    else:
                        print(f"[RESCUE] WARNING: Group {rc_idx+1} could not be rescued")
                
                # Final count
                still_empty = sum(1 for t in self.translated if not t.strip())
                if still_empty:
                    print(f"[RESCUE] WARNING: {still_empty} paragraphs still empty after rescue")
                else:
                    print(f"[RESCUE] All paragraphs translated successfully")
            
            # ===== POST-PROCESSING =====
            print(f"[POST-PROCESS] Enforcing glossary ({len(self.glossary)} terms)...")
            self.translated = enforce_glossary(self.translated, self.glossary)
            
            # Conditional quote normalization — only if non-Greek quotes detected
            has_bad_quotes = any(
                '"' in t or '"' in t or '"' in t or "'" in t or "'" in t
                for t in self.translated if t
            )
            if has_bad_quotes:
                print(f"[POST-PROCESS] Non-Greek quotes detected, fixing...")
                self.translated = fix_quotation_marks(self.translated)
            else:
                print(f"[POST-PROCESS] Quotes OK, skipping normalization")
            
            print(f"[POST-PROCESS] Fixing stray mid-word capitals...")
            self.translated = fix_midword_capitals(self.translated)
            
            # Check for untranslated paragraphs
            untranslated = [i for i, t in enumerate(self.translated) if not t]
            if untranslated:
                print(f"[TRANSLATE] WARNING: {len(untranslated)} paragraphs untranslated: {untranslated[:10]}...")
            
            self.status = "ready"
            self.save_progress()
            
        except Exception as e:
            self.status = "error"
            self.error_message = str(e)
            self.save_progress()
            raise
    
    async def consistency_check(self, api_key: str, model_id: str) -> Dict:
        if not self.glossary or not self.translated:
            return {"issues": [], "message": "No glossary or translation to check"}
        
        full_translation = "\n\n".join(t for t in self.translated if t)
        issues = []
        
        for entry in self.glossary:
            target = entry.get("target", "")
            source = entry.get("source", "")
            if not source or not target:
                continue
            
            occurrences = set()
            for match in re.finditer(re.escape(target), full_translation, re.IGNORECASE):
                occurrences.add(match.group())
            
            if len(occurrences) > 1:
                issues.append({
                    "source": source,
                    "variants": list(occurrences),
                    "expected": target,
                })
        
        return {
            "issues": issues,
            "total_glossary_terms": len(self.glossary),
            "inconsistent_terms": len(issues),
        }
    
    def generate_output(self) -> Path:
        if self.is_docx:
            return self._generate_docx()
        else:
            return self._generate_txt()
    
    def _generate_docx(self) -> Path:
        if not docx:
            raise RuntimeError("python-docx not installed")
        
        # Re-extract paragraphs if needed
        if not self.paragraphs and self.original_path.exists():
            self.paragraphs = extract_docx_paragraphs(self.original_path)
        
        shutil.copy(self.original_path, self.output_path)
        doc = docx.Document(str(self.output_path))
        
        for i, para_data in enumerate(self.paragraphs):
            if i >= len(self.translated) or not self.translated[i]:
                continue
            original_idx = para_data["index"]
            if original_idx >= len(doc.paragraphs):
                continue
            tagged_text_to_runs(doc.paragraphs[original_idx], self.translated[i])
        
        doc.save(str(self.output_path))
        return self.output_path
    
    def _generate_txt(self) -> Path:
        txt_path = self.session_dir / "translated.txt"
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write("\n\n".join(t if t else "" for t in self.translated))
        return txt_path


# ============== SESSION STORAGE ==============

translation_sessions: Dict[str, TranslationSession] = {}

def create_session(session_id: str = None) -> TranslationSession:
    if not session_id:
        session_id = str(uuid.uuid4())[:8]
    session = TranslationSession(session_id)
    translation_sessions[session_id] = session
    return session

def get_session(session_id: str) -> Optional[TranslationSession]:
    if session_id in translation_sessions:
        return translation_sessions[session_id]
    session = TranslationSession(session_id)
    if session.load_progress():
        translation_sessions[session_id] = session
        return session
    return None
